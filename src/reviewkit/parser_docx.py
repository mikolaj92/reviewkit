"""DOCX parser that builds the internal review hierarchy."""

from __future__ import annotations

import itertools
import re
from collections import defaultdict
from collections.abc import Iterator
from dataclasses import dataclass
from pathlib import Path
from typing import assert_never
from zipfile import BadZipFile, ZipFile

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxtor import AddressableSpan, DocxDocument as AddressableDocxDocument
from lxml import etree

from reviewkit.comments import (
    DocxComment,
    _comment_markers_are_complete,
    _comment_thread_ids_are_complete,
    comments_for_locator,
    read_comments,
)
from reviewkit.document import ParagraphNode, ReviewDocument, SectionNode, SentenceNode
from reviewkit.markup_purity import MarkupReport, has_tracked_revisions, inspect_markup
from reviewkit.models import (
    RevisionCoverageState,
    RevisionLedger,
    SourceRevision,
    SourceRevisionKind,
)

# Terminators span Latin (. ! ?) and common non-Latin sentence enders so the sentence
# tier does not silently disappear for non-Latin-script documents: CJK (。！？), the
# horizontal ellipsis (…), the Arabic question mark (؟) and the Devanagari danda (।).
_SENTENCE_PUNCT_RE = re.compile(r"[.!?。！？…؟।]+", re.UNICODE)
# The non-Latin enders above are unambiguous, script-specific sentence terminators with
# no abbreviation/decimal role, and CJK writes no space between sentences - so they end a
# sentence regardless of what follows, unlike the whitespace-gated Latin punctuation.
_STRONG_TERMINATORS = frozenset("。！？…؟।")
_TRAILING_WORD_RE = re.compile(r"(\w+)$", re.UNICODE)


def load_docx(path: str | Path) -> ReviewDocument:
    source_path = Path(path)
    docx = DocxDocument(str(source_path))
    comments = read_comments(source_path)
    addressable_document = AddressableDocxDocument.open(source_path)
    effective_texts, revision_ledger = _project_revision_input(addressable_document.spans)
    markup_report = inspect_markup(source_path)
    tracked_revisions = has_tracked_revisions(source_path)
    if (
        _revision_coverage_is_incomplete(source_path, markup_report, revision_ledger)
        or _comment_ids_are_ambiguous(comments)
        or not _comment_markers_are_complete(source_path, comments)
        or not _comment_thread_ids_are_complete(source_path)
        or any(_comment_anchor_is_unresolved(comment, comments) for comment in comments)
    ):
        revision_ledger = revision_ledger.model_copy(
            update={"coverage": RevisionCoverageState.INCOMPLETE}
        )

    # Section/paragraph id counters are shared across the body walk and the synthetic
    # header/footer sections so every node keeps a globally unique id. "s1" is reserved
    # for the implicit leading body section, so section numbering starts at 2.
    section_ids = itertools.count(2)
    paragraph_ids = itertools.count(1)

    sections: list[SectionNode] = []
    current = SectionNode(id="s1")

    # Walk the body in true document order so a table interleaves with the paragraphs
    # around it and lands under its authoring heading, instead of every table being
    # appended to whatever section happened to be open at the end of the body.
    for docx_paragraph, locator, source in _iter_body_sources(docx):
        text = effective_texts.get(locator, str(getattr(docx_paragraph, "text", ""))).strip()
        if not text:
            continue

        if _is_heading(docx_paragraph):
            if current.title or current.paragraphs:
                sections.append(current)
                current = SectionNode(
                    id=f"s{next(section_ids)}",
                    title=text,
                    locator=locator,
                    metadata={"source": source},
                )
            else:
                current = SectionNode(
                    id=current.id,
                    title=text,
                    locator=locator,
                    metadata={"source": source},
                )
            continue

        current.paragraphs.append(
            _paragraph_node(
                f"p{next(paragraph_ids)}",
                text,
                current.id,
                locator,
                source,
                _opaque_ranges(docx_paragraph),
                comments_for_locator(comments, locator),
            )
        )

    if current.title or current.paragraphs or not sections:
        sections.append(current)

    # Header/footer paragraphs get their own synthetic sections keyed by source so they
    # are not misread as body prose tacked onto the trailing body section. Locator strings
    # ("header:S:p:P"/"footer:S:p:P") are unchanged, so rendering resolves them identically.
    sections.extend(
        _header_footer_sections(docx, section_ids, paragraph_ids, comments, effective_texts)
    )

    metadata = {
        "paragraph_count": str(sum(len(section.paragraphs) for section in sections)),
        "table_count": str(len(docx.tables)),
        "comment_count": str(len(comments)),
        "tracked_revisions_detected": str(tracked_revisions).lower(),
    }
    return ReviewDocument(
        source_path=source_path,
        sections=sections,
        metadata=metadata,
        comments=comments,
        revision_ledger=revision_ledger,
    )


def _project_revision_input(
    spans: tuple[AddressableSpan, ...],
) -> tuple[dict[str, str], RevisionLedger]:
    effective_parts: dict[str, list[str]] = {}
    entries: list[SourceRevision] = []
    coverage = RevisionCoverageState.COMPLETE
    for span in spans:
        locator = _reviewkit_locator(span.container_id)
        match span.role:
            case "insertion":
                effective_parts.setdefault(locator, []).append(span.text)
                entries.append(_source_revision(span, locator, SourceRevisionKind.INSERTED))
            case "deletion":
                entries.append(_source_revision(span, locator, SourceRevisionKind.DELETED))
            case "run":
                effective_parts.setdefault(locator, []).append(span.text)
            case "hyperlink":
                effective_parts.setdefault(locator, []).append(span.text)
                if span.revision_id is not None:
                    coverage = RevisionCoverageState.INCOMPLETE
            case unexpected:
                assert_never(unexpected)
    return (
        {locator: "".join(parts) for locator, parts in effective_parts.items()},
        RevisionLedger(coverage=coverage, entries=tuple(entries)),
    )


_SUPPORTED_REVISION_KINDS = frozenset({"ins", "del"})
_BLOCK_REVISION_CHILDREN = frozenset({"p", "tbl", "tr", "tc"})


def _revision_coverage_is_incomplete(
    source_path: Path, markup_report: MarkupReport, ledger: RevisionLedger
) -> bool:
    """Return whether every text-bearing source revision is represented by the typed projection.

    Word commonly emits empty revision wrappers for paragraph/property bookkeeping and
    nests a deletion inside an insertion for replacement text. Those wrappers do not own
    independent text spans, so comparing every wrapper count with Docxtor's text spans
    falsely marks otherwise addressable documents as incomplete. Compare the direct text
    owned by each revision identity instead; nested revision text belongs to the inner
    identity and is deliberately excluded from its parent.
    """
    if not markup_report.has_tracked_revisions:
        return False
    if set(markup_report.revision_kinds) - _SUPPORTED_REVISION_KINDS:
        return True

    source_inventory: dict[tuple[str, str | None, str | None, str | None], str] = defaultdict(str)
    with ZipFile(source_path) as bundle:
        for name in bundle.namelist():
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            root = etree.fromstring(bundle.read(name))
            for element in root.iter():
                kind = etree.QName(element).localname
                if kind not in _SUPPORTED_REVISION_KINDS:
                    continue
                if any(
                    etree.QName(descendant).localname in _BLOCK_REVISION_CHILDREN
                    for descendant in element.iter()
                    if descendant is not element
                ):
                    return True
                direct_text = _direct_revision_text(element)
                if direct_text:
                    source_inventory[
                        (
                            kind,
                            element.get(f"{{{_W_NS}}}id"),
                            element.get(f"{{{_W_NS}}}author"),
                            element.get(f"{{{_W_NS}}}date"),
                        )
                    ] += direct_text

    ledger_inventory: dict[tuple[str, str | None, str | None, str | None], str] = defaultdict(str)
    for entry in ledger.entries:
        ledger_inventory[
            (
                "ins" if entry.kind is SourceRevisionKind.INSERTED else "del",
                entry.revision_id,
                entry.author,
                entry.date,
            )
        ] += entry.text
    return source_inventory != ledger_inventory


def _direct_revision_text(element: etree._Element) -> str:
    """Return text/control characters owned directly by one revision wrapper."""
    parts: list[str] = []
    for node in element.iter():
        local = etree.QName(node).localname
        if local not in {"t", "delText", "tab", "br", "cr"}:
            continue
        parent = node.getparent()
        nested = False
        while parent is not None and parent is not element:
            if etree.QName(parent).localname in _SUPPORTED_REVISION_KINDS:
                nested = True
                break
            parent = parent.getparent()
        if nested:
            continue
        if local in {"t", "delText"}:
            parts.append(node.text or "")
        elif local == "tab":
            parts.append("\t")
        else:
            parts.append("\n")
    return "".join(parts)


def _source_revision(
    span: AddressableSpan,
    locator: str,
    kind: SourceRevisionKind,
) -> SourceRevision:
    return SourceRevision(
        kind=kind,
        text=span.text,
        locator=locator,
        span_id=span.span_id,
        start_offset=span.start_offset,
        end_offset=span.end_offset,
        revision_id=span.revision_id,
        author=span.revision_author,
        date=span.revision_date,
    )


def _reviewkit_locator(container_id: str) -> str:
    parts = container_id.split(":")
    if len(parts) == 8 and parts[0] == "table" and parts[2] == "r" and parts[4] == "c":
        return f"table:{parts[1]}:row:{parts[3]}:cell:{parts[5]}:p:{parts[7]}"
    return container_id


def _comment_anchor_is_unresolved(comment: DocxComment, comments: list[DocxComment]) -> bool:
    """Return whether a source comment has no usable story anchor.

    Word replies normally have no range markers of their own. A reply is anchored through
    its parent comment when that parent has a stable locator; only an unanchored standalone
    comment (or a reply whose parent is missing/unanchored) makes revision coverage
    incomplete.
    """
    if comment.locator is not None:
        return False
    if comment.parent_id is None:
        return True
    parent = next((candidate for candidate in comments if candidate.id == comment.parent_id), None)
    return parent is None or parent.locator is None


def _comment_ids_are_ambiguous(comments: list[DocxComment]) -> bool:
    return len({comment.id for comment in comments}) != len(comments)


_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_FOOTNOTES_PART = "word/footnotes.xml"
# The two footnotes Word always stores alongside real notes: the horizontal rule that
# separates footnotes from body text and its continuation variant. They carry the
# separator glyph, never document prose, so they are not content footnotes.
_STRUCTURAL_FOOTNOTE_TYPES = frozenset({"separator", "continuationSeparator"})


@dataclass(frozen=True)
class DocxFootnote:
    """One content footnote read from a ``.docx`` package: its ``w:id`` and visible text."""

    id: str
    text: str


def read_footnotes(path: str | Path) -> list[DocxFootnote]:
    """Extract every content footnote's visible text from a ``.docx`` package, in order.

    python-docx models no footnotes, so this reads ``word/footnotes.xml`` directly -- the
    package-level read reviewkit owns so consumers never hand-roll OOXML. Text is assembled
    the same way the renderer reads visible text (``w:t`` verbatim, ``w:tab`` -> tab,
    ``w:br``/``w:cr`` -> newline). The structural separator footnotes Word stores next to
    real notes are skipped. Returns an empty list when the package carries no footnotes part
    (or cannot be opened as a zip).
    """
    try:
        with ZipFile(str(path)) as bundle:
            raw = bundle.read(_FOOTNOTES_PART)
    except (KeyError, OSError, BadZipFile):
        return []
    root = etree.fromstring(raw)
    footnotes: list[DocxFootnote] = []
    for element in root.findall(f"{{{_W_NS}}}footnote"):
        if element.get(f"{{{_W_NS}}}type") in _STRUCTURAL_FOOTNOTE_TYPES:
            continue
        note_id = element.get(f"{{{_W_NS}}}id")
        if note_id is None:
            continue
        footnotes.append(DocxFootnote(id=note_id, text=_footnote_visible_text(element)))
    return footnotes


def _footnote_visible_text(element: object) -> str:
    parts: list[str] = []
    for node in element.iter():  # type: ignore[attr-defined]
        tag = node.tag
        if tag == f"{{{_W_NS}}}t":
            parts.append(node.text or "")
        elif tag == f"{{{_W_NS}}}tab":
            parts.append("\t")
        elif tag in (f"{{{_W_NS}}}br", f"{{{_W_NS}}}cr"):
            parts.append("\n")
    return "".join(parts)


def split_sentences(text: str) -> list[str]:
    return [sentence for sentence, _start, _end in split_sentences_with_spans(text)]


def split_sentences_with_spans(text: str) -> list[tuple[str, int, int]]:
    """Split ``text`` into sentences, keeping each sentence's char span within ``text``.

    The returned offsets refer to the stripped sentence as it appears inside
    ``text`` so callers can rebase sentence-relative locators into paragraph
    coordinates.
    """

    spans: list[tuple[str, int, int]] = []
    segment_start = 0
    for match in _SENTENCE_PUNCT_RE.finditer(text):
        if not _is_sentence_boundary(text, match.start(), match.end()):
            continue
        _append_span(spans, text, segment_start, match.end())
        segment_start = match.end()
    _append_span(spans, text, segment_start, len(text))
    if spans:
        return spans
    stripped = text.strip()
    if not stripped:
        return []
    start = text.find(stripped)
    return [(stripped, start, start + len(stripped))]


def _append_span(
    spans: list[tuple[str, int, int]], text: str, start: int, end: int
) -> None:
    segment = text[start:end]
    stripped = segment.strip()
    if not stripped:
        return
    lead = len(segment) - len(segment.lstrip())
    span_start = start + lead
    spans.append((stripped, span_start, span_start + len(stripped)))


def _is_sentence_boundary(text: str, punct_start: int, punct_end: int) -> bool:
    """Decide whether the punctuation run at ``[punct_start:punct_end]`` ends a sentence.

    Language- and domain-neutral heuristics avoid the classic over-splits:
    - non-Latin terminators (``。！？…؟।``) are unambiguous sentence enders with no
      abbreviation role, and CJK writes no inter-sentence space, so a run containing one
      is always a boundary regardless of the following character;
    - otherwise (Latin ``.!?``) a boundary must be followed by whitespace or end-of-text,
      so ``3.14`` and the inner dots of ``o.o.`` are never boundaries;
    - a period preceded by a single-letter token is treated as an initial/abbreviation
      (``J. R. R.``, the trailing ``o.``);
    - a period followed by a lowercase word is treated as an abbreviation (``Sp. z``).
    ``!`` and ``?`` are always strong boundaries when followed by whitespace/end.
    """
    if any(char in _STRONG_TERMINATORS for char in text[punct_start:punct_end]):
        return True
    if punct_end < len(text) and not text[punct_end].isspace():
        return False
    if "!" in text[punct_start:punct_end] or "?" in text[punct_start:punct_end]:
        return True
    trailing = _TRAILING_WORD_RE.search(text[:punct_start])
    if trailing is not None and len(trailing.group(1)) == 1 and trailing.group(1).isalpha():
        return False
    following = _next_non_space_char(text, punct_end)
    if following and following.islower():
        return False
    return True


def _next_non_space_char(text: str, index: int) -> str:
    while index < len(text) and text[index].isspace():
        index += 1
    return text[index] if index < len(text) else ""


def _is_heading(docx_paragraph: object) -> bool:
    style = getattr(docx_paragraph, "style", None)
    if style is None:
        return False
    # ``style_id`` is the language-independent internal identifier Word assigns to
    # built-in styles ("Heading1", "Heading2", "Title", ...), so heading detection
    # stays domain- and language-agnostic regardless of the document's UI language.
    style_id = str(getattr(style, "style_id", "") or "")
    return style_id.startswith("Heading") or style_id == "Title"


def _paragraph_node(
    paragraph_id: str,
    text: str,
    section_id: str,
    locator: str,
    source: str,
    opaque_ranges: list[tuple[int, int]] | None = None,
    comments: list[DocxComment] | None = None,
) -> ParagraphNode:
    sentences = [
        SentenceNode(
            id=f"{paragraph_id}.s{index}",
            text=sentence,
            paragraph_id=paragraph_id,
            char_start=start,
            char_end=end,
            locator=f"{locator}:s:{index - 1}",
            metadata={"source": source},
        )
        for index, (sentence, start, end) in enumerate(split_sentences_with_spans(text), start=1)
    ]
    return ParagraphNode(
        id=paragraph_id,
        text=text,
        section_id=section_id,
        locator=locator,
        metadata={"source": source},
        sentences=sentences,
        opaque_ranges=opaque_ranges or [],
        comments=comments or [],
    )


def _opaque_ranges(docx_paragraph: object) -> list[tuple[int, int]]:
    """Spans of the paragraph's STRIPPED text contributed by non-editable content.

    ``Paragraph.text`` includes visible characters the renderers treat as opaque:
    tabs/breaks inside runs and the text of non-run inline children (hyperlinks,
    fields, ...). Locators use the stripped node text, so the returned coordinates
    are shifted by the leading whitespace and clipped to the stripped length.

    Fail-open on any mismatch with python-docx's notion of the paragraph text
    (unknown layouts): returning [] just skips the prepare-time demotion; the
    renderers' own integrity guards still fail closed at render time.
    """
    p_element = getattr(docx_paragraph, "_p", None)
    text = str(getattr(docx_paragraph, "text", ""))
    if p_element is None:
        return []

    parts: list[tuple[str, bool]] = []  # (visible chunk, editable?)
    for child in p_element:
        tag = child.tag
        if tag == qn("w:pPr"):
            continue
        if tag == qn("w:r"):
            for run_child in child:
                run_tag = run_child.tag
                if run_tag == qn("w:t"):
                    parts.append((run_child.text or "", True))
                elif run_tag == qn("w:tab"):
                    parts.append(("\t", False))
                elif run_tag in (qn("w:br"), qn("w:cr")):
                    parts.append(("\n", False))
            continue
        parts.append((_descendant_visible_text(child), False))

    if "".join(chunk for chunk, _editable in parts) != text:
        return []

    lead = len(text) - len(text.lstrip())
    stripped_length = len(text.strip())
    ranges: list[tuple[int, int]] = []
    offset = 0
    for chunk, editable in parts:
        if not editable and chunk:
            start = max(offset - lead, 0)
            end = min(offset + len(chunk) - lead, stripped_length)
            if start < end:
                ranges.append((start, end))
        offset += len(chunk)
    return ranges


def _descendant_visible_text(element: object) -> str:
    parts: list[str] = []
    for node in element.iter():  # type: ignore[attr-defined]
        if node.tag == qn("w:t"):
            parts.append(node.text or "")
        elif node.tag == qn("w:tab"):
            parts.append("\t")
        elif node.tag in (qn("w:br"), qn("w:cr")):
            parts.append("\n")
    return "".join(parts)


def _iter_body_sources(docx: object) -> Iterator[tuple[object, str, str]]:
    # ``iter_inner_content`` yields body ``<w:p>`` and ``<w:tbl>`` children in true document
    # order, so a table is emitted at its real position (between the paragraphs that surround
    # it) rather than after all paragraphs. Separate paragraph/table counters keep the emitted
    # locators identical to the previous scheme: ``paragraph_index`` matches ``docx.paragraphs``
    # (which excludes table paragraphs) and ``table_index`` matches ``docx.tables``.
    paragraph_index = 0
    table_index = 0
    for block in docx.iter_inner_content():  # type: ignore[attr-defined]
        if isinstance(block, Paragraph):
            yield block, f"body:p:{paragraph_index}", "body"
            paragraph_index += 1
        elif isinstance(block, Table):
            yield from _iter_table_sources(block, table_index)
            table_index += 1


def _iter_table_sources(table: Table, table_index: int) -> Iterator[tuple[object, str, str]]:
    # A merged cell is yielded by ``row.cells`` once per grid position it spans (across
    # columns AND rows), so walk each underlying ``<w:tc>`` exactly once at its first grid
    # position - otherwise merged cells (ubiquitous in forms and contracts) are reviewed
    # twice, edited twice, and inflate paragraph_count. The set holds the lxml element
    # proxies (not ``id()``, whose value is reused after GC) so identity is stable and
    # unique per physical cell.
    seen_cells: set[object] = set()
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            if cell._tc in seen_cells:
                continue
            seen_cells.add(cell._tc)
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                locator = f"table:{table_index}:row:{row_index}:cell:{cell_index}:p:{paragraph_index}"
                yield paragraph, locator, "table"


def _header_footer_sections(
    docx: object,
    section_ids: Iterator[int],
    paragraph_ids: Iterator[int],
    comments: list[DocxComment],
    effective_texts: dict[str, str],
) -> list[SectionNode]:
    grouped: dict[str, list[tuple[object, str]]] = {}
    for docx_paragraph, locator, source in _iter_header_footer_sources(docx):
        grouped.setdefault(source, []).append((docx_paragraph, locator))

    sections: list[SectionNode] = []
    for source, entries in grouped.items():
        non_empty = [
            (paragraph, locator)
            for paragraph, locator in entries
            if effective_texts.get(locator, str(getattr(paragraph, "text", ""))).strip()
        ]
        if not non_empty:
            continue
        section_id = f"s{next(section_ids)}"
        paragraphs = [
            _paragraph_node(
                f"p{next(paragraph_ids)}",
                effective_texts.get(locator, str(getattr(paragraph, "text", ""))).strip(),
                section_id,
                locator,
                source,
                _opaque_ranges(paragraph),
                comments_for_locator(comments, locator),
            )
            for paragraph, locator in non_empty
        ]
        sections.append(
            SectionNode(
                id=section_id,
                # No fabricated title: capitalizing the source ("Header"/"Footer") injected an
                # English word into the reviewable tree, which the LLM would see as document
                # prose -- a language leak in the language-blind core. The header/footer
                # distinction is preserved in metadata["source"] below.
                title=None,
                metadata={"source": source},
                paragraphs=paragraphs,
            )
        )
    return sections


def _iter_header_footer_sources(docx: object) -> Iterator[tuple[object, str, str]]:
    for section_index, section in enumerate(getattr(docx, "sections", [])):
        for paragraph_index, paragraph in enumerate(section.header.paragraphs):
            yield paragraph, f"header:{section_index}:p:{paragraph_index}", "header"
        for paragraph_index, paragraph in enumerate(section.footer.paragraphs):
            yield paragraph, f"footer:{section_index}:p:{paragraph_index}", "footer"
