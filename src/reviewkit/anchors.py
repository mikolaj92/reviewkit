"""Body-anchor grammar and body-only paragraph resolution for DOCX documents.

The supported insertion-anchor grammar is ``body:p:<digits>`` and
``body:p:last``. Numeric indices are body-local positions in
``document.paragraphs``; anything outside the body fails closed instead of
resolving into tables, headers, or footers. Insertion, validation, and any
caller-side planning should share these helpers so the grammar cannot drift
between pipeline stages.
"""

from __future__ import annotations

import re
from collections.abc import Sequence
from dataclasses import dataclass, field

from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph

ANCHOR_LAST = "body:p:last"

# Signature/closing-section markers are only meaningful in the document tail:
# only paragraphs within the last _SIGNATURE_TAIL_WINDOW non-empty body
# paragraphs may start a signature block.
_SIGNATURE_TAIL_WINDOW = 5


def _compile_signature_keywords(keywords: Sequence[str]) -> tuple[re.Pattern[str], ...]:
    patterns: list[re.Pattern[str]] = []
    for keyword in keywords:
        stem = keyword.endswith("*")
        token = (keyword[:-1] if stem else keyword).lower()
        if not token:
            continue  # "" / "*" would match every word: skip, fail closed
        escaped = re.escape(token)
        patterns.append(re.compile(rf"\b{escaped}\w*" if stem else rf"\b{escaped}\b"))
    return tuple(patterns)


@dataclass(frozen=True, eq=False)
class SignatureBlockStart:
    """Opaque handle to the first body element of the trailing signature block.

    Callers only get the paragraph's visible ``text``; the underlying OOXML
    element stays in-package so all element plumbing remains reviewkit's.
    Identity semantics (``eq=False``): two handles are equal only when they are
    the same handle.
    """

    _element: CT_P = field(repr=False)

    @property
    def text(self) -> str:
        """The visible text of the paragraph starting the signature block."""
        return self._element.text or ""


def find_signature_block_start(
    document: DocxDocument,
    ignore_texts: Sequence[str] = (),
    *,
    signature_keywords: Sequence[str] = (),
) -> SignatureBlockStart | None:
    """Return a handle to the start of the trailing signature block, or ``None``.

    ``signature_keywords`` use the lib-owned keyword grammar: each keyword is
    matched case-insensitively against paragraph text at a leading word
    boundary; a trailing ``*`` marks a stem (``\\b<kw>\\w*``, matching any
    inflection), otherwise the keyword matches as an exact word
    (``\\b<kw>\\b``, so e.g. ``date`` never matches inside ``update``).
    Keywords are ``re.escape``'d, so regex metacharacters are literal; empty
    and ``*``-only keywords are skipped fail-closed. With no keywords there is
    no signature detection and the scan returns ``None``.

    Insertion and validation should share one scan so insertion bounds and any
    post-insertion placement gate cannot drift apart. Paragraphs whose text
    contains any of ``ignore_texts`` (already-inserted clauses) are treated
    like whitespace: they neither start the block nor end the bottom-up scan,
    so a clause inserted into the tail cannot split the signature block.
    """
    element = _find_signature_block_element(
        document, ignore_texts, _compile_signature_keywords(signature_keywords)
    )
    return None if element is None else SignatureBlockStart(element)


def _find_signature_block_element(
    document: DocxDocument,
    ignore_texts: Sequence[str],
    signature_patterns: Sequence[re.Pattern[str]],
) -> CT_P | None:
    """Scan the body tail for the element starting the signature block, or ``None``."""
    body = document.element.body
    if body is None:
        # A document whose w:body is missing has no tail to scan; the
        # corruption is check_document_integrity's to report, not a crash here.
        return None
    children = list(body.iterchildren())
    lowered_ignores = [ignored.lower() for ignored in ignore_texts if ignored]

    block_start_index = len(children)
    scanned_without_match = 0
    for i in range(len(children) - 1, -1, -1):
        element = children[i]
        # Comment/PI nodes expose a callable .tag, not a string (lxml).
        if not isinstance(element.tag, str) or not element.tag.endswith("}p"):
            continue
        text_content = " ".join(element.itertext()).lower()
        if not text_content.strip():
            # Whitespace-only paragraphs neither start a signature block
            # nor consume the tail window.
            continue
        if any(ignored in text_content for ignored in lowered_ignores):
            continue
        if any(pattern.search(text_content) for pattern in signature_patterns):
            block_start_index = i
        elif block_start_index < len(children):
            # First non-matching paragraph above the signature block ends the scan.
            break
        else:
            scanned_without_match += 1
            if scanned_without_match >= _SIGNATURE_TAIL_WINDOW:
                # No signature block starts within the tail window.
                break

    if block_start_index < len(children):
        return children[block_start_index]
    return None


def _is_ascii_digits(token: str) -> bool:
    """True for non-empty ASCII-decimal tokens.

    ``str.isdigit`` alone also accepts Unicode digits (e.g. ``"²"``, ``"①"``)
    that ``int()`` rejects; such tokens must fail closed, not raise.
    """
    return token.isascii() and token.isdigit()


def parse_body_anchor_index(anchor: str) -> int | None:
    """Return the body-paragraph index for ``body:p:<digits>`` anchors, else ``None``."""
    parts = anchor.split(":")
    if len(parts) == 3 and parts[0] == "body" and parts[1] == "p" and _is_ascii_digits(parts[2]):
        return int(parts[2])
    return None


def is_supported_anchor(anchor: str) -> bool:
    """True when the anchor uses the grammar the insertion engine can actually consume."""
    return anchor == ANCHOR_LAST or parse_body_anchor_index(anchor) is not None


def find_body_paragraph(document: DocxDocument, index: int) -> Paragraph | None:
    """Resolve a body-local paragraph index; ``None`` outside the body (fail closed)."""
    paragraphs = document.paragraphs
    if 0 <= index < len(paragraphs):
        return paragraphs[index]
    return None


def find_paragraph_by_locator(document: DocxDocument, locator: str | None) -> Paragraph | None:
    """Resolve a parser locator (``body|table|header|footer ...:p:<n>``) to its paragraph.

    Locator indices are local to their section, matching the segment ids the
    DOCX parser emits. Unknown or out-of-range locators return ``None``; the
    grammar is strict (digits only, no trailing tokens), so malformed locators
    fail closed instead of resolving to a nearby paragraph.
    """
    if not locator:
        return None
    parts = locator.split(":")
    if len(parts) < 3 or parts[-2] != "p" or not _is_ascii_digits(parts[-1]):
        return None
    index = int(parts[-1])
    prefix = parts[:-2]

    paragraphs: list[Paragraph] | None = None
    if prefix == ["body"]:
        paragraphs = list(document.paragraphs)
    elif (
        len(prefix) == 6
        and prefix[0] == "table"
        and prefix[2] == "row"
        and prefix[4] == "cell"
        and _is_ascii_digits(prefix[1])
        and _is_ascii_digits(prefix[3])
        and _is_ascii_digits(prefix[5])
    ):
        table_index, row_index, cell_index = int(prefix[1]), int(prefix[3]), int(prefix[5])
        if table_index < len(document.tables):
            table = document.tables[table_index]
            if row_index < len(table.rows):
                row = table.rows[row_index]
                if cell_index < len(row.cells):
                    paragraphs = list(row.cells[cell_index].paragraphs)
    elif len(prefix) == 2 and prefix[0] in ("header", "footer") and _is_ascii_digits(prefix[1]):
        section_index = int(prefix[1])
        if section_index < len(document.sections):
            section = document.sections[section_index]
            part = section.header if prefix[0] == "header" else section.footer
            paragraphs = list(part.paragraphs)

    if paragraphs is None or not 0 <= index < len(paragraphs):
        return None
    return paragraphs[index]
