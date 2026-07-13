"""Minimal DOCX renderers for reviewed and corrected output."""

from __future__ import annotations

from collections.abc import Callable
from copy import deepcopy
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Literal

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reviewkit.actions import (
    _actions_in_text_application_order,
    _find_anchor_application_order,
    actions_for_paragraph,
    apply_action_to_text,
    should_apply_to_corrected,
)
from reviewkit.docx_package import normalize_docx_timestamps
from reviewkit.document import ReviewDocument
from reviewkit.models import ActionStatus, ReviewAction, ReviewActionType
from reviewkit.policy import WRITING_ACTIONS
from doctotext import (
    InlineSegment as _InlineSegment,  # base mechanical only (text/opaque + rpr/element)
    paragraph_to_inline_segments,
    _advances_offset as _base_advances_offset,
    _visible_text as _base_visible_text,
    _visible_len as _base_visible_len,
    _split_visible_offset as _base_split_visible_offset,
    _insert_visible as _base_insert_visible,
    _replace_visible_range as _base_replace_visible_range,
    _rpr_at as _base_rpr_at,
    _index_at_visible_offset as _base_index_at_visible_offset,
    _copy_segment as _base_copy_segment,
)
# ------------------------------------------------------------------
# Delegation to DocToText for pure mechanical DOCX (addressing, run splitting,
# visible offset math, rPr/opaque preservation). reviewkit owns only the
# review overlay (ins/del markup as decision trace, comments, revision ids,
# apply_to_corrected, integrity, policy).
# ------------------------------------------------------------------
def _inline_to_review(seg: _InlineSegment, **review_fields: Any) -> _Segment:
    """Lift a base mechanical InlineSegment to a review _Segment, attaching review metadata."""
    return _Segment(
        kind=seg.kind,  # type: ignore[arg-type]
        text=seg.text,
        rpr=seg.rpr,
        element=seg.element,
        **review_fields,
    )


def _review_to_inline(seg: _Segment) -> _InlineSegment:
    return _InlineSegment(
        kind=seg.kind,  # type: ignore[arg-type]
        text=seg.text,
        rpr=seg.rpr,
        element=seg.element,
    )


def _lift_list(base_list: list[_InlineSegment]) -> list[_Segment]:
    return [_inline_to_review(s) for s in base_list]


def _to_base_list(review_list: list[_Segment]) -> list[_InlineSegment]:
    return [_review_to_inline(s) for s in review_list]

_SegmentKind = Literal["text", "ins", "del", "opaque"]

# Deterministic default stamped on tracked-change revisions so ``reviewed.docx`` is
# reproducible byte-for-byte. Callers that want wall-clock revision dates opt in by
# passing ``revision_timestamp=datetime.now(UTC)`` explicitly.
_DEFAULT_REVISION_DATE = "1970-01-01T00:00:00+00:00"


class RenderIntegrityError(RuntimeError):
    """A writing action that must leave a trace in an artifact could not be rendered.

    After ``prepare_actions`` validation every writing action anchors deterministically,
    so a miss at render time is an internal inconsistency (unknown node_id, a paragraph
    locator that no longer resolves, stale char offsets, drifted text) - never a
    legitimate skip. Raising converts what used to be silent data loss (an artifact
    quietly missing edits) into a hard error, so consumers can fail closed on a raised
    render call and trust a successful one without re-parsing the DOCX.
    """


def _describe_action(action: ReviewAction) -> str:
    locator = action.locator
    if locator is not None and locator.char_start is not None and locator.char_end is not None:
        where = f"chars [{locator.char_start}, {locator.char_end})"
    else:
        where = "no char range"
    return (
        f"action {action.id!r} ({action.action_type.value}, status={action.status.value}) "
        f"on node {action.node_id!r} ({where}, original_text={action.original_text!r})"
    )


def _assert_writing_actions_reach_a_paragraph(
    document: ReviewDocument,
    actions: list[ReviewAction],
    must_render: Callable[[ReviewAction], bool],
    artifact: str,
) -> None:
    """Fail closed when a writing action's node_id routes to no paragraph.

    ``actions_for_paragraph`` silently ignores an action whose node_id names no
    paragraph or sentence and is not a section/document scope that can anchor
    (original_text present and the scope has at least one paragraph); such an
    action would leave no trace in the artifact.
    """
    paragraph_ids = {paragraph.id for paragraph in document.iter_paragraphs()}
    sentence_ids = {sentence.id for sentence in document.iter_sentences()}

    def _routes(action: ReviewAction) -> bool:
        node_id = action.node_id
        if node_id in paragraph_ids or node_id in sentence_ids:
            return True
        if not action.original_text:
            return False
        if node_id == document.id:
            return bool(paragraph_ids)
        section = next((s for s in document.sections if s.id == node_id), None)
        return section is not None and bool(section.paragraphs)

    dropped = [action for action in actions if must_render(action) and not _routes(action)]
    if dropped:
        raise RenderIntegrityError(
            f"{artifact} would silently drop writing action(s) whose node_id routes to "
            "no paragraph: " + "; ".join(_describe_action(action) for action in dropped)
        )


def _anchored_in_text(text: str, action: ReviewAction) -> bool:
    """Whether ``action`` has a usable anchor against the pristine paragraph ``text``."""
    action_range = _locator_action_range(action)
    if action_range is not None:
        return action_range[1] <= len(text)
    if action.original_text:
        return action.original_text in text
    # Insertions without an anchor legitimately default to paragraph start/end.
    return action.action_type in {
        ReviewActionType.INSERT_TEXT,
        ReviewActionType.INSERT_BEFORE,
        ReviewActionType.INSERT_AFTER,
    }


@dataclass(frozen=True)
class _ReviewerIdentity:
    """Author identity stamped onto tracked-change revisions and Word comments."""

    author: str = "Reviewer"
    initials: str = "RV"
    revision_date: str = _DEFAULT_REVISION_DATE


@dataclass
class _Segment:
    kind: _SegmentKind
    text: str
    rpr: Any | None = None
    action_id: str | None = None
    revision_id: int | None = None
    start_comments: list[int] = field(default_factory=list)
    end_comments: list[int] = field(default_factory=list)
    # For ``opaque`` segments: the original inline element (image, hyperlink, tab,
    # break, field, pre-existing revision, ...) re-emitted verbatim on rebuild. Its
    # ``text`` is the visible characters it contributes to the paragraph so char
    # offsets around it stay aligned with the parser's coordinate system.
    element: Any | None = None


def _advances_offset(segment: _Segment) -> bool:
    return _base_advances_offset(_review_to_inline(segment))


def render_reviewed_docx(
    document: ReviewDocument,
    actions: list[ReviewAction],
    output_path: str | Path,
    *,
    comment_author: str = "Reviewer",
    comment_initials: str = "RV",
    revision_timestamp: datetime | None = None,
) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    _assert_writing_actions_reach_a_paragraph(
        document, actions, _is_trackable_edit, "reviewed.docx"
    )
    docx = DocxDocument(str(document.source_path)) if document.source_path else DocxDocument()
    revision_date = (
        _DEFAULT_REVISION_DATE
        if revision_timestamp is None
        else revision_timestamp.replace(microsecond=0).isoformat()
    )
    reviewer = _ReviewerIdentity(
        author=comment_author, initials=comment_initials, revision_date=revision_date
    )
    revision_id = 1
    # Stand-alone clause insertions (``new_paragraph``) splice NEW sibling paragraphs
    # into the body. _paragraph_for_locator indexes ``docx.paragraphs`` positionally,
    # so splicing mid-pass would shift every later paragraph's index and desync the
    # remaining locators. We hold each resolved anchor element and splice against those
    # stable references only after every locator has been resolved (just before save).
    deferred_block_inserts: list[tuple[Any, ReviewAction]] = []

    for section in document.sections:
        for paragraph in section.paragraphs:
            paragraph_actions = actions_for_paragraph(document, paragraph, actions)
            block_inserts: list[ReviewAction] = []
            inline_actions: list[ReviewAction] = []
            for action in paragraph_actions:
                (block_inserts if _is_block_paragraph_insert(action) else inline_actions).append(
                    action
                )
            docx_paragraph = _paragraph_for_locator(docx, paragraph.locator)
            if docx_paragraph is None:
                trackable = [a for a in paragraph_actions if _is_trackable_edit(a)]
                if trackable and document.source_path is not None:
                    # With a source document every parsed locator must resolve; landing
                    # tracked edits in an appended duplicate paragraph would leave the
                    # real paragraph untouched.
                    raise RenderIntegrityError(
                        f"reviewed.docx: paragraph locator {paragraph.locator!r} does not "
                        "resolve in the source document; tracked edits would land in a "
                        "detached paragraph: "
                        + "; ".join(_describe_action(a) for a in trackable)
                    )
                docx_paragraph = docx.add_paragraph(paragraph.text)
            revision_id = _add_reviewed_runs(
                docx,
                docx_paragraph,
                inline_actions,
                revision_id,
                reviewer,
            )
            deferred_block_inserts.extend((docx_paragraph, action) for action in block_inserts)

        section_comments = [
            _comment_text(action)
            for action in actions
            if action.node_id == section.id and not action.original_text
        ]
        for comment in section_comments:
            if comment:
                target = section.locator or (
                    section.paragraphs[0].locator if section.paragraphs else None
                )
                section_paragraph = _paragraph_for_locator(docx, target)
                if section_paragraph is None:
                    section_paragraph = docx.add_paragraph(section.title or section.id)
                if not _add_comment(docx, section_paragraph, comment, reviewer):
                    docx.add_paragraph(comment)

    document_comments = [
        _comment_text(action)
        for action in actions
        if action.node_id == document.id and not action.original_text
    ]
    if document_comments:
        first_paragraph = next(document.iter_paragraphs(), None)
        for comment in document_comments:
            if comment:
                document_paragraph = _paragraph_for_locator(
                    docx,
                    first_paragraph.locator if first_paragraph else None,
                )
                if document_paragraph is None:
                    document_paragraph = docx.add_paragraph("Document-level review")
                if not _add_comment(docx, document_paragraph, comment, reviewer):
                    docx.add_paragraph(comment)

    # Every locator is resolved now; splice the stand-alone clause paragraphs against
    # the anchor elements we held onto (see deferred_block_inserts above).
    for anchor_paragraph, action in deferred_block_inserts:
        revision_id = _insert_tracked_block_paragraphs(
            anchor_paragraph, action, reviewer, revision_id
        )

    docx.save(str(path))
    # python-docx stamps every zip entry with the wall-clock mtime, which alone makes an
    # otherwise byte-identical reviewed.docx differ on every run; pin them so the promised
    # byte-for-byte reproducibility actually holds at the package level.
    normalize_docx_timestamps(path)
    return path


def render_corrected_docx(
    document: ReviewDocument,
    actions: list[ReviewAction],
    output_path: str | Path,
) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    _assert_writing_actions_reach_a_paragraph(
        document, actions, should_apply_to_corrected, "corrected.docx"
    )
    # Apply the accepted edits onto a copy of the ORIGINAL document so tables,
    # headers/footers, images, styles and section structure are preserved. Only
    # fall back to a blank document when there is no source to copy from.
    docx = DocxDocument(str(document.source_path)) if document.source_path else DocxDocument()
    has_source = document.source_path is not None

    for section in document.sections:
        if not has_source and section.title:
            docx.add_heading(section.title, level=1)
        for paragraph in section.paragraphs:
            paragraph_actions = [
                action
                for action in actions_for_paragraph(document, paragraph, actions)
                if should_apply_to_corrected(action)
            ]
            docx_paragraph = _paragraph_for_locator(docx, paragraph.locator) if has_source else None
            if docx_paragraph is None:
                if has_source and paragraph_actions:
                    # With a source document every parsed locator must resolve; applying
                    # the edits to an appended duplicate paragraph would leave the real
                    # paragraph uncorrected.
                    raise RenderIntegrityError(
                        f"corrected.docx: paragraph locator {paragraph.locator!r} does not "
                        "resolve in the source document; APPLIED edits would land in a "
                        "detached paragraph: "
                        + "; ".join(_describe_action(a) for a in paragraph_actions)
                    )
                # Check each anchor against the EVOLVING text right before applying, in
                # the exact application order: an earlier APPLIED edit can consume a
                # later edit's original_text, and str.replace would then silently no-op
                # while the report still claims the edit APPLIED.
                corrected = paragraph.text
                for action in _actions_in_text_application_order(paragraph_actions):
                    if not _anchored_in_text(corrected, action):
                        raise RenderIntegrityError(
                            "corrected.docx would silently drop an APPLIED edit that fails "
                            "to anchor in its paragraph: " + _describe_action(action)
                        )
                    corrected = apply_action_to_text(corrected, action)
                docx.add_paragraph(corrected)
                continue
            _apply_clean_corrections(docx_paragraph, paragraph_actions)

    docx.save(str(path))
    # See render_reviewed_docx: pin python-docx's wall-clock zip timestamps so corrected.docx
    # is reproducible byte-for-byte too.
    normalize_docx_timestamps(path)
    return path


def _apply_clean_corrections(paragraph: Any, actions: list[ReviewAction]) -> None:
    """Rewrite ``paragraph``'s runs with the accepted edits applied, no tracked changes."""
    if not actions:
        return
    segments = _paragraph_segments(paragraph)
    actions = _align_locators_to_visible_text(segments, actions)
    revision_id = 1
    for action in _trackable_actions_in_application_order(actions):
        segments, revision_id, tracked = _track_action(segments, action, revision_id)
        if not tracked:
            # Every action here is APPLIED (should_apply_to_corrected filtered): a miss
            # means corrected.docx would ship without an edit the report claims.
            raise RenderIntegrityError(
                "corrected.docx would silently drop an APPLIED edit that fails to anchor "
                "in its paragraph: " + _describe_action(action)
            )

    parent = paragraph._p
    for child in list(parent):
        if child.tag != qn("w:pPr"):
            parent.remove(child)
    for segment in segments:
        if segment.kind == "del":
            continue
        if segment.kind == "opaque":
            if segment.element is not None:
                parent.append(deepcopy(segment.element))
            continue
        if not segment.text:
            continue
        # ``ins`` segments are accepted corrections: they land as ordinary text.
        _append_text_run(parent, segment.text, segment.rpr)


def _paragraph_for_locator(docx: Any, locator: str | None) -> Any | None:
    if not locator:
        return None

    parts = locator.split(":")
    try:
        if parts[:2] == ["body", "p"]:
            return docx.paragraphs[int(parts[2])]
        if parts[:1] == ["table"]:
            table = docx.tables[int(parts[1])]
            row = table.rows[int(parts[3])]
            cell = row.cells[int(parts[5])]
            return cell.paragraphs[int(parts[7])]
        if parts[:1] == ["header"]:
            return docx.sections[int(parts[1])].header.paragraphs[int(parts[3])]
        if parts[:1] == ["footer"]:
            return docx.sections[int(parts[1])].footer.paragraphs[int(parts[3])]
    except (IndexError, ValueError):
        return None
    return None


def _is_trackable_edit(action: ReviewAction) -> bool:
    if action.status == ActionStatus.CONFLICT:
        return False
    if action.metadata.get("blocked_from_corrected") is True:
        return False
    return action.action_type in WRITING_ACTIONS


def _is_block_paragraph_insert(action: ReviewAction) -> bool:
    # A stand-alone clause insert: a trackable INSERT_BEFORE/INSERT_AFTER explicitly
    # flagged new_paragraph. It becomes a NEW tracked paragraph sibling of the anchor
    # (not inline runs patched into the anchor's text), so that accepting the markup
    # yields a real stand-alone paragraph. The flag is explicit rather than inferred
    # from a missing locator because the degraded escalation path emits the same shape
    # but must stay a comment.
    return (
        action.new_paragraph
        and _is_trackable_edit(action)
        and action.action_type
        in {ReviewActionType.INSERT_BEFORE, ReviewActionType.INSERT_AFTER}
    )


def _add_reviewed_runs(
    docx: Any,
    paragraph: Any,
    actions: list[ReviewAction],
    revision_id: int,
    reviewer: _ReviewerIdentity,
) -> int:
    if not actions:
        return revision_id

    segments = _paragraph_segments(paragraph)
    actions = _align_locators_to_visible_text(segments, actions)
    pristine_text = _visible_text(segments)
    first_revision_id = revision_id
    for action in _trackable_actions_in_application_order(actions):
        segments, revision_id, tracked = _track_action(segments, action, revision_id)
        if not tracked and (
            action.status == ActionStatus.APPLIED
            or not _anchored_in_text(pristine_text, action)
        ):
            # An APPLIED edit must ALWAYS appear as a tracked change (it is what
            # corrected.docx applies), and any edit that cannot even anchor in the
            # paragraph's pristine text is an internal inconsistency. Only a
            # non-APPLIED suggestion whose anchor was consumed by an earlier
            # overlapping tracked change keeps the documented degrade to a labelled
            # comment below.
            raise RenderIntegrityError(
                "reviewed.docx would silently drop a tracked edit that fails to anchor "
                "in its paragraph: " + _describe_action(action)
            )

    # Tracking assigns ids in APPLICATION order (right-to-left, inserts first), not
    # document order. Renumber in document order so revision ids read monotonically
    # left to right, keeping reviewed.docx deterministic and diff-friendly.
    revision_id = first_revision_id
    for segment in segments:
        if segment.kind in ("ins", "del"):
            segment.revision_id = revision_id
            revision_id += 1

    for action in actions:
        if not _comment_text(action):
            continue
        if _is_trackable_edit(action) and _mark_action_segments(docx, segments, action, reviewer):
            continue
        if action.original_text and _mark_text_comment(
            docx, segments, action.original_text, action, reviewer
        ):
            continue
        _mark_whole_paragraph_comment(docx, segments, action, reviewer)

    _replace_paragraph_children(paragraph, segments, reviewer)
    return revision_id


def _trackable_actions_in_application_order(actions: list[ReviewAction]) -> list[ReviewAction]:
    locator_actions: list[tuple[int, int, int, ReviewAction]] = []
    other_actions: list[ReviewAction] = []
    for index, action in enumerate(actions):
        if not _is_trackable_edit(action):
            continue
        action_range = _locator_action_range(action)
        if action_range is None:
            other_actions.append(action)
            continue
        start, end = action_range
        locator_actions.append((start, end, index, action))
    locator_actions.sort(key=lambda item: (-item[0], -item[1], item[2]))
    # Find-based actions re-anchor against the evolving visible text, so order them the
    # same way the string path does (zero-width inserts first, APPLIED before
    # suggestions) or a tracked replace/delete consumes another prepare-approved
    # APPLIED action's anchor and the render aborts.
    return [item[3] for item in locator_actions] + _find_anchor_application_order(other_actions)


def _track_action(
    segments: list[_Segment],
    action: ReviewAction,
    revision_id: int,
) -> tuple[list[_Segment], int, bool]:
    """Apply one trackable edit; the third element reports whether it left a trace.

    ``False`` means the action anchored nowhere (or its range produced no revision
    segments at all) and the segments are unchanged - the caller decides whether
    that miss is a hard :class:`RenderIntegrityError` or a documented degrade.
    """
    if action.action_type in {
        ReviewActionType.REPLACE_TEXT,
        ReviewActionType.DELETE_TEXT,
        ReviewActionType.REPLACE,
        ReviewActionType.DELETE,
    }:
        action_range = _action_range(segments, action)
        if action_range is None:
            return segments, revision_id, False
        start, end = action_range
        deleted = _visible_slice(segments, start, end, "del", action.id)
        # The revision must cover EXACTLY the text the action targets. A shorter (or
        # different) selection means part of the range is opaque inline content
        # (tab/break/hyperlink/field/...) or the text drifted: marking only the
        # editable remainder would ship a wrong edit, so report a miss and let the
        # caller fail closed.
        selected_text = "".join(segment.text for segment in deleted)
        expected = (
            (action.locator.original_text if action.locator else None)
            or action.original_text
            or None
        )
        if expected is not None:
            if selected_text != expected:
                return segments, revision_id, False
        elif len(selected_text) != end - start:
            return segments, revision_id, False
        replacement: list[_Segment] = deleted
        if action.action_type in {
            ReviewActionType.REPLACE_TEXT,
            ReviewActionType.REPLACE,
        } and action.replacement_text:
            replacement.append(
                _Segment("ins", action.replacement_text, _rpr_at(segments, start), action.id)
            )
        if not replacement:
            # A zero-width range with no insertion either: nothing would mark this edit.
            return segments, revision_id, False
        revision_id = _assign_revision_ids(replacement, revision_id)
        return _replace_visible_range(segments, start, end, replacement), revision_id, True

    if action.action_type in {
        ReviewActionType.INSERT_TEXT,
        ReviewActionType.INSERT_BEFORE,
        ReviewActionType.INSERT_AFTER,
    }:
        locator = action.locator
        if locator and locator.char_start is not None and locator.char_end is not None:
            # Honor the locator exactly as apply_action_to_text does: INSERT_AFTER
            # inserts at char_end, INSERT_TEXT/INSERT_BEFORE at char_start. Offsets are
            # already aligned to visible-text coordinates by the alignment pass above.
            # find()-by-text below would land on the wrong occurrence when the anchor
            # repeats, diverging the reviewed markup from the corrected text.
            if action.action_type == ReviewActionType.INSERT_AFTER:
                offset = locator.char_end
            else:
                offset = locator.char_start
        elif action.original_text:
            start = _visible_text(segments).find(action.original_text)
            if start < 0:
                return segments, revision_id, False
            offset = start
            if action.action_type == ReviewActionType.INSERT_AFTER:
                offset += len(action.original_text)
        else:
            offset = 0 if action.action_type == ReviewActionType.INSERT_BEFORE else _visible_len(
                segments
            )
        insert = _Segment(
            "ins",
            action.replacement_text or "",
            _rpr_at(segments, offset),
            action.id,
            revision_id,
        )
        return _insert_visible(segments, offset, insert), revision_id + 1, True

    return segments, revision_id, True


def _align_locators_to_visible_text(
    segments: list[_Segment], actions: list[ReviewAction]
) -> list[ReviewAction]:
    """Shift char offsets to the renderer's raw-run coordinate system.

    The parser strips paragraph text, so validated char offsets are relative to
    the stripped text, while the renderer rebuilds visible text from the raw
    (unstripped) runs. Any leading whitespace therefore skews every offset; shift
    locator ranges by that amount so edits land on the intended characters.
    """
    visible = _visible_text(segments)
    leading = len(visible) - len(visible.lstrip())
    if leading == 0:
        return actions
    aligned: list[ReviewAction] = []
    for action in actions:
        locator = action.locator
        if locator is not None and locator.char_start is not None and locator.char_end is not None:
            aligned.append(
                action.model_copy(
                    update={
                        "locator": locator.model_copy(
                            update={
                                "char_start": locator.char_start + leading,
                                "char_end": locator.char_end + leading,
                            }
                        )
                    }
                )
            )
        else:
            aligned.append(action)
    return aligned


def _paragraph_segments(paragraph: Any) -> list[_Segment]:
    base = paragraph_to_inline_segments(paragraph)
    return _lift_list(base)



def _replace_visible_range(
    segments: list[_Segment],
    start: int,
    end: int,
    replacement: list[_Segment],
) -> list[_Segment]:
    base_rep = _to_base_list(replacement)
    base = _base_replace_visible_range(_to_base_list(segments), start, end, base_rep)
    return _lift_list(base)


def _insert_visible(segments: list[_Segment], offset: int, insert: _Segment) -> list[_Segment]:
    base_ins = _review_to_inline(insert)
    base = _base_insert_visible(_to_base_list(segments), offset, base_ins)
    return _lift_list(base)


def _split_visible_offset(segments: list[_Segment], offset: int) -> list[_Segment]:
    base = _base_split_visible_offset(_to_base_list(segments), offset)
    return _lift_list(base)


def _index_at_visible_offset(segments: list[_Segment], offset: int) -> int:
    return _base_index_at_visible_offset(_to_base_list(segments), offset)


def _visible_slice(
    segments: list[_Segment],
    start: int,
    end: int,
    kind: Literal["ins", "del"],
    action_id: str,
) -> list[_Segment]:
    segments = _split_visible_offset(_split_visible_offset(segments, end), start)
    selected: list[_Segment] = []
    offset = 0
    for segment in segments:
        next_offset = offset + (len(segment.text) if _advances_offset(segment) else 0)
        if segment.kind == "text" and start <= offset and next_offset <= end:
            selected.append(_Segment(kind, segment.text, deepcopy(segment.rpr), action_id))
        offset = next_offset
    return selected


def _assign_revision_ids(segments: list[_Segment], revision_id: int) -> int:
    for segment in segments:
        if segment.kind != "text":
            segment.revision_id = revision_id
            revision_id += 1
    return revision_id


def _visible_text(segments: list[_Segment]) -> str:
    return _base_visible_text(_to_base_list(segments))


def _action_range(segments: list[_Segment], action: ReviewAction) -> tuple[int, int] | None:
    locator_range = _locator_action_range(action)
    if locator_range is not None:
        return locator_range
    if not action.original_text:
        return None
    start = _visible_text(segments).find(action.original_text)
    if start < 0:
        return None
    return start, start + len(action.original_text)


def _locator_action_range(action: ReviewAction) -> tuple[int, int] | None:
    if not action.locator:
        return None
    if action.locator.char_start is None or action.locator.char_end is None:
        return None
    return action.locator.char_start, action.locator.char_end


def _visible_len(segments: list[_Segment]) -> int:
    return _base_visible_len(_to_base_list(segments))


def _rpr_at(segments: list[_Segment], offset: int) -> Any | None:
    return _base_rpr_at(_to_base_list(segments), offset)


def _copy_segment(segment: _Segment, text: str) -> _Segment:
    base = _base_copy_segment(_review_to_inline(segment), text)
    return _inline_to_review(
        base,
        action_id=segment.action_id,
        revision_id=segment.revision_id,
        start_comments=list(segment.start_comments),
        end_comments=list(segment.end_comments),
    )


def _replace_paragraph_children(
    paragraph: Any, segments: list[_Segment], reviewer: _ReviewerIdentity
) -> None:
    parent = paragraph._p
    for child in list(parent):
        if child.tag != qn("w:pPr"):
            parent.remove(child)

    for segment in segments:
        if not segment.text and segment.kind != "opaque":
            continue
        for comment_id in segment.start_comments:
            parent.append(_comment_range("w:commentRangeStart", comment_id))
        if segment.kind == "text":
            _append_text_run(parent, segment.text, segment.rpr)
        elif segment.kind == "opaque":
            if segment.element is not None:
                parent.append(deepcopy(segment.element))
        else:
            _append_revision(
                parent, segment.kind, segment.text, reviewer, segment.rpr, segment.revision_id
            )
        for comment_id in reversed(segment.end_comments):
            parent.append(_comment_range("w:commentRangeEnd", comment_id))
            parent.append(_comment_reference_run(comment_id))


def _append_text_run(parent: Any, text: str, rpr: Any | None = None) -> None:
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(deepcopy(rpr))
    text_element = OxmlElement("w:t")
    _set_text(text_element, text)
    run.append(text_element)
    parent.append(run)


def _append_revision(
    parent: Any,
    kind: Literal["ins", "del"],
    text: str,
    reviewer: _ReviewerIdentity,
    rpr: Any | None = None,
    revision_id: int | None = None,
) -> None:
    revision = OxmlElement(f"w:{kind}")
    revision.set(qn("w:id"), str(revision_id or 0))
    revision.set(qn("w:author"), reviewer.author)
    revision.set(qn("w:date"), reviewer.revision_date)

    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(deepcopy(rpr))
    text_element = OxmlElement("w:t" if kind == "ins" else "w:delText")
    _set_text(text_element, text)
    run.append(text_element)
    revision.append(run)
    parent.append(revision)


def _insert_tracked_block_paragraphs(
    anchor_paragraph: Any,
    action: ReviewAction,
    reviewer: _ReviewerIdentity,
    revision_id: int,
) -> int:
    # Materialise a new_paragraph insert as one or more NEW tracked paragraphs spliced
    # next to the anchor. Each line of replacement_text becomes its own paragraph whose
    # paragraph mark AND runs are wrapped in w:ins, so Word (and accept_all_revisions)
    # treats it as a genuinely inserted stand-alone paragraph.
    text = action.replacement_text
    if not text:
        raise RenderIntegrityError(
            "reviewed.docx: new_paragraph insert carries no replacement_text: "
            + _describe_action(action)
        )
    lines = text.split("\n")
    if len(lines) > 1 and lines[-1] == "":
        # Drop the empty tail a trailing newline produces; keep interior blank lines.
        lines = lines[:-1]

    new_paragraphs: list[Any] = []
    for line in lines:
        paragraph_element, revision_id = _build_tracked_paragraph(line, reviewer, revision_id)
        new_paragraphs.append(paragraph_element)

    anchor = anchor_paragraph._p
    if action.action_type == ReviewActionType.INSERT_BEFORE:
        # addprevious keeps each new paragraph immediately before the anchor, so the
        # list stays in order: [line0, line1, ..., anchor].
        for paragraph_element in new_paragraphs:
            anchor.addprevious(paragraph_element)
    else:
        # addnext inserts after a moving cursor: anchor -> line0 -> line1 -> ...
        cursor = anchor
        for paragraph_element in new_paragraphs:
            cursor.addnext(paragraph_element)
            cursor = paragraph_element
    return revision_id


def _build_tracked_paragraph(
    text: str, reviewer: _ReviewerIdentity, revision_id: int
) -> tuple[Any, int]:
    paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    run_properties = OxmlElement("w:rPr")
    # The paragraph mark itself is an insertion: <w:pPr><w:rPr><w:ins/></w:rPr></w:pPr>.
    # Without this, accepting changes would merge this paragraph's text up into the
    # previous one instead of keeping it stand-alone.
    mark = OxmlElement("w:ins")
    mark.set(qn("w:id"), str(revision_id))
    mark.set(qn("w:author"), reviewer.author)
    mark.set(qn("w:date"), reviewer.revision_date)
    run_properties.append(mark)
    properties.append(run_properties)
    paragraph.append(properties)
    revision_id += 1

    _append_revision(paragraph, "ins", text, reviewer, None, revision_id)
    return paragraph, revision_id + 1


def _set_text(element: Any, text: str) -> None:
    if text[:1].isspace() or text[-1:].isspace():
        element.set(qn("xml:space"), "preserve")
    element.text = text


def _mark_action_segments(
    docx: Any, segments: list[_Segment], action: ReviewAction, reviewer: _ReviewerIdentity
) -> bool:
    indexes = [index for index, segment in enumerate(segments) if segment.action_id == action.id]
    return _mark_comment_indexes(docx, segments, indexes, action, reviewer)


def _mark_text_comment(
    docx: Any,
    segments: list[_Segment],
    original_text: str,
    action: ReviewAction,
    reviewer: _ReviewerIdentity,
) -> bool:
    start = _visible_text(segments).find(original_text)
    if start < 0:
        return False
    end = start + len(original_text)
    segments[:] = _split_visible_offset(_split_visible_offset(segments, end), start)
    indexes: list[int] = []
    offset = 0
    for index, segment in enumerate(segments):
        # ``start``/``end`` come from _visible_text, which counts every _advances_offset
        # segment (text AND opaque). Advance this accumulator the same way, or an opaque
        # segment before the quote desyncs the two coordinate systems, leaves ``indexes``
        # empty, and the anchor silently degrades to a whole-paragraph comment.
        next_offset = offset + (len(segment.text) if _advances_offset(segment) else 0)
        if segment.kind == "text" and start <= offset and next_offset <= end:
            indexes.append(index)
        offset = next_offset
    return _mark_comment_indexes(docx, segments, indexes, action, reviewer)


def _mark_whole_paragraph_comment(
    docx: Any, segments: list[_Segment], action: ReviewAction, reviewer: _ReviewerIdentity
) -> bool:
    indexes = [index for index, segment in enumerate(segments) if segment.text]
    return _mark_comment_indexes(docx, segments, indexes, action, reviewer)


def _mark_comment_indexes(
    docx: Any,
    segments: list[_Segment],
    indexes: list[int],
    action: ReviewAction,
    reviewer: _ReviewerIdentity,
) -> bool:
    comment = _comment_text(action)
    if comment is None or not indexes:
        return False
    comment_id = _create_comment(docx, comment, reviewer)
    segments[indexes[0]].start_comments.append(comment_id)
    segments[indexes[-1]].end_comments.append(comment_id)
    return True


def _create_comment(docx: Any, text: str, reviewer: _ReviewerIdentity) -> int:
    comment = docx.comments.add_comment(
        text=text, author=reviewer.author, initials=reviewer.initials
    )
    _stamp_comment_date(comment, reviewer)
    return int(comment.comment_id)


def _stamp_comment_date(comment: Any, reviewer: _ReviewerIdentity) -> None:
    # python-docx stamps ``w:date`` with wall-clock ``datetime.now(utc)`` inside
    # ``add_comment``, which alone makes reviewed.docx non-reproducible. Overwrite it
    # with the same deterministic revision date used for tracked changes so identical
    # inputs yield a byte-identical comments.xml (and comments share the ins/del date).
    comment._comment_elm.set(qn("w:date"), reviewer.revision_date)


def _comment_range(tag: str, comment_id: int) -> Any:
    element = OxmlElement(tag)
    element.set(qn("w:id"), str(comment_id))
    return element


def _comment_reference_run(comment_id: int) -> Any:
    run = OxmlElement("w:r")
    reference = OxmlElement("w:commentReference")
    reference.set(qn("w:id"), str(comment_id))
    run.append(reference)
    return run


def _comment_text(action: ReviewAction) -> str | None:
    label = _comment_label(action)
    parts = [f"{label}: {action.comment or action.reason or action.policy_reason or ''}".rstrip()]
    if action.original_text:
        parts.append(f"Original: {action.original_text!r}")
    if action.replacement_text:
        parts.append(f"Replacement: {action.replacement_text!r}")
    if action.category:
        parts.append(f"Category: {action.category}")
    if action.policy_reason:
        parts.append(f"Policy: {action.policy_reason}")
    if action.references:
        refs = ", ".join(reference.label or reference.source for reference in action.references)
        parts.append(f"References: {refs}")
    if action.evidence_refs:
        evidence = ", ".join(
            ref.locator or ref.segment_id or ref.source or "evidence"
            for ref in action.evidence_refs
        )
        parts.append(f"Evidence: {evidence}")
    parts.append(f"Status: {action.status.value}")
    return "\n".join(parts)


def _comment_label(action: ReviewAction) -> str:
    if action.action_type in {
        ReviewActionType.REPLACE_TEXT,
        ReviewActionType.DELETE_TEXT,
        ReviewActionType.INSERT_TEXT,
        ReviewActionType.REPLACE,
        ReviewActionType.DELETE,
        ReviewActionType.INSERT_BEFORE,
        ReviewActionType.INSERT_AFTER,
    }:
        if action.status == ActionStatus.APPLIED:
            return "CORRECTION"
        if action.status == ActionStatus.CONFLICT:
            return "CONFLICT"
        if action.status == ActionStatus.NEEDS_HUMAN_DECISION:
            return "HUMAN_DECISION"
        return "SUGGESTION"
    if action.action_type == ReviewActionType.QUESTION:
        return "QUESTION"
    if action.action_type == ReviewActionType.RISK:
        return "RISK"
    if action.action_type == ReviewActionType.SUGGESTION:
        return "SUGGESTION"
    if action.action_type == ReviewActionType.PRAISE:
        return "PRAISE"
    if action.action_type == ReviewActionType.SUMMARY:
        return "SUMMARY"
    return "COMMENT"


def _add_comment(
    docx: Any, paragraph: Any, text: str, reviewer: _ReviewerIdentity
) -> bool:
    try:
        runs = getattr(paragraph, "runs")
        if not runs:
            paragraph.add_run("")
            runs = getattr(paragraph, "runs")
        comment = docx.comments.add_comment(
            text=text, author=reviewer.author, initials=reviewer.initials
        )
        _stamp_comment_date(comment, reviewer)
        runs[0].mark_comment_range(runs[-1], comment.comment_id)
    except AttributeError:
        return False
    return True
