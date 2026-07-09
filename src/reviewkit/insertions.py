"""Anchored paragraph-insertion engine for corrected DOCX generation.

Inserts whole paragraphs (fix clauses or suggestion markers) into an existing
DOCX at ``body:p:<n>`` / ``body:p:last`` anchors (see :mod:`reviewkit.anchors`),
with deterministic placement rules:

- every anchor resolves against the pristine document before any mutation, so
  earlier insertions cannot shift later targets;
- actions sharing an anchor chain in batch order;
- a paragraph directly followed by a table is a lead-in for that table, so
  insertions anchored on it land after the table, never between the two;
- end-of-document insertions stay above a caller-defined trailing signature
  block instead of landing below the signature lines.

The engine is domain-blind: what to insert, where, and which texts mark a
signature block are entirely caller decisions.
"""

from __future__ import annotations

from collections.abc import Callable, Sequence
from dataclasses import dataclass
from typing import Any, Literal

from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.text.paragraph import Paragraph

from reviewkit.anchors import (
    ANCHOR_LAST,
    _compile_signature_keywords,
    _find_signature_block_element,
    parse_body_anchor_index,
)

InsertionKind = Literal["insert", "suggest"]

# The literal prefix every ``suggest`` insertion carries. It is reviewkit's own marker
# (produced just below), so any consumer that needs to detect a surviving suggestion --
# e.g. a clean-copy purity gate -- matches against this constant / :func:`contains_suggestion_marker`
# rather than hard-coding the string, and stays correct if the marker format ever changes.
SUGGESTION_MARKER_PREFIX = "[SUGGESTION"


def format_suggestion_text(reason: str, text: str) -> str:
    """Render the paragraph text a ``suggest`` action inserts into the document."""
    return f"{SUGGESTION_MARKER_PREFIX}: {reason}]\n{text}"


def contains_suggestion_marker(text: str) -> bool:
    """Whether ``text`` carries a suggestion marker produced by :func:`format_suggestion_text`.

    Accepts the visible document text or the raw ``word/document.xml`` -- the marker prefix
    is emitted verbatim as run text, so it survives intact in both.
    """
    return SUGGESTION_MARKER_PREFIX in text


@dataclass(frozen=True)
class InsertionAction:
    """One requested paragraph insertion.

    ``insert`` actions insert ``text`` verbatim; ``suggest`` actions insert the
    :func:`format_suggestion_text` rendering of ``reason`` and ``text``.
    ``action_id`` is an opaque caller key used to correlate results (and
    :class:`ParagraphInserter`'s ``resolve_last_anchor`` calls) back to caller
    state.
    """

    action_id: str
    anchor: str
    text: str = ""
    kind: InsertionKind = "insert"
    reason: str = ""

    def rendered_text(self) -> str:
        if self.kind == "suggest":
            return format_suggestion_text(self.reason, self.text)
        return self.text


@dataclass
class InsertionResult:
    """A recorded insertion whose predecessor stays adjacency-true.

    ``predecessor`` is the body element the inserted paragraph directly
    follows — usually a paragraph, but it can be e.g. a table when an end
    insert lands below one. Later insertions that land in that slot remap it,
    so the recorded ``applied_anchor`` (``body:p:<n>`` of the predecessor in
    the final document; ``None`` when the inserted paragraph has no preceding
    body paragraph — a plain append into an empty body, an insert at the top,
    or a non-paragraph predecessor) never claims an adjacency another
    insertion took over. ``action`` keeps the original request, including the
    originally requested ``anchor``.
    """

    action: InsertionAction
    element: CT_P
    predecessor: BaseOxmlElement | None
    applied_anchor: str | None = None


# eq=False: list fields would make the generated __eq__/__hash__ unusable
# (frozen + lists means hash() raises); identity semantics match the lxml
# elements the results carry.
@dataclass(frozen=True, eq=False)
class InsertionReport:
    """Outcome of one insertion batch: applied results and failed actions."""

    applied: list[InsertionResult]
    failed: list[InsertionAction]


@dataclass(frozen=True)
class _ResolvedAction:
    """An action resolved to its insertion target before any mutation.

    ``anchor_paragraph`` holds a live lxml-backed paragraph reference that
    survives later insertions. For signature-aware end inserts it is ``None``
    and ``insert_before`` carries the pre-resolved signature-block start
    (``None`` there means plain append at the end of the body).
    """

    text: str
    anchor_paragraph: Paragraph | None
    insert_before: CT_P | None = None


def _skip_following_tables(element: CT_P) -> BaseOxmlElement:
    """Return the last element of the adjacency unit starting at ``element``.

    A paragraph directly followed by one or more ``w:tbl`` siblings is a
    table lead-in; insertions anchored on it must land after the table(s),
    never between the lead-in and its table.
    """
    last: BaseOxmlElement = element
    sibling = element.getnext()
    while sibling is not None and isinstance(sibling.tag, str) and sibling.tag.endswith("}tbl"):
        last = sibling
        sibling = sibling.getnext()
    return last


def _intended_anchor(
    paragraph_positions: dict[int, int], predecessor: BaseOxmlElement | None
) -> str | None:
    """Return the anchor of the paragraph the insertion was intended to follow.

    The predecessor is captured at apply time from the resolved target, never
    from where the new element actually landed, so post-insertion validation
    can detect misplacement instead of blessing it.
    """
    if predecessor is None:
        return None
    index = paragraph_positions.get(id(predecessor))
    if index is None:
        return None
    return f"body:p:{index}"


class ParagraphInserter:
    """Applies one insertion batch to one document.

    Anchors resolve against the pristine document, so an instance supports a
    single mutating call; further calls fail loudly instead of silently
    re-resolving against shifted indices.

    ``resolve_last_anchor`` is called for actions requesting ``body:p:last``
    BEFORE any suggestion-text formatting, with the raw action; it may return
    a concrete ``body:p:<n>`` anchor (contextual placement) or ``body:p:last``
    to keep the end-of-document insert. ``signature_keywords`` use the keyword
    grammar of :func:`reviewkit.anchors.find_signature_block_start`
    (case-insensitive whole words; trailing ``*`` marks a stem) and feed the
    shared signature scan.
    """

    def __init__(
        self,
        document: DocxDocument,
        *,
        signature_keywords: Sequence[str] = (),
        resolve_last_anchor: Callable[[InsertionAction], str] | None = None,
    ) -> None:
        self.document = document
        self._signature_patterns = _compile_signature_keywords(signature_keywords)
        self._resolve_last_anchor = resolve_last_anchor
        self._mutated = False

    def apply_actions(self, actions: Sequence[InsertionAction]) -> InsertionReport:
        """Apply a batch and return applied results plus failed actions.

        Every anchor is resolved against the pristine document before any
        mutation, so earlier insertions cannot shift later targets, and
        actions sharing an anchor are chained in batch order. That guarantee
        only holds within one batch: once this inserter has mutated the
        document, further calls raise instead of resolving against shifted
        paragraph indices.
        """
        if self._mutated:
            raise ValueError(
                "ParagraphInserter has already mutated the document; anchors "
                "resolve against the pristine document, so further "
                "insertions require a new inserter instance"
            )
        failed: list[InsertionAction] = []
        # Snapshot once: python-docx rebuilds the paragraph wrapper list on
        # every .paragraphs access, which is O(document) per call.
        body_paragraphs = list(self.document.paragraphs)
        resolved_actions = [
            (action, self._resolve_action(action, body_paragraphs)) for action in actions
        ]

        chained_anchors: dict[int, CT_P] = {}
        records: list[InsertionResult] = []
        for action, resolved in resolved_actions:
            if resolved is None:
                failed.append(action)
                continue
            element, predecessor = self._apply_resolved(resolved, chained_anchors)
            if predecessor is not None:
                # The new element now directly follows `predecessor`, so any
                # earlier insertion recorded against that slot is no longer
                # adjacent to it — remap its predecessor to the new element
                # so post-validation checks real adjacency.
                for record in records:
                    if record.predecessor is predecessor:
                        record.predecessor = element
            records.append(InsertionResult(action, element, predecessor))
        if records:
            self._mutated = True
        paragraph_positions = {
            id(paragraph._element): index
            for index, paragraph in enumerate(self.document.paragraphs)
        }
        for record in records:
            record.applied_anchor = _intended_anchor(paragraph_positions, record.predecessor)
        return InsertionReport(applied=records, failed=failed)

    def _resolve_action(
        self, action: InsertionAction, body_paragraphs: list[Paragraph]
    ) -> _ResolvedAction | None:
        text = action.rendered_text()

        anchor = action.anchor
        contextual = anchor == ANCHOR_LAST
        if contextual and self._resolve_last_anchor is not None:
            anchor = self._resolve_last_anchor(action)

        index = parse_body_anchor_index(anchor)
        if index is not None:
            # Body-local resolution only: indices beyond the body must fail the
            # action instead of landing in tables or headers.
            if not 0 <= index < len(body_paragraphs):
                return None
            if contextual:
                # The signature-tail guard is an upper bound on contextual
                # resolution: a heuristic match inside the signature block
                # (e.g. an annex title) must not drop the clause below the
                # signature lines.
                signature_start = self._resolve_end_insertion_point()
                if signature_start is not None and not self._precedes(
                    body_paragraphs[index]._element, signature_start
                ):
                    return _ResolvedAction(
                        text=text, anchor_paragraph=None, insert_before=signature_start
                    )
            return _ResolvedAction(text=text, anchor_paragraph=body_paragraphs[index])
        if anchor == ANCHOR_LAST:
            return _ResolvedAction(
                text=text,
                anchor_paragraph=None,
                insert_before=self._resolve_end_insertion_point(),
            )

        return None

    def _apply_resolved(
        self, resolved: _ResolvedAction, chained_anchors: dict[int, CT_P]
    ) -> tuple[CT_P, BaseOxmlElement | None]:
        if resolved.anchor_paragraph is None:
            return self._insert_at_end(resolved.text, resolved.insert_before)
        anchor_element = resolved.anchor_paragraph._element
        # Later actions on the same anchor insert after the previous insertion
        # so the final document keeps batch order.
        target = chained_anchors.get(id(anchor_element), anchor_element)
        # A paragraph immediately followed by a table is a lead-in for that
        # table; the pair is an adjacency unit. Inserting between them would
        # split the lead-in from its table, so the clause goes after the
        # table instead. In the paragraph sequence (tables are not
        # paragraphs) the clause still directly follows the lead-in, so the
        # recorded predecessor stays adjacency-true for validation.
        insert_after = _skip_following_tables(target)
        new_element = self._insert_after_element(insert_after, resolved.text)
        chained_anchors[id(anchor_element)] = new_element
        return new_element, target

    def _resolve_end_insertion_point(self) -> CT_P | None:
        """Return the element an end insert must precede, or ``None`` to append.

        Runs at resolve time, against the not-yet-mutated body, so clauses
        inserted by earlier actions in the same batch are never mistaken for
        signature blocks.
        """
        return _find_signature_block_element(self.document, (), self._signature_patterns)

    def _precedes(self, element: CT_P, reference: CT_P) -> bool:
        """True when ``element`` appears strictly before ``reference`` in the body."""
        for child in self.document.element.body.iterchildren():
            # Reference first: anchoring on the signature-start element itself
            # would still insert inside the block, so identity is "not before".
            if child is reference:
                return False
            if child is element:
                return True
        return False

    def _insert_at_end(
        self, text: str, insert_before: CT_P | None
    ) -> tuple[CT_P, BaseOxmlElement | None]:
        # Capture the intended predecessor before mutating, for applied-anchor
        # bookkeeping.
        if insert_before is not None:
            predecessor = insert_before.getprevious()
        else:
            predecessor = None
            for child in self.document.element.body.iterchildren():
                if isinstance(child.tag, str) and child.tag.endswith("}p"):
                    predecessor = child

        new_paragraph = self.document.add_paragraph()
        new_paragraph.add_run(text)
        new_element = new_paragraph._element

        if insert_before is not None:
            self.document.element.body.remove(new_element)
            insert_before.addprevious(new_element)
        # Otherwise the paragraph is already in the right place:
        # python-docx inserts new paragraphs before w:sectPr, which must stay
        # the final body child (ECMA-376).

        return new_element, predecessor

    def _insert_after_element(self, target_element: BaseOxmlElement, text: str) -> CT_P:
        body = self.document.element.body

        new_paragraph = self.document.add_paragraph()
        new_paragraph.add_run(text)
        new_element = new_paragraph._element

        body.remove(new_element)
        target_element.addnext(new_element)

        return new_element


class InsertionValidator:
    """Structural post-insertion checks over a saved-and-reopened document.

    Validates per action that the inserted paragraph is actually present at
    its recorded anchor (or, without one, at a position consistent with the
    originally requested anchor), that no insertion landed at or below the
    trailing signature block, and that the document body is structurally
    sound. The per-action checks return booleans and action lists — error
    wording for those is a caller concern — while
    :meth:`check_document_integrity` returns a structural summary dict with
    fixed error strings.

    ``signature_keywords`` use the keyword grammar of
    :func:`reviewkit.anchors.find_signature_block_start` (case-insensitive
    whole words; trailing ``*`` marks a stem) and feed the shared signature
    scan.
    """

    def __init__(
        self,
        document: DocxDocument,
        *,
        signature_keywords: Sequence[str] = (),
    ) -> None:
        self.document = document
        # Snapshot once: python-docx rebuilds the paragraph wrapper list on
        # every .paragraphs access; validation never mutates the document.
        try:
            self._paragraphs = list(document.paragraphs)
        except AttributeError:
            # A document whose w:body is missing raises here; keep the
            # validator constructible so check_document_integrity can report
            # the corruption instead of the constructor crashing.
            self._paragraphs = []
        self._signature_patterns = _compile_signature_keywords(signature_keywords)

    def action_applied(self, action: InsertionAction, applied_anchor: str | None = None) -> bool:
        """True when ``action``'s rendered text sits where the insertion claimed.

        With an ``applied_anchor`` the text must directly follow that
        paragraph in the paragraph sequence. Without one, the check falls
        back to the originally requested anchor: ``suggest`` actions and
        ``body:p:last`` degrade to a document-wide text-exists check,
        ``body:p:<n>`` to the same paragraph-sequence adjacency check.
        """
        text = action.rendered_text()

        if applied_anchor is not None:
            return self._text_follows_anchor(applied_anchor, text)

        if action.kind == "suggest":
            return self._text_exists_in_document(text)

        anchor = action.anchor
        if anchor == ANCHOR_LAST:
            return self._text_exists_in_document(text)

        if parse_body_anchor_index(anchor) is not None:
            # Paragraph-sequence adjacency, same as the applied_anchor path:
            # an XML-sibling check would falsely fail table lead-in
            # placements, where the clause follows the anchor's table rather
            # than the anchor element itself.
            return self._text_follows_anchor(anchor, text)

        return False

    def misplaced_actions(self, actions: Sequence[InsertionAction]) -> list[InsertionAction]:
        """Actions whose inserted text sits at/below the signature block start.

        The inserter's own signature-tail guard defines such placement as
        wrong, so validation must reject it instead of blessing it. Inserted
        texts are excluded from the signature scan so a misplaced clause
        cannot split the block and hide its own misplacement.
        """
        inserted_texts = [
            (action, text)
            for action in actions
            if (text := self._inserted_text(action)) is not None
        ]
        if not inserted_texts:
            return []

        signature_start = _find_signature_block_element(
            self.document,
            [text for _, text in inserted_texts],
            self._signature_patterns,
        )
        if signature_start is None:
            return []
        start_index = next(
            (
                index
                for index, paragraph in enumerate(self._paragraphs)
                if paragraph._element is signature_start
            ),
            None,
        )
        if start_index is None:
            return []

        return [
            action
            for action, text in inserted_texts
            if any(text in paragraph.text for paragraph in self._paragraphs[start_index:])
        ]

    def check_document_integrity(self) -> dict[str, Any]:
        errors: list[str] = []

        if len(self._paragraphs) == 0:
            errors.append("Document has no paragraphs")

        try:
            body = self.document.element.body
        except AttributeError:
            # Narrow catch: a malformed OOXML root may not expose ``body`` at all
            # (lxml raises AttributeError). Any other failure is a real bug and
            # must propagate, not be masked as "corrupted".
            body = None
        if body is None:
            # python-docx returns ``None`` (rather than raising) when the
            # ``w:body`` element is missing; treat that as corruption too, so the
            # validator reports a fail-closed result instead of crashing on the
            # ``NoneType.iterchildren`` access below.
            errors.append("Document body is corrupted")
        else:
            children = list(body.iterchildren())
            for position, child in enumerate(children):
                # Comment/PI nodes expose a callable .tag, not a string (lxml).
                if not isinstance(child.tag, str):
                    continue
                if child.tag.endswith("}sectPr") and position != len(children) - 1:
                    errors.append("Invalid OOXML structure: w:sectPr is not the final body child")
                    break

        return {
            "valid": len(errors) == 0,
            "errors": errors,
            "paragraph_count": len(self._paragraphs),
        }

    def _inserted_text(self, action: InsertionAction) -> str | None:
        if action.kind == "insert":
            return action.text or None
        return action.rendered_text()

    def _text_follows_anchor(self, anchor: str, text: str) -> bool:
        """Assert placement: the text must sit directly after its anchor paragraph."""
        anchor_index = parse_body_anchor_index(anchor)
        if anchor_index is None:
            return False
        paragraphs = self._paragraphs
        inserted_index = anchor_index + 1
        if inserted_index >= len(paragraphs):
            return False
        return text in paragraphs[inserted_index].text

    def _text_exists_in_document(self, text: str) -> bool:
        for paragraph in self._paragraphs:
            if text in paragraph.text:
                return True
        return False
