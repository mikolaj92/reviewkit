from io import BytesIO

from docx import Document

from reviewkit import (
    PortableReviewTrailProfile,
    append_portable_review_trail,
    has_portable_review_trail,
    strip_portable_review_trail,
)

PROFILE = PortableReviewTrailProfile("Review trail", "Portable comments")


def test_portable_trail_no_comments_is_noop() -> None:
    doc = Document()
    doc.add_paragraph("body")
    stream = BytesIO()
    doc.save(stream)
    source = stream.getvalue()
    assert append_portable_review_trail(source, profile=PROFILE) == source
    assert not has_portable_review_trail(source, profile=PROFILE)


def test_strip_absent_trail_is_noop() -> None:
    doc = Document()
    doc.add_paragraph("body")
    stream = BytesIO()
    doc.save(stream)
    result, changed = strip_portable_review_trail(stream.getvalue(), profile=PROFILE)
    assert not changed and result == stream.getvalue()
