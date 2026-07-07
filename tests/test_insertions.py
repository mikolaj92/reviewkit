import io
import re
import zipfile

import pytest
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn

from reviewkit.anchors import ANCHOR_LAST
from reviewkit.insertions import (
    SUGGESTION_MARKER_PREFIX,
    ClauseInserter,
    InsertionAction,
    InsertionValidator,
    contains_suggestion_marker,
    format_suggestion_text,
)

_SIGNATURE_PATTERNS = (
    re.compile(r"\bsignature\b"),
    re.compile(r"\bdate\b"),
    re.compile(r"\battachment\b"),
)


def _make_document(texts: list[str]) -> DocxDocument:
    document = Document()
    for text in texts:
        document.add_paragraph(text)
    return document


def _texts(document: DocxDocument) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def test_insert_places_clause_after_anchor() -> None:
    document = _make_document(["First.", "Second."])
    inserter = ClauseInserter(document)
    report = inserter.apply_actions(
        [InsertionAction(action_id="a1", anchor="body:p:0", text="Inserted clause.")]
    )
    assert not report.failed
    assert _texts(document) == ["First.", "Inserted clause.", "Second."]
    assert report.applied[0].applied_anchor == "body:p:0"


def test_batch_anchors_resolve_against_pristine_document() -> None:
    document = _make_document(["A.", "B.", "C."])
    inserter = ClauseInserter(document)
    report = inserter.apply_actions(
        [
            InsertionAction(action_id="a1", anchor="body:p:0", text="X."),
            InsertionAction(action_id="a2", anchor="body:p:1", text="Y."),
        ]
    )
    assert not report.failed
    assert _texts(document) == ["A.", "X.", "B.", "Y.", "C."]
    assert report.applied[0].applied_anchor == "body:p:0"
    assert report.applied[1].applied_anchor == "body:p:2"


def test_same_anchor_actions_chain_in_batch_order() -> None:
    document = _make_document(["A.", "B."])
    inserter = ClauseInserter(document)
    report = inserter.apply_actions(
        [
            InsertionAction(action_id="a1", anchor="body:p:0", text="First insert."),
            InsertionAction(action_id="a2", anchor="body:p:0", text="Second insert."),
        ]
    )
    assert _texts(document) == ["A.", "First insert.", "Second insert.", "B."]
    assert report.applied[0].applied_anchor == "body:p:0"
    assert report.applied[1].applied_anchor == "body:p:1"


def test_out_of_range_and_unsupported_anchors_fail_in_order() -> None:
    document = _make_document(["Only paragraph."])
    inserter = ClauseInserter(document)
    out_of_range = InsertionAction(action_id="a1", anchor="body:p:9", text="X.")
    unsupported = InsertionAction(action_id="a2", anchor="header:0:p:0", text="Y.")
    report = inserter.apply_actions([out_of_range, unsupported])
    assert report.applied == []
    assert report.failed == [out_of_range, unsupported]
    assert _texts(document) == ["Only paragraph."]


def test_unicode_digit_anchor_fails_action_not_batch() -> None:
    document = _make_document(["A."])
    inserter = ClauseInserter(document)
    bad = InsertionAction(action_id="a1", anchor="body:p:²", text="X.")
    good = InsertionAction(action_id="a2", anchor="body:p:0", text="Y.")
    report = inserter.apply_actions([bad, good])
    assert report.failed == [bad]
    assert _texts(document) == ["A.", "Y."]


def test_failed_only_batch_keeps_inserter_reusable() -> None:
    document = _make_document(["Only paragraph."])
    inserter = ClauseInserter(document)
    first = inserter.apply_actions([InsertionAction(action_id="a1", anchor="body:p:9", text="X.")])
    assert first.applied == []
    second = inserter.apply_actions([InsertionAction(action_id="a2", anchor="body:p:0", text="Y.")])
    assert not second.failed
    assert _texts(document) == ["Only paragraph.", "Y."]


def test_second_apply_after_mutation_raises() -> None:
    document = _make_document(["Only paragraph."])
    inserter = ClauseInserter(document)
    inserter.apply_actions([InsertionAction(action_id="a1", anchor="body:p:0", text="X.")])
    with pytest.raises(ValueError, match="pristine"):
        inserter.apply_actions([InsertionAction(action_id="a2", anchor="body:p:0", text="Y.")])


def test_anchor_last_appends_without_signature_block() -> None:
    document = _make_document(["A.", "B."])
    inserter = ClauseInserter(document)
    report = inserter.apply_actions(
        [InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="Appended clause.")]
    )
    assert _texts(document) == ["A.", "B.", "Appended clause."]
    assert report.applied[0].applied_anchor == "body:p:1"


def test_anchor_last_inserts_above_signature_block() -> None:
    document = _make_document(["Body clause.", "Signature: ____", "Date: ____"])
    inserter = ClauseInserter(document, signature_patterns=_SIGNATURE_PATTERNS)
    report = inserter.apply_actions(
        [InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="New clause.")]
    )
    assert _texts(document) == ["Body clause.", "New clause.", "Signature: ____", "Date: ____"]
    assert report.applied[0].applied_anchor == "body:p:0"


def test_end_insert_batch_keeps_order_even_when_clause_matches_signature_pattern() -> None:
    # "date" in the first clause matches a signature pattern; the signature
    # scan runs at resolve time against the pristine document, so a clause
    # inserted earlier in the batch must never be mistaken for the block.
    document = _make_document(["Body clause.", "Signature: ____"])
    inserter = ClauseInserter(document, signature_patterns=_SIGNATURE_PATTERNS)
    report = inserter.apply_actions(
        [
            InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="Valid until date X."),
            InsertionAction(action_id="a2", anchor=ANCHOR_LAST, text="Second clause."),
        ]
    )
    assert not report.failed
    assert _texts(document) == [
        "Body clause.",
        "Valid until date X.",
        "Second clause.",
        "Signature: ____",
    ]
    assert [result.applied_anchor for result in report.applied] == ["body:p:0", "body:p:1"]


def test_suggestion_action_inserts_formatted_text() -> None:
    document = _make_document(["A."])
    inserter = ClauseInserter(document)
    inserter.apply_actions(
        [
            InsertionAction(
                action_id="a1",
                anchor="body:p:0",
                text="Proposed wording.",
                kind="suggest",
                reason="Missing clause",
            )
        ]
    )
    assert _texts(document) == ["A.", "[SUGGESTION: Missing clause]\nProposed wording."]
    assert format_suggestion_text("r", "t") == "[SUGGESTION: r]\nt"


def test_contains_suggestion_marker_detects_formatted_text() -> None:
    # The detector is the contract a clean-copy purity gate matches against, so anything
    # format_suggestion_text produces must register -- prefix constant and all.
    assert contains_suggestion_marker(format_suggestion_text("Missing clause", "Proposed."))
    assert contains_suggestion_marker(f"prefix {SUGGESTION_MARKER_PREFIX}: r]\nt suffix")


def test_contains_suggestion_marker_is_false_for_plain_text() -> None:
    assert not contains_suggestion_marker("An ordinary paragraph with no marker.")
    assert not contains_suggestion_marker("")


def test_resolver_receives_raw_action_before_formatting() -> None:
    document = _make_document(["A.", "B."])
    seen: list[InsertionAction] = []

    def resolver(action: InsertionAction) -> str:
        seen.append(action)
        return "body:p:0"

    inserter = ClauseInserter(document, resolve_last_anchor=resolver)
    action = InsertionAction(
        action_id="a1",
        anchor=ANCHOR_LAST,
        text="Raw text.",
        kind="suggest",
        reason="Why",
    )
    report = inserter.apply_actions([action])
    assert seen == [action]
    assert seen[0].text == "Raw text."
    assert _texts(document) == ["A.", "[SUGGESTION: Why]\nRaw text.", "B."]
    assert report.applied[0].action.anchor == ANCHOR_LAST


def test_resolver_is_not_called_for_concrete_anchors() -> None:
    document = _make_document(["A."])

    def resolver(action: InsertionAction) -> str:
        raise AssertionError("resolver must only run for body:p:last anchors")

    inserter = ClauseInserter(document, resolve_last_anchor=resolver)
    report = inserter.apply_actions([InsertionAction(action_id="a1", anchor="body:p:0", text="X.")])
    assert not report.failed


def test_contextual_resolution_into_signature_block_falls_back_above_it() -> None:
    document = _make_document(["Body clause.", "Signature: ____", "Date: ____"])
    inserter = ClauseInserter(
        document,
        signature_patterns=_SIGNATURE_PATTERNS,
        resolve_last_anchor=lambda action: "body:p:1",
    )
    report = inserter.apply_actions(
        [InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="New clause.")]
    )
    assert _texts(document) == ["Body clause.", "New clause.", "Signature: ____", "Date: ____"]
    assert report.applied[0].applied_anchor == "body:p:0"


def test_contextual_resolution_above_signature_block_is_honored() -> None:
    document = _make_document(["P0.", "P1.", "Signature: ____"])
    inserter = ClauseInserter(
        document,
        signature_patterns=_SIGNATURE_PATTERNS,
        resolve_last_anchor=lambda action: "body:p:0",
    )
    report = inserter.apply_actions(
        [InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="X.")]
    )
    assert _texts(document) == ["P0.", "X.", "P1.", "Signature: ____"]
    assert report.applied[0].applied_anchor == "body:p:0"


def test_predecessor_remap_keeps_applied_anchors_adjacency_true() -> None:
    document = _make_document(["P0.", "P1.", "P2."])
    inserter = ClauseInserter(document)
    appended = InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="Appended.")
    indexed = InsertionAction(action_id="a2", anchor="body:p:2", text="Indexed.")
    report = inserter.apply_actions([appended, indexed])
    assert _texts(document) == ["P0.", "P1.", "P2.", "Indexed.", "Appended."]
    assert report.applied[0].applied_anchor == "body:p:3"
    assert report.applied[1].applied_anchor == "body:p:2"

    validator = InsertionValidator(document)
    assert validator.action_applied(appended, report.applied[0].applied_anchor)
    assert validator.action_applied(indexed, report.applied[1].applied_anchor)


def test_table_leadin_insertion_lands_after_table() -> None:
    document = _make_document(["Lead-in."])
    document.add_table(rows=1, cols=1)
    document.add_paragraph("After.")
    inserter = ClauseInserter(document)
    report = inserter.apply_actions(
        [InsertionAction(action_id="a1", anchor="body:p:0", text="Clause.")]
    )
    assert _texts(document) == ["Lead-in.", "Clause.", "After."]
    tags = [
        child.tag.rsplit("}", 1)[-1]
        for child in document.element.body.iterchildren()
        if isinstance(child.tag, str)
    ]
    assert tags == ["p", "tbl", "p", "p", "sectPr"]
    assert report.applied[0].applied_anchor == "body:p:0"
    validator = InsertionValidator(document)
    assert validator.action_applied(report.applied[0].action, report.applied[0].applied_anchor)


def test_multi_table_leadin_insertion_lands_after_all_tables() -> None:
    document = _make_document(["Lead-in."])
    document.add_table(rows=1, cols=1)
    document.add_table(rows=1, cols=1)
    document.add_paragraph("After.")
    inserter = ClauseInserter(document)
    report = inserter.apply_actions(
        [InsertionAction(action_id="a1", anchor="body:p:0", text="Clause.")]
    )
    assert _texts(document) == ["Lead-in.", "Clause.", "After."]
    tags = [
        child.tag.rsplit("}", 1)[-1]
        for child in document.element.body.iterchildren()
        if isinstance(child.tag, str)
    ]
    assert tags == ["p", "tbl", "tbl", "p", "p", "sectPr"]
    assert report.applied[0].applied_anchor == "body:p:0"


def test_two_runs_produce_identical_document_xml() -> None:
    def run() -> bytes:
        document = _make_document(["Body clause.", "Signature: ____"])
        inserter = ClauseInserter(document, signature_patterns=_SIGNATURE_PATTERNS)
        inserter.apply_actions(
            [
                InsertionAction(action_id="a1", anchor="body:p:0", text="Clause."),
                InsertionAction(
                    action_id="a2",
                    anchor=ANCHOR_LAST,
                    text="Note.",
                    kind="suggest",
                    reason="Check",
                ),
            ]
        )
        buffer = io.BytesIO()
        document.save(buffer)
        # Zip member metadata carries wall-clock timestamps, so whole-archive
        # bytes are not comparable across runs; the document part must be.
        with zipfile.ZipFile(io.BytesIO(buffer.getvalue())) as archive:
            return archive.read("word/document.xml")

    assert run() == run()


def test_validator_rejects_wrong_or_missing_placement() -> None:
    document = _make_document(["A.", "Inserted.", "B."])
    validator = InsertionValidator(document)
    action = InsertionAction(action_id="a1", anchor="body:p:0", text="Inserted.")
    assert validator.action_applied(action, "body:p:0")
    assert not validator.action_applied(action, "body:p:1")
    assert not validator.action_applied(action, "body:p:9")
    assert not validator.action_applied(action, "not-an-anchor")
    missing = InsertionAction(action_id="a2", anchor="body:p:0", text="Nonexistent.")
    assert not validator.action_applied(missing, "body:p:0")


def test_validator_fallback_accepts_table_leadin_placement() -> None:
    # The clause sits after the anchor's table, so it directly follows the
    # anchor in the paragraph sequence even though it is not its XML sibling.
    document = _make_document(["Lead-in."])
    document.add_table(rows=1, cols=1)
    document.add_paragraph("Inserted.")
    validator = InsertionValidator(document)
    action = InsertionAction(action_id="a1", anchor="body:p:0", text="Inserted.")
    assert validator.action_applied(action)


def test_validator_fallbacks_without_applied_anchor() -> None:
    document = _make_document(["A.", "Inserted.", "B."])
    validator = InsertionValidator(document)
    indexed = InsertionAction(action_id="a1", anchor="body:p:0", text="Inserted.")
    assert validator.action_applied(indexed)
    shifted = InsertionAction(action_id="a2", anchor="body:p:1", text="Inserted.")
    assert not validator.action_applied(shifted)
    out_of_range = InsertionAction(action_id="a3", anchor="body:p:9", text="Inserted.")
    assert not validator.action_applied(out_of_range)
    last = InsertionAction(action_id="a4", anchor=ANCHOR_LAST, text="Inserted.")
    assert validator.action_applied(last)
    absent_last = InsertionAction(action_id="a5", anchor=ANCHOR_LAST, text="Nonexistent.")
    assert not validator.action_applied(absent_last)


def test_validator_suggest_fallback_checks_formatted_text() -> None:
    document = _make_document(["A.", "[SUGGESTION: Why]\nProposed."])
    validator = InsertionValidator(document)
    action = InsertionAction(
        action_id="a1", anchor="body:p:0", text="Proposed.", kind="suggest", reason="Why"
    )
    assert validator.action_applied(action)
    wrong_reason = InsertionAction(
        action_id="a2", anchor="body:p:0", text="Proposed.", kind="suggest", reason="Other"
    )
    assert not validator.action_applied(wrong_reason)


def test_validator_suggest_with_applied_anchor_checks_adjacency() -> None:
    document = _make_document(["A.", "[SUGGESTION: Why]\nProposed.", "B."])
    validator = InsertionValidator(document)
    action = InsertionAction(
        action_id="a1", anchor=ANCHOR_LAST, text="Proposed.", kind="suggest", reason="Why"
    )
    assert validator.action_applied(action, "body:p:0")
    assert not validator.action_applied(action, "body:p:1")


def test_validator_flags_suggestion_below_signature_start() -> None:
    document = _make_document(
        ["Body clause.", "Signature: ____", "[SUGGESTION: Why]\nProposed.", "Date: ____"]
    )
    validator = InsertionValidator(document, signature_patterns=_SIGNATURE_PATTERNS)
    misplaced = InsertionAction(
        action_id="a1", anchor=ANCHOR_LAST, text="Proposed.", kind="suggest", reason="Why"
    )
    assert validator.misplaced_actions([misplaced]) == [misplaced]


def test_validator_flags_clause_at_or_below_signature_start() -> None:
    document = _make_document(
        ["Body clause.", "Signature: ____", "Misplaced clause.", "Date: ____"]
    )
    validator = InsertionValidator(document, signature_patterns=_SIGNATURE_PATTERNS)
    misplaced = InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="Misplaced clause.")
    assert validator.misplaced_actions([misplaced]) == [misplaced]


def test_validator_accepts_clause_above_signature_block() -> None:
    document = _make_document(["Body clause.", "Placed clause.", "Signature: ____", "Date: ____"])
    validator = InsertionValidator(document, signature_patterns=_SIGNATURE_PATTERNS)
    placed = InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="Placed clause.")
    assert validator.misplaced_actions([placed]) == []
    assert validator.misplaced_actions([]) == []


def test_validator_without_patterns_flags_nothing() -> None:
    document = _make_document(["Body clause.", "Signature: ____", "Misplaced clause."])
    validator = InsertionValidator(document)
    misplaced = InsertionAction(action_id="a1", anchor=ANCHOR_LAST, text="Misplaced clause.")
    assert validator.misplaced_actions([misplaced]) == []


def test_check_document_integrity_valid() -> None:
    document = _make_document(["A."])
    result = InsertionValidator(document).check_document_integrity()
    assert result == {"valid": True, "errors": [], "paragraph_count": 1}


def test_check_document_integrity_empty_document() -> None:
    document = Document()
    result = InsertionValidator(document).check_document_integrity()
    assert not result["valid"]
    assert "Document has no paragraphs" in result["errors"]
    assert result["paragraph_count"] == 0


def test_check_document_integrity_missing_body() -> None:
    document = Document()
    root = document.element
    root.remove(root.find(qn("w:body")))
    validator = InsertionValidator(document)
    result = validator.check_document_integrity()
    assert not result["valid"]
    assert "Document body is corrupted" in result["errors"]


def test_misplaced_actions_missing_body_returns_empty() -> None:
    document = Document()
    root = document.element
    root.remove(root.find(qn("w:body")))
    validator = InsertionValidator(document, signature_patterns=_SIGNATURE_PATTERNS)
    action = InsertionAction(action_id="a1", anchor="body:p:0", text="X.")
    assert validator.misplaced_actions([action]) == []


def test_check_document_integrity_sectpr_not_last() -> None:
    document = _make_document(["A."])
    body = document.element.body
    sect_pr = next(
        child
        for child in body.iterchildren()
        if isinstance(child.tag, str) and child.tag.endswith("}sectPr")
    )
    body.remove(sect_pr)
    body.insert(0, sect_pr)
    result = InsertionValidator(document).check_document_integrity()
    assert not result["valid"]
    assert "Invalid OOXML structure: w:sectPr is not the final body child" in result["errors"]
