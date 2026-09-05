from pathlib import Path

from docx import Document as DocxDocument
from docxtor import OperationStatus, inventory_review_markup

from reviewkit import DocumentParser
from reviewkit.docxtor_adapter import apply_review_actions
from reviewkit.models import ActionStatus, ReviewAction, ReviewActionType, ReviewScope
from reviewkit.parser_docx import DocxDocumentParser, load_docx


def test_docx_parser_implements_public_adapter_boundary(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    docx = DocxDocument()
    docx.add_paragraph("Old clause")
    docx.save(source)
    parser: DocumentParser = DocxDocumentParser()

    document = parser.parse(source)

    assert document.source_path == source
    assert document.sections[0].paragraphs[0].locator == "body:p:0"


def test_effective_replace_maps_to_typed_docxtor_receipt(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    docx = DocxDocument()
    docx.add_paragraph("Old clause")
    docx.save(source)
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        id="replace-1",
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE_TEXT,
        node_id=paragraph.id,
        original_text="Old",
        replacement_text="New",
        status=ActionStatus.APPLIED,
    )

    result = apply_review_actions(source.read_bytes(), document, [action])

    assert result.operation_ids_by_action == {"replace-1": ("replace-1:revision",)}
    assert len(result.receipt.operations) == 1
    assert result.receipt.operations[0].status is OperationStatus.APPLIED
    inventory = inventory_review_markup(result.data)
    assert len(inventory.revisions) == 2


def test_conflicted_write_keeps_reviewkit_policy_and_is_not_mutated(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    docx = DocxDocument()
    docx.add_paragraph("Old clause")
    docx.save(source)
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        id="conflict-1",
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.DELETE_TEXT,
        node_id=paragraph.id,
        original_text="Old",
        status=ActionStatus.CONFLICT,
    )

    result = apply_review_actions(source.read_bytes(), document, [action])

    assert result.operation_ids_by_action == {}
    assert result.data == source.read_bytes()
