from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from reviewkit import DocxComment
from reviewkit.document import ReviewDocument
from reviewkit.models import RevisionCoverageError, RevisionCoverageState, RevisionLedger
from reviewkit.parser_docx import load_docx
from reviewkit.renderer_docx import render_corrected_docx, render_reviewed_docx

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def test_effective_projection_and_typed_ledger(tmp_path: Path) -> None:
    # Given: a Word document with one anchored source comment and both revision kinds.
    source = tmp_path / "revision-input.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph()
    plain_run = paragraph.add_run("Plain ")
    document.add_comment(
        runs=plain_run,
        text="Source note.",
        author="Source reviewer",
        initials="SR",
    )
    document.save(source)
    _append_revisions(source)

    # When: ReviewKit reads the DOCX through its public parser contract.
    review_document = load_docx(source)

    # Then: the parser exposes the requested typed revision-aware public result.
    ledger = review_document.revision_ledger
    assert review_document.text == "Plain Inserted."
    assert [(entry.kind, entry.text, entry.locator) for entry in ledger.entries] == [
        ("inserted", "Inserted.", "body:p:0"),
        ("deleted", "Deleted.", "body:p:0"),
    ]
    assert ledger.entries[0].revision_id == "1"
    assert ledger.entries[0].author == "Source reviewer"
    assert review_document.comments == [
        DocxComment(
            id="0",
            author="Source reviewer",
            initials="SR",
            text="Source note.",
            locator="body:p:0",
            anchor_text="Plain ",
        )
    ]


def test_supported_revision_coverage_ignores_empty_and_nested_wrappers(tmp_path: Path) -> None:
    # Given: supported revisions whose OOXML wrapper count is larger than their text spans.
    source = tmp_path / "wrapper-noise.docx"
    document = DocxDocument()
    document.add_paragraph("Plain ")
    document.save(source)
    _append_empty_nested_and_split_revisions(source)

    # When: ReviewKit projects effective text and the typed revision ledger.
    review_document = load_docx(source)

    # Then: direct revision-owned text is complete despite empty/nested wrappers and split spans.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert review_document.text == "Plain Inserted."
    assert [(entry.kind.value, entry.revision_id, entry.text) for entry in review_document.revision_ledger.entries] == [
        ("inserted", "1", "Inserted."),
        ("deleted", "4", "Deleted"),
    ]
    output = tmp_path / "reviewed.docx"
    render_reviewed_docx(review_document, [], output)
    assert output.exists()


def test_supported_revision_coverage_matches_text_controls(tmp_path: Path) -> None:
    # Given: a supported insertion whose visible text includes Word control characters.
    source = tmp_path / "revision-controls.docx"
    document = DocxDocument()
    document.add_paragraph("Plain ")
    document.save(source)
    _append_revision_controls(source)

    # When: ReviewKit projects effective text and the typed revision ledger.
    review_document = load_docx(source)

    # Then: tabs, line breaks, and carriage returns are compared as owned text, not wrapper counts.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert [(entry.kind.value, entry.revision_id, entry.text) for entry in review_document.revision_ledger.entries] == [
        ("inserted", "5", "A\tB\nC\nD"),
    ]


@pytest.mark.parametrize("revision_kind", ["pPrChange", "rPrChange"])
def test_empty_formatting_property_changes_do_not_poison_coverage(
    tmp_path: Path, revision_kind: str
) -> None:
    # Given: a DOCX carrying an empty formatting-only property-change snapshot.
    source = tmp_path / f"{revision_kind}.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_property_revision(source, revision_kind)

    # When: ReviewKit parses the source and callers ask either renderer to publish it.
    review_document = load_docx(source)

    # Then: bookkeeping property changes without owned text stay complete and publishable.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert review_document.text == "Plain text."
    for renderer, filename in (
        (render_reviewed_docx, "reviewed.docx"),
        (render_corrected_docx, "corrected.docx"),
    ):
        output = tmp_path / filename
        renderer(review_document, [], output)
        assert output.exists()


def test_empty_formatting_property_change_with_supported_revisions_is_complete(
    tmp_path: Path,
) -> None:
    # Given: text-bearing ins/del plus an empty pPrChange snapshot, as Word emits them together.
    source = tmp_path / "property-and-text-revisions.docx"
    document = DocxDocument()
    document.add_paragraph("Plain ")
    document.save(source)
    _append_revisions(source)
    _append_property_revision(source, "pPrChange")

    # When: ReviewKit projects effective text and the typed revision ledger.
    review_document = load_docx(source)

    # Then: the formatting-only record does not mask or poison the supported text inventory.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert review_document.text == "Plain Inserted."
    assert [(entry.kind.value, entry.text) for entry in review_document.revision_ledger.entries] == [
        ("inserted", "Inserted."),
        ("deleted", "Deleted."),
    ]
    output = tmp_path / "reviewed.docx"
    render_reviewed_docx(review_document, [], output)
    assert output.exists()


def test_rprchange_snapshot_empty_del_toggle_is_complete(tmp_path: Path) -> None:
    """#224: empty w:del inside rPrChange/rPr is a Word property toggle, not text."""
    source = tmp_path / "rprchange-del-toggle.docx"
    document = DocxDocument()
    document.add_paragraph("Plain ")
    document.save(source)
    _append_revisions(source)
    _append_rprchange_with_empty_del_toggle(source)

    review_document = load_docx(source)

    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert review_document.text == "Plain Inserted."
    assert [(entry.kind.value, entry.text) for entry in review_document.revision_ledger.entries] == [
        ("inserted", "Inserted."),
        ("deleted", "Deleted."),
    ]
    output = tmp_path / "reviewed.docx"
    render_reviewed_docx(review_document, [], output)
    assert output.exists()


def test_rprchange_snapshot_property_toggle_with_fonts_is_complete(tmp_path: Path) -> None:
    """#224: rFonts/sz plus empty w:del in the rPr snapshot is still formatting-only."""
    source = tmp_path / "rprchange-fonts-del-toggle.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_rprchange_with_fonts_and_empty_del(source)

    review_document = load_docx(source)

    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    output = tmp_path / "reviewed.docx"
    render_reviewed_docx(review_document, [], output)
    assert output.exists()


@pytest.mark.parametrize("revision_kind", ["pPrChange", "rPrChange"])
def test_property_change_that_owns_text_remains_fail_closed(
    tmp_path: Path, revision_kind: str
) -> None:
    # Given: a property-change record that owns visible text instead of a formatting snapshot.
    source = tmp_path / f"{revision_kind}-owns-text.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_property_revision_with_text(source, revision_kind)

    # When: ReviewKit parses the source and callers ask either renderer to publish it.
    review_document = load_docx(source)

    # Then: a text-bearing property change stays incomplete and unpublished.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    for renderer, filename in (
        (render_reviewed_docx, "reviewed.docx"),
        (render_corrected_docx, "corrected.docx"),
    ):
        output = tmp_path / filename
        with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
            renderer(review_document, [], output)
        assert not output.exists()


def test_property_change_with_nested_revision_remains_fail_closed(tmp_path: Path) -> None:
    # Given: a pPrChange whose snapshot hides a nested insertion.
    source = tmp_path / "property-nested-revision.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_property_revision_with_nested_insertion(source)

    # When: ReviewKit builds its effective input projection.
    review_document = load_docx(source)

    # Then: nested revision grammar inside a property change is refused before output exists.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    output = tmp_path / "reviewed.docx"
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


@pytest.mark.parametrize("revision_kind", ["pPrChange", "rPrChange"])
def test_malformed_property_change_without_snapshot_remains_fail_closed(
    tmp_path: Path, revision_kind: str
) -> None:
    # Given: a property-change wrapper with no matching pPr/rPr snapshot child.
    source = tmp_path / f"{revision_kind}-malformed.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_malformed_property_revision(source, revision_kind)

    # When: ReviewKit parses the source and callers ask either renderer to publish it.
    review_document = load_docx(source)

    # Then: unexpected property-change children stay fail-closed.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    output = tmp_path / "reviewed.docx"
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_incomplete_revision_coverage_refuses_reviewed_output(tmp_path: Path) -> None:
    # Given: a parsed document whose source revision roles cannot be fully projected.
    document = ReviewDocument(
        revision_ledger=RevisionLedger(coverage=RevisionCoverageState.INCOMPLETE)
    )
    output = tmp_path / "reviewed.docx"

    # When / Then: publication fails before it creates an artifact.
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(document, [], output)
    assert not output.exists()


def test_incomplete_revision_coverage_refuses_corrected_output(tmp_path: Path) -> None:
    document = ReviewDocument(
        revision_ledger=RevisionLedger(coverage=RevisionCoverageState.INCOMPLETE)
    )
    output = tmp_path / "corrected.docx"
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_corrected_docx(document, [], output)
    assert not output.exists()


def test_unresolved_source_comment_refuses_reviewed_output(tmp_path: Path) -> None:
    # Given: a source comment body with no Word range anchor.
    source = tmp_path / "unresolved-comment.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.comments.add_comment(text="Unresolved note.", author="Source", initials="S")
    document.save(source)

    # When / Then: parsing marks coverage incomplete and blocks publication.
    review_document = load_docx(source)
    output = tmp_path / "reviewed.docx"
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_mixed_supported_and_unsupported_revisions_fail_closed(tmp_path: Path) -> None:
    # Given: a source with one projected insertion and one unsupported move revision.
    source = tmp_path / "mixed-revisions.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_revisions(source)
    _append_unsupported_move(source)

    # When: ReviewKit builds its effective input projection.
    review_document = load_docx(source)

    # Then: one supported entry cannot mask the unsupported source grammar.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    output = tmp_path / "reviewed.docx"
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_empty_custom_xml_range_bookmarks_do_not_poison_coverage(tmp_path: Path) -> None:
    # Given: a text-bearing insertion plus empty customXml ins/del range bookmarks.
    source = tmp_path / "custom-xml-range.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_revisions(source)
    _append_custom_xml_range(source)

    # When: ReviewKit builds its effective input projection.
    review_document = load_docx(source)

    # Then: identity-only content-control bookmarks are formatting-only (#226).
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE


def test_custom_xml_range_marker_with_children_remains_fail_closed(tmp_path: Path) -> None:
    # Given: a customXml ins-range bookmark that owns a child run.
    source = tmp_path / "custom-xml-range-child.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_custom_xml_range_with_child(source)

    # When: ReviewKit builds its effective input projection.
    review_document = load_docx(source)

    # Then: a range marker that owns children stays fail-closed.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    output = tmp_path / "reviewed.docx"
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_comment_reference_recovers_locator_when_range_start_is_outside_paragraph(
    tmp_path: Path,
) -> None:
    # Given: a comment whose rangeStart sits on sdtContent and rangeEnd inside ins.
    source = tmp_path / "sdt-comment.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Anchored.")
    document.add_comment(
        runs=run,
        text="Source note.",
        author="Source reviewer",
        initials="SR",
    )
    document.save(source)
    _move_comment_start_onto_sdt_content(source)

    # When: ReviewKit reads comments.
    review_document = load_docx(source)

    # Then: commentReference still supplies a paragraph locator (#226).
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert len(review_document.comments) == 1
    assert review_document.comments[0].locator is not None


def test_nested_sdt_comment_paragraph_still_gets_a_locator(tmp_path: Path) -> None:
    # Given: the whole comment paragraph lives inside nested w:sdt, so python-docx
    # omits it from docx.paragraphs / iter_inner_content.
    source = tmp_path / "nested-sdt-comment.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Anchored.")
    document.add_comment(
        runs=run,
        text="Source note.",
        author="Source reviewer",
        initials="SR",
    )
    document.save(source)
    _wrap_comment_paragraph_in_nested_sdt(source)

    # When: ReviewKit reads comments from package XML (#228).
    review_document = load_docx(source)

    # Then: complete markers still yield a locator even though python-docx hid the paragraph.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert len(review_document.comments) == 1
    assert review_document.comments[0].locator is not None
    paragraphs = list(review_document.iter_paragraphs())
    assert [paragraph.text for paragraph in paragraphs] == ["Anchored."]
    assert paragraphs[0].locator == "body:p:0"
    assert paragraphs[0].comments[0].text == "Source note."


def test_deletion_only_paragraph_is_absent_from_effective_review_text(tmp_path: Path) -> None:
    source = tmp_path / "deletion-only.docx"
    document = DocxDocument()
    document.add_paragraph("Visible.")
    deleted = document.add_paragraph("Deleted.")
    run = deleted.runs[0]._r
    deleted._p.remove(run)
    wrapper = OxmlElement("w:del")
    wrapper.set(qn("w:id"), "1")
    wrapper.set(qn("w:author"), "Reviewer")
    text = run.find(qn("w:t"))
    assert text is not None
    text.tag = qn("w:delText")
    wrapper.append(run)
    deleted._p.append(wrapper)
    document.save(source)

    review_document = load_docx(source)

    assert [paragraph.text for paragraph in review_document.iter_paragraphs()] == ["Visible."]
    assert review_document.revision_ledger.coverage == RevisionCoverageState.COMPLETE
    assert [entry.locator for entry in review_document.revision_ledger.entries] == ["body:p:1"]


def test_indirectly_nested_block_revision_remains_fail_closed(tmp_path: Path) -> None:
    # Given: a supported revision wrapper hiding an empty block through a customXml node.
    source = tmp_path / "nested-block-revision.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_indirectly_nested_block_revision(source)

    # When: ReviewKit builds its effective input projection.
    review_document = load_docx(source)

    # Then: unsupported block-level revision grammar is refused before any output exists.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    reviewed_output = tmp_path / "reviewed.docx"
    corrected_output = tmp_path / "corrected.docx"
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], reviewed_output)
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_corrected_docx(review_document, [], corrected_output)
    assert not reviewed_output.exists()
    assert not corrected_output.exists()


def test_direct_block_revision_remains_fail_closed(tmp_path: Path) -> None:
    # Given: a supported revision wrapper containing a block-level paragraph directly.
    source = tmp_path / "direct-block-revision.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_direct_block_revision(source)

    # When: ReviewKit builds its effective input projection.
    review_document = load_docx(source)

    # Then: direct block-level revision grammar is refused before either output exists.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    reviewed_output = tmp_path / "reviewed.docx"
    corrected_output = tmp_path / "corrected.docx"
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], reviewed_output)
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_corrected_docx(review_document, [], corrected_output)
    assert not reviewed_output.exists()
    assert not corrected_output.exists()


def test_duplicate_source_comment_ids_fail_closed(tmp_path: Path) -> None:
    # Given: a source whose two comment bodies claim the same Word identifier.
    source = tmp_path / "duplicate-comment-ids.docx"
    document = DocxDocument()
    first = document.add_paragraph()
    first_run = first.add_run("First.")
    document.add_comment(runs=first_run, text="First note.", author="A", initials="A")
    second = document.add_paragraph()
    second_run = second.add_run("Second.")
    document.add_comment(runs=second_run, text="Second note.", author="B", initials="B")
    document.save(source)
    _duplicate_second_comment_id(source)

    # When: ReviewKit reads the source and attempts to publish a review artifact.
    review_document = load_docx(source)
    output = tmp_path / "reviewed.docx"

    # Then: ambiguous source-comment identity is refused before output creation.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_orphan_source_comment_markers_fail_closed(tmp_path: Path) -> None:
    # Given: a source with balanced comment markers but no matching comment body.
    source = tmp_path / "orphan-comment-markers.docx"
    document = DocxDocument()
    document.add_paragraph("Plain text.")
    document.save(source)
    _append_orphan_comment_markers(source)

    # When: ReviewKit reads the source and attempts to publish a review artifact.
    review_document = load_docx(source)
    output = tmp_path / "reviewed.docx"

    # Then: orphan markers are refused before output creation.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_source_comment_without_reference_marker_fails_closed(tmp_path: Path) -> None:
    # Given: a source comment range whose reference marker was removed.
    source = tmp_path / "missing-comment-reference.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph("Plain text.")
    document.add_comment(runs=paragraph.runs[0], text="Source note.", author="A", initials="A")
    document.save(source)
    _remove_comment_reference(source)

    # When: ReviewKit reads the source and attempts to publish a review artifact.
    review_document = load_docx(source)
    output = tmp_path / "reviewed.docx"

    # Then: a range without its reference is refused before output creation.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_duplicate_source_thread_paragraph_ids_fail_closed(tmp_path: Path) -> None:
    # Given: a threaded source where two comment bodies share one paragraph identity.
    source = tmp_path / "duplicate-thread-paragraph-ids.docx"
    document = DocxDocument()
    for label in ("First", "Second", "Third"):
        paragraph = document.add_paragraph(f"{label}.")
        document.add_comment(
            runs=paragraph.runs[0], text=f"{label} note.", author=label, initials=label[0]
        )
    document.save(source)
    _duplicate_thread_paragraph_id(source)

    # When: ReviewKit reads the source and attempts to publish a review artifact.
    review_document = load_docx(source)
    output = tmp_path / "reviewed.docx"

    # Then: ambiguous thread provenance is refused before output creation.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def test_duplicate_source_thread_entries_fail_closed(tmp_path: Path) -> None:
    # Given: a threaded source with two sidecar entries for one child paragraph identity.
    source = tmp_path / "duplicate-thread-entries.docx"
    document = DocxDocument()
    for label in ("First", "Second"):
        paragraph = document.add_paragraph(f"{label}.")
        document.add_comment(
            runs=paragraph.runs[0], text=f"{label} note.", author=label, initials=label[0]
        )
    document.save(source)
    _duplicate_thread_comment_ex(source)

    # When: ReviewKit reads the source and attempts to publish a review artifact.
    review_document = load_docx(source)
    output = tmp_path / "reviewed.docx"

    # Then: duplicate sidecar identity is refused before output creation.
    assert review_document.revision_ledger.coverage == RevisionCoverageState.INCOMPLETE
    with pytest.raises(RevisionCoverageError, match="coverage is incomplete"):
        render_reviewed_docx(review_document, [], output)
    assert not output.exists()


def _append_revisions(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    paragraph.append(_revision("ins", "Inserted.", "1"))
    paragraph.append(_revision("del", "Deleted.", "2"))
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_empty_nested_and_split_revisions(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    paragraph.append(_revision("ins", "In", "1"))
    paragraph.append(_revision("ins", "serted.", "1"))
    paragraph.append(etree.Element(f"{_W}ins", {f"{_W}id": "2", f"{_W}author": "Source reviewer"}))
    outer = etree.Element(f"{_W}ins", {f"{_W}id": "3", f"{_W}author": "Source reviewer"})
    outer.append(_revision("del", "Deleted", "4"))
    paragraph.append(outer)
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_revision_controls(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    insertion = etree.SubElement(
        paragraph,
        f"{_W}ins",
        {f"{_W}id": "5", f"{_W}author": "Source reviewer"},
    )
    run = etree.SubElement(insertion, f"{_W}r")
    first = etree.SubElement(run, f"{_W}t")
    first.text = "A"
    etree.SubElement(run, f"{_W}tab")
    second = etree.SubElement(run, f"{_W}t")
    second.text = "B"
    etree.SubElement(run, f"{_W}br")
    third = etree.SubElement(run, f"{_W}t")
    third.text = "C"
    etree.SubElement(run, f"{_W}cr")
    fourth = etree.SubElement(run, f"{_W}t")
    fourth.text = "D"
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_property_revision(path: Path, kind: str) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    attributes = {
        f"{_W}id": "7",
        f"{_W}author": "Source reviewer",
    }
    snapshot_kind = "pPr" if kind == "pPrChange" else "rPr"
    if kind == "pPrChange":
        properties = paragraph.find(f"{_W}pPr")
        if properties is None:
            properties = etree.Element(f"{_W}pPr")
            paragraph.insert(0, properties)
    else:
        run = paragraph.find(f"{_W}r")
        assert run is not None
        properties = run.find(f"{_W}rPr")
        if properties is None:
            properties = etree.Element(f"{_W}rPr")
            run.insert(0, properties)
    change = etree.SubElement(properties, f"{_W}{kind}", attributes)
    snapshot = etree.SubElement(change, f"{_W}{snapshot_kind}")
    etree.SubElement(snapshot, f"{_W}{'jc' if kind == 'pPrChange' else 'b'}")
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_rprchange_with_empty_del_toggle(path: Path) -> None:
    """Word run-property toggle: empty w:del inside rPrChange/rPr (#224)."""
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    run = paragraph.find(f"{_W}r")
    assert run is not None
    properties = run.find(f"{_W}rPr")
    if properties is None:
        properties = etree.Element(f"{_W}rPr")
        run.insert(0, properties)
    change = etree.SubElement(
        properties,
        f"{_W}rPrChange",
        {f"{_W}id": "12", f"{_W}author": "Source reviewer"},
    )
    snapshot = etree.SubElement(change, f"{_W}rPr")
    etree.SubElement(
        snapshot,
        f"{_W}del",
        {f"{_W}id": "13", f"{_W}author": "Source reviewer"},
    )
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_rprchange_with_fonts_and_empty_del(path: Path) -> None:
    """rPr snapshot with rFonts/sz plus empty w:del toggle (#224)."""
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    run = paragraph.find(f"{_W}r")
    assert run is not None
    properties = run.find(f"{_W}rPr")
    if properties is None:
        properties = etree.Element(f"{_W}rPr")
        run.insert(0, properties)
    change = etree.SubElement(
        properties,
        f"{_W}rPrChange",
        {f"{_W}id": "14", f"{_W}author": "Source reviewer"},
    )
    snapshot = etree.SubElement(change, f"{_W}rPr")
    etree.SubElement(
        snapshot,
        f"{_W}del",
        {f"{_W}id": "15", f"{_W}author": "Source reviewer"},
    )
    etree.SubElement(snapshot, f"{_W}rFonts", {f"{_W}ascii": "Calibri"})
    etree.SubElement(snapshot, f"{_W}sz", {f"{_W}val": "24"})
    etree.SubElement(snapshot, f"{_W}szCs", {f"{_W}val": "24"})
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_property_revision_with_text(path: Path, kind: str) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    change = etree.SubElement(
        paragraph,
        f"{_W}{kind}",
        {f"{_W}id": "8", f"{_W}author": "Source reviewer"},
    )
    snapshot = etree.SubElement(change, f"{_W}{'pPr' if kind == 'pPrChange' else 'rPr'}")
    run = etree.SubElement(snapshot, f"{_W}r")
    text_node = etree.SubElement(run, f"{_W}t")
    text_node.text = "Hidden."
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_property_revision_with_nested_insertion(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    properties = paragraph.find(f"{_W}pPr")
    if properties is None:
        properties = etree.Element(f"{_W}pPr")
        paragraph.insert(0, properties)
    change = etree.SubElement(
        properties,
        f"{_W}pPrChange",
        {f"{_W}id": "9", f"{_W}author": "Source reviewer"},
    )
    snapshot = etree.SubElement(change, f"{_W}pPr")
    snapshot.append(_revision("ins", "Hidden.", "10"))
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_malformed_property_revision(path: Path, kind: str) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    if kind == "pPrChange":
        properties = paragraph.find(f"{_W}pPr")
        if properties is None:
            properties = etree.Element(f"{_W}pPr")
            paragraph.insert(0, properties)
    else:
        run = paragraph.find(f"{_W}r")
        assert run is not None
        properties = run.find(f"{_W}rPr")
        if properties is None:
            properties = etree.Element(f"{_W}rPr")
            run.insert(0, properties)
    etree.SubElement(
        properties,
        f"{_W}{kind}",
        {f"{_W}id": "11", f"{_W}author": "Source reviewer"},
    )
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_unsupported_move(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    move_from = etree.Element(f"{_W}moveFrom", {f"{_W}id": "3"})
    run = etree.SubElement(move_from, f"{_W}r")
    text_node = etree.SubElement(run, f"{_W}t")
    text_node.text = "Moved."
    paragraph.append(move_from)
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_custom_xml_range(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    paragraph.append(
        etree.Element(
            f"{_W}customXmlInsRangeStart",
            {f"{_W}id": "9", f"{_W}author": "Source reviewer", f"{_W}date": "2026-01-01T00:00:00Z"},
        )
    )
    paragraph.append(etree.Element(f"{_W}customXmlInsRangeEnd", {f"{_W}id": "9"}))
    paragraph.append(
        etree.Element(
            f"{_W}customXmlDelRangeStart",
            {f"{_W}id": "10", f"{_W}author": "Source reviewer", f"{_W}date": "2026-01-01T00:00:00Z"},
        )
    )
    paragraph.append(etree.Element(f"{_W}customXmlDelRangeEnd", {f"{_W}id": "10"}))
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_custom_xml_range_with_child(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    start = etree.Element(f"{_W}customXmlInsRangeStart", {f"{_W}id": "9"})
    run = etree.SubElement(start, f"{_W}r")
    text_node = etree.SubElement(run, f"{_W}t")
    text_node.text = "Hidden."
    paragraph.append(start)
    paragraph.append(etree.Element(f"{_W}customXmlInsRangeEnd", {f"{_W}id": "9"}))
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _wrap_comment_paragraph_in_nested_sdt(path: Path) -> None:
    """Hide the comment paragraph inside nested w:sdt so python-docx omits it."""
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    reference = root.find(f".//{_W}commentReference")
    assert reference is not None
    paragraph = reference
    while paragraph is not None and etree.QName(paragraph).localname != "p":
        paragraph = paragraph.getparent()
    assert paragraph is not None
    parent = paragraph.getparent()
    assert parent is not None
    index = list(parent).index(paragraph)
    inner = etree.Element(f"{_W}sdt")
    inner_content = etree.SubElement(inner, f"{_W}sdtContent")
    outer = etree.Element(f"{_W}sdt")
    outer_content = etree.SubElement(outer, f"{_W}sdtContent")
    parent.remove(paragraph)
    inner_content.append(paragraph)
    outer_content.append(inner)
    parent.insert(index, outer)
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _move_comment_start_onto_sdt_content(path: Path) -> None:
    """Park commentRangeStart on sdtContent; leave end and reference on the paragraph."""
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    start = root.find(f".//{_W}commentRangeStart")
    assert start is not None
    body = root.find(f"{_W}body")
    assert body is not None
    sdt = etree.Element(f"{_W}sdt")
    content = etree.SubElement(sdt, f"{_W}sdtContent")
    parent = start.getparent()
    assert parent is not None
    parent.remove(start)
    content.append(start)
    body.insert(0, sdt)
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_indirectly_nested_block_revision(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    insertion = etree.SubElement(
        paragraph,
        f"{_W}ins",
        {f"{_W}id": "6", f"{_W}author": "Source reviewer"},
    )
    custom_xml = etree.SubElement(insertion, f"{_W}customXml")
    etree.SubElement(custom_xml, f"{_W}p")
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _append_direct_block_revision(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    insertion = etree.SubElement(
        paragraph,
        f"{_W}ins",
        {f"{_W}id": "8", f"{_W}author": "Source reviewer"},
    )
    etree.SubElement(insertion, f"{_W}p")
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _duplicate_second_comment_id(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    comments_xml = next(data for info, data in entries if info.filename == "word/comments.xml")
    root = etree.fromstring(comments_xml)
    comments = root.findall(f"{_W}comment")
    assert len(comments) == 2
    comments[1].set(f"{_W}id", comments[0].get(f"{_W}id") or "0")
    revised_comments_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_comments_xml if info.filename == "word/comments.xml" else data,
            )


def _append_orphan_comment_markers(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    paragraph.extend(
        [
            etree.Element(f"{_W}commentRangeStart", {f"{_W}id": "99"}),
            etree.Element(f"{_W}commentRangeEnd", {f"{_W}id": "99"}),
            etree.Element(f"{_W}commentReference", {f"{_W}id": "99"}),
        ]
    )
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _remove_comment_reference(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    document_xml = next(data for info, data in entries if info.filename == "word/document.xml")
    root = etree.fromstring(document_xml)
    references = root.findall(f".//{_W}commentReference")
    assert len(references) == 1
    references[0].getparent().remove(references[0])
    revised_document_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_document_xml if info.filename == "word/document.xml" else data,
            )


def _duplicate_thread_paragraph_id(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    comments_xml = next(data for info, data in entries if info.filename == "word/comments.xml")
    root = etree.fromstring(comments_xml)
    comments = root.findall(f"{_W}comment")
    assert len(comments) == 3
    for comment, para_id in zip(comments, ("AAAA0001", "BBBB0002", "AAAA0001"), strict=True):
        paragraph = comment.find(f"{_W}p")
        assert paragraph is not None
        paragraph.set("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", para_id)
    revised_comments_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    extended_xml = (
        b'<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
        b'<w15:commentEx w15:paraId="BBBB0002" w15:paraIdParent="AAAA0001"/>'
        b"</w15:commentsEx>"
    )
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_comments_xml if info.filename == "word/comments.xml" else data,
            )
        archive.writestr("word/commentsExtended.xml", extended_xml)


def _duplicate_thread_comment_ex(path: Path) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    comments_xml = next(data for info, data in entries if info.filename == "word/comments.xml")
    root = etree.fromstring(comments_xml)
    comments = root.findall(f"{_W}comment")
    assert len(comments) == 2
    for comment, para_id in zip(comments, ("AAAA0001", "BBBB0002"), strict=True):
        paragraph = comment.find(f"{_W}p")
        assert paragraph is not None
        paragraph.set("{http://schemas.microsoft.com/office/word/2010/wordml}paraId", para_id)
    revised_comments_xml = etree.tostring(root, encoding="UTF-8", xml_declaration=True)
    extended_xml = (
        b'<w15:commentsEx xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml">'
        b'<w15:commentEx w15:paraId="AAAA0001"/> '
        b'<w15:commentEx w15:paraId="BBBB0002" w15:paraIdParent="AAAA0001"/> '
        b'<w15:commentEx w15:paraId="BBBB0002" w15:paraIdParent="AAAA0001"/>'
        b"</w15:commentsEx>"
    )
    with ZipFile(path, "w") as archive:
        for info, data in entries:
            archive.writestr(
                info,
                revised_comments_xml if info.filename == "word/comments.xml" else data,
            )
        archive.writestr("word/commentsExtended.xml", extended_xml)


def _revision(kind: str, text: str, revision_id: str) -> etree._Element:
    revision = etree.Element(f"{_W}{kind}")
    revision.set(f"{_W}id", revision_id)
    revision.set(f"{_W}author", "Source reviewer")
    run = etree.SubElement(revision, f"{_W}r")
    text_node = etree.SubElement(run, f"{_W}{'t' if kind == 'ins' else 'delText'}")
    text_node.text = text
    return revision
