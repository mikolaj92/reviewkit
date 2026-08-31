"""Existing Word comments are review input: structured parse + render preserve."""

from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree

from docx import Document as DocxDocument
from lxml import etree

from reviewkit.comments import DocxComment, comments_for_locator, read_comments
from reviewkit.models import ReviewAction, ReviewActionType, ReviewScope
from reviewkit.parser_docx import load_docx
from reviewkit.renderer_docx import render_reviewed_docx

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W14 = "{http://schemas.microsoft.com/office/word/2010/wordml}"
_REL = "{http://schemas.openxmlformats.org/package/2006/relationships}"
_CT = "{http://schemas.openxmlformats.org/package/2006/content-types}"
_W15 = "http://schemas.microsoft.com/office/word/2012/wordml"
_COMMENTS_EXTENDED_TYPE = "application/vnd.ms-word.commentsExtended+xml"
_COMMENTS_EXTENDED_REL = (
    "http://schemas.microsoft.com/office/2011/relationships/commentsExtended"
)
_PEOPLE_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.people+xml"
)
_PEOPLE_REL = (
    "http://schemas.microsoft.com/office/2011/relationships/people"
)

_CLAUSE = (
    "Umowa zostaje zawarta na czas nieokreślony. "
    "Okres wypowiedzenia wynosi 3 miesiące."
)
_LAWYER_NOTE = "sprawdzić czy to jest w umowie"


def _clause_docx(
    path: Path,
    *,
    clause: str = _CLAUSE,
    note: str = _LAWYER_NOTE,
    author: str = "Prawnik",
    initials: str = "PR",
) -> Path:
    docx = DocxDocument()
    paragraph = docx.add_paragraph(clause)
    docx.add_comment(runs=paragraph.runs[0], text=note, author=author, initials=initials)
    docx.add_paragraph("Wynagrodzenie wynosi 10000 PLN.")
    docx.save(path)
    return path


def _inject_comment_thread(path: Path) -> Path:
    """Give the source a parent/reply thread part the way Word stores replies."""
    with ZipFile(path) as bundle:
        comments_xml = bundle.read("word/comments.xml")
        document_rels = bundle.read("word/_rels/document.xml.rels")
        content_types = bundle.read("[Content_Types].xml")
        other = {
            name: bundle.read(name)
            for name in bundle.namelist()
            if name
            not in {"word/comments.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"}
        }

    comments_root = etree.fromstring(comments_xml)
    comment_paras = comments_root.findall(f"{_W}comment/{_W}p")
    assert comment_paras, "expected at least one comment paragraph"
    comment_paras[0].set(f"{_W14}paraId", "AAAA0001")

    reply = etree.SubElement(comments_root, f"{_W}comment")
    reply.set(f"{_W}id", "1")
    reply.set(f"{_W}author", "Partner")
    reply.set(f"{_W}initials", "PA")
    reply_p = etree.SubElement(reply, f"{_W}p")
    reply_p.set(f"{_W14}paraId", "BBBB0002")
    reply_r = etree.SubElement(reply_p, f"{_W}r")
    reply_t = etree.SubElement(reply_r, f"{_W}t")
    reply_t.text = "zgadzam się, do weryfikacji"

    extended = (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        f'<w15:commentsEx xmlns:w15="{_W15}">'
        '<w15:commentEx w15:paraId="AAAA0001" w15:done="0"/>'
        '<w15:commentEx w15:paraId="BBBB0002" w15:paraIdParent="AAAA0001" w15:done="0"/>'
        "</w15:commentsEx>"
    ).encode("utf-8")

    rels_root = etree.fromstring(document_rels)
    rel = etree.SubElement(rels_root, f"{_REL}Relationship")
    rel.set("Id", "rIdCommentEx")
    rel.set("Type", _COMMENTS_EXTENDED_REL)
    rel.set("Target", "commentsExtended.xml")

    types_root = etree.fromstring(content_types)
    override = etree.SubElement(types_root, f"{_CT}Override")
    override.set("PartName", "/word/commentsExtended.xml")
    override.set("ContentType", _COMMENTS_EXTENDED_TYPE)

    with ZipFile(path, "w") as bundle:
        for name, data in other.items():
            bundle.writestr(name, data)
        bundle.writestr("word/comments.xml", etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8"))
        bundle.writestr("word/_rels/document.xml.rels", etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8"))
        bundle.writestr("[Content_Types].xml", etree.tostring(types_root, xml_declaration=True, encoding="UTF-8"))
        bundle.writestr("word/commentsExtended.xml", extended)
    return path


def _comment_markers(path: Path) -> list[tuple[str, str]]:
    root = ElementTree.fromstring(ZipFile(path).read("word/document.xml"))
    markers: list[tuple[str, str]] = []
    for element in root.iter():
        local = element.tag.split("}")[-1]
        if local in {"commentRangeStart", "commentRangeEnd", "commentReference"}:
            markers.append((local, element.get(_W + "id") or ""))
    return markers


def test_read_comments_exposes_anchor_text_author_and_locator(tmp_path: Path) -> None:
    path = _clause_docx(tmp_path / "clause.docx")

    comments = read_comments(path)

    assert len(comments) == 1
    comment = comments[0]
    assert comment.text == _LAWYER_NOTE
    assert comment.author == "Prawnik"
    assert comment.initials == "PR"
    assert comment.locator == "body:p:0"
    assert comment.anchor_text == _CLAUSE
    assert "w:comment" not in comment.text


def test_load_docx_attaches_comment_to_the_clause_paragraph(tmp_path: Path) -> None:
    path = _clause_docx(tmp_path / "clause.docx")

    document = load_docx(path)
    clause = document.sections[0].paragraphs[0]
    other = document.sections[0].paragraphs[1]

    assert document.comments
    assert document.comments[0].text == _LAWYER_NOTE
    assert clause.text == _CLAUSE
    assert [comment.text for comment in clause.comments] == [_LAWYER_NOTE]
    assert clause.comments[0].anchor_text == _CLAUSE
    assert other.comments == []
    assert comments_for_locator(document.comments, clause.locator) == clause.comments


def test_render_reviewed_docx_keeps_existing_comment_and_adds_review(tmp_path: Path) -> None:
    path = _clause_docx(tmp_path / "clause.docx")
    document = load_docx(path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.RISK,
        node_id=document.sections[0].paragraphs[0].id,
        original_text="3 miesiące",
        comment="Duration may be unenforceable.",
    )

    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")

    rendered = DocxDocument(str(reviewed))
    texts = [comment.text for comment in rendered.comments]
    assert _LAWYER_NOTE in texts
    assert any("Duration may be unenforceable." in text for text in texts)
    ids = {marker[1] for marker in _comment_markers(reviewed)}
    assert "0" in ids
    assert "1" in ids
    starts = [comment_id for kind, comment_id in _comment_markers(reviewed) if kind == "commentRangeStart"]
    ends = [comment_id for kind, comment_id in _comment_markers(reviewed) if kind == "commentRangeEnd"]
    assert starts.count("0") == 1
    assert ends.count("0") == 1


def test_render_reviewed_docx_preserves_comment_thread_parts(tmp_path: Path) -> None:
    path = _inject_comment_thread(_clause_docx(tmp_path / "thread.docx"))
    comments = read_comments(path)
    by_id = {comment.id: comment for comment in comments}
    assert by_id["0"].text == _LAWYER_NOTE
    assert by_id["1"].text == "zgadzam się, do weryfikacji"
    assert by_id["1"].parent_id == "0"
    assert by_id["0"].parent_id is None

    document = load_docx(path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.COMMENT,
        node_id=document.sections[0].paragraphs[0].id,
        comment="Reviewer follows the existing thread.",
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "thread-reviewed.docx")

    with ZipFile(reviewed) as bundle:
        names = set(bundle.namelist())
        assert "word/commentsExtended.xml" in names
        extended = bundle.read("word/commentsExtended.xml")
    assert b"BBBB0002" in extended
    assert b"AAAA0001" in extended
    rendered = DocxDocument(str(reviewed))
    texts = [comment.text for comment in rendered.comments]
    assert _LAWYER_NOTE in texts
    assert "zgadzam się, do weryfikacji" in texts
    assert any("Reviewer follows the existing thread." in text for text in texts)


def test_read_comments_table_cell_gets_table_locator(tmp_path: Path) -> None:
    path = tmp_path / "table.docx"
    docx = DocxDocument()
    table = docx.add_table(rows=1, cols=1)
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_paragraph.add_run("Treść w tabeli.")
    docx.add_comment(runs=cell_paragraph.runs[0], text="Uwaga do komórki.", author="Prawnik")
    docx.save(path)

    comments = read_comments(path)
    assert len(comments) == 1
    assert comments[0].locator == "table:0:r:0:c:0:p:0"
    assert comments[0].anchor_text == "Treść w tabeli."


def test_comments_for_locator_filters(tmp_path: Path) -> None:
    comments = [
        DocxComment(id="0", author="A", initials="A", text="one", locator="body:p:0", anchor_text="x"),
        DocxComment(id="1", author="B", initials="B", text="two", locator="body:p:1", anchor_text="y"),
    ]
    assert [comment.id for comment in comments_for_locator(comments, "body:p:1")] == ["1"]
    assert comments_for_locator(comments, None) == []

def _people_only_docx(path: Path) -> Path:
    """Tracked-revision package with people.xml and no comments.xml (#222)."""
    docx = DocxDocument()
    docx.add_paragraph(_CLAUSE)
    docx.save(path)
    with ZipFile(path) as bundle:
        members = {name: bundle.read(name) for name in bundle.namelist()}
    people = (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        f'<w15:people xmlns:w15="{_W15}">'
        '<w15:person w15:author="Reviewer">'
        '<w15:presenceInfo w15:providerId="None" w15:userId="reviewer"/>'
        "</w15:person>"
        "</w15:people>"
    ).encode("utf-8")
    rels_root = etree.fromstring(members["word/_rels/document.xml.rels"])
    rel = etree.SubElement(rels_root, f"{_REL}Relationship")
    rel.set("Id", "rIdPeople")
    rel.set("Type", _PEOPLE_REL)
    rel.set("Target", "people.xml")
    types_root = etree.fromstring(members["[Content_Types].xml"])
    override = etree.SubElement(types_root, f"{_CT}Override")
    override.set("PartName", "/word/people.xml")
    override.set("ContentType", _PEOPLE_TYPE)
    members["word/people.xml"] = people
    members["word/_rels/document.xml.rels"] = etree.tostring(
        rels_root, xml_declaration=True, encoding="UTF-8"
    )
    members["[Content_Types].xml"] = etree.tostring(
        types_root, xml_declaration=True, encoding="UTF-8"
    )
    with ZipFile(path, "w") as bundle:
        for name, data in members.items():
            bundle.writestr(name, data)
    return path


def test_render_reviewed_docx_creates_comments_when_source_has_people_only(
    tmp_path: Path,
) -> None:
    """#222: people.xml without comments.xml is valid Word, not a KeyError."""
    path = _people_only_docx(tmp_path / "people-only.docx")
    with ZipFile(path) as bundle:
        names = set(bundle.namelist())
    assert "word/people.xml" in names
    assert "word/comments.xml" not in names

    document = load_docx(path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.COMMENT,
        node_id=document.sections[0].paragraphs[0].id,
        comment="Reviewer note on a source that never had comments.",
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "people-reviewed.docx")

    with ZipFile(reviewed) as bundle:
        names = set(bundle.namelist())
        comments_xml = bundle.read("word/comments.xml")
        people_xml = bundle.read("word/people.xml")
    assert "word/people.xml" in names
    assert b"Reviewer note on a source that never had comments." in comments_xml
    assert b"Reviewer" in people_xml
    rendered = DocxDocument(str(reviewed))
    assert any("never had comments" in comment.text for comment in rendered.comments)
