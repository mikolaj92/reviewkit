from pathlib import Path
from zipfile import ZIP_DEFLATED, BadZipFile, ZipFile

import pytest
import reviewkit.revision_package as revision_package
import reviewkit.revisions as revisions_module
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reviewkit import RejectRevisionsError, reject_all_revisions
from reviewkit.markup_purity import inspect_markup
from reviewkit.models import ReviewAction, ReviewActionType, ReviewScope
from reviewkit.parser_docx import load_docx
from reviewkit.renderer_docx import render_reviewed_docx


def _saved_docx(path: Path, text: str) -> Path:
    document = DocxDocument()
    document.add_paragraph(text)
    document.save(path)
    return path


def _reviewed_replacement(path: Path, output: Path) -> Path:
    document = load_docx(path)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE_TEXT,
        node_id=paragraph.id,
        original_text="old",
        replacement_text="new",
    )
    return render_reviewed_docx(document, [action], output)


def _add_inserted_text(paragraph, text: str) -> None:
    insertion = OxmlElement("w:ins")
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    insertion.append(run)
    paragraph._p.append(insertion)


def _add_numbering(paragraph) -> None:
    numbering = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    numbering_id = OxmlElement("w:numId")
    numbering_id.set(qn("w:val"), "1")
    numbering.extend((level, numbering_id))
    paragraph._p.get_or_add_pPr().append(numbering)


def _replace_document_xml(path: Path, document_xml: bytes) -> None:
    with ZipFile(path) as archive:
        entries = [(info, archive.read(info.filename)) for info in archive.infolist()]
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        for info, data in entries:
            archive.writestr(info, document_xml if info.filename == "word/document.xml" else data)


def test_reject_all_revisions_is_byte_reproducible(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path / "source.docx", "old clause")
    reviewed = _reviewed_replacement(source, tmp_path / "reviewed.docx")

    first = reject_all_revisions(reviewed, tmp_path / "first.docx")
    second = reject_all_revisions(reviewed, tmp_path / "second.docx")

    assert first.read_bytes() == second.read_bytes()


def test_reject_all_revisions_in_place(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path / "source.docx", "old clause")
    reviewed = _reviewed_replacement(source, tmp_path / "reviewed.docx")

    result = reject_all_revisions(reviewed, reviewed)

    assert result == reviewed
    assert inspect_markup(reviewed).is_clean
    assert DocxDocument(reviewed).paragraphs[0].text == "old clause"


def test_reject_all_revisions_comment_policy(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path / "source.docx", "Clause")
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.COMMENT,
        node_id=paragraph.id,
        comment="Reviewer note",
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")

    stripped = reject_all_revisions(reviewed, tmp_path / "stripped.docx")
    preserved = reject_all_revisions(reviewed, tmp_path / "preserved.docx", drop_comments=False)

    assert inspect_markup(stripped).is_clean
    assert inspect_markup(preserved).has_comments


@pytest.mark.parametrize("change_name", ["cellIns", "cellMerge"])
def test_reject_all_revisions_fails_closed_on_cell_structure(
    tmp_path: Path, change_name: str
) -> None:
    path = tmp_path / "table.docx"
    document = DocxDocument()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "cell"
    cell._tc.get_or_add_tcPr().append(OxmlElement(f"w:{change_name}"))
    document.save(path)

    with pytest.raises(RejectRevisionsError, match=change_name):
        reject_all_revisions(path, tmp_path / "out.docx")

    assert not (tmp_path / "out.docx").exists()


def test_reject_all_revisions_fails_closed_on_unsupported_property_change(
    tmp_path: Path,
) -> None:
    path = tmp_path / "table.docx"
    document = DocxDocument()
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    change = OxmlElement("w:tblPrChange")
    original = OxmlElement("w:tblPr")
    style = OxmlElement("w:tblStyle")
    style.set(qn("w:val"), "OriginalStyle")
    original.append(style)
    change.append(original)
    table._tbl.tblPr.append(change)
    document.save(path)

    with pytest.raises(RejectRevisionsError, match="tblPrChange"):
        reject_all_revisions(path, tmp_path / "out.docx")

    assert not (tmp_path / "out.docx").exists()


def test_reject_all_revisions_enforces_uncompressed_package_budget(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = _saved_docx(tmp_path / "source.docx", "old clause")
    monkeypatch.setattr(revision_package, "MAX_TOTAL_UNCOMPRESSED_BYTES", 1)

    with pytest.raises(RejectRevisionsError, match="uncompressed size"):
        reject_all_revisions(source, tmp_path / "out.docx")

    assert not (tmp_path / "out.docx").exists()


def test_reject_all_revisions_rejects_doctype(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path / "source.docx", "old clause")
    reviewed = _reviewed_replacement(source, tmp_path / "reviewed.docx")
    with ZipFile(reviewed) as archive:
        document_xml = archive.read("word/document.xml")
    document_xml = document_xml.replace(
        b"<w:document",
        b'<!DOCTYPE document [<!ENTITY payload "expanded">]><w:document',
        1,
    ).replace(b">new<", b">&payload;<", 1)
    _replace_document_xml(reviewed, document_xml)

    with pytest.raises(RejectRevisionsError, match="DOCTYPE"):
        reject_all_revisions(reviewed, tmp_path / "out.docx")

    assert not (tmp_path / "out.docx").exists()


def test_parse_xml_rejects_utf16_doctype() -> None:
    xml = (
        '<?xml version="1.0" encoding="UTF-16"?>'
        '<!DOCTYPE document [<!ENTITY payload "expanded">]>'
        "<document>&payload;</document>"
    ).encode("utf-16")

    with pytest.raises(revision_package.RevisionPackageError, match="DOCTYPE"):
        revision_package.parse_xml(xml)


def test_reject_all_revisions_handles_arbitrary_namespace_prefix(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path / "source.docx", "old clause")
    reviewed = _reviewed_replacement(source, tmp_path / "reviewed.docx")
    with ZipFile(reviewed) as archive:
        document_xml = archive.read("word/document.xml")
    document_xml = document_xml.replace(b"xmlns:w=", b"xmlns:x=").replace(b"w:", b"x:")
    _replace_document_xml(reviewed, document_xml)

    restored = reject_all_revisions(reviewed, tmp_path / "out.docx")

    assert inspect_markup(restored).is_clean
    assert DocxDocument(restored).paragraphs[0].text == "old clause"


def test_accept_all_revisions_preserves_existing_output_on_late_failure(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = _saved_docx(tmp_path / "source.docx", "old clause")
    reviewed = _reviewed_replacement(source, tmp_path / "reviewed.docx")
    destination = tmp_path / "destination.docx"
    destination.write_bytes(b"SENTINEL")

    def fail_inspection(_path: Path) -> None:
        raise RuntimeError("late validation failure")

    monkeypatch.setattr(revisions_module, "inspect_markup", fail_inspection)

    with pytest.raises(RuntimeError, match="late validation failure"):
        revisions_module.accept_all_revisions(reviewed, destination)

    assert destination.read_bytes() == b"SENTINEL"


def test_accept_all_revisions_preserves_in_place_source_on_late_failure(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    source = _saved_docx(tmp_path / "source.docx", "old clause")
    reviewed = _reviewed_replacement(source, tmp_path / "reviewed.docx")
    before = reviewed.read_bytes()

    def fail_inspection(_path: Path) -> None:
        raise RuntimeError("late validation failure")

    monkeypatch.setattr(revisions_module, "inspect_markup", fail_inspection)

    with pytest.raises(RuntimeError, match="late validation failure"):
        revisions_module.accept_all_revisions(reviewed, reviewed)

    assert reviewed.read_bytes() == before


@pytest.mark.parametrize("change_name", ["cellIns", "cellMerge"])
def test_accept_all_revisions_fails_closed_on_cell_structure(
    tmp_path: Path, change_name: str
) -> None:
    path = tmp_path / "table.docx"
    document = DocxDocument()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "cell"
    cell._tc.get_or_add_tcPr().append(OxmlElement(f"w:{change_name}"))
    document.save(path)

    with pytest.raises(revisions_module.AcceptRevisionsError, match=change_name):
        revisions_module.accept_all_revisions(path, tmp_path / "out.docx")

    assert not (tmp_path / "out.docx").exists()


def test_accept_all_revisions_fails_closed_on_unmergeable_paragraph_mark(
    tmp_path: Path,
) -> None:
    path = tmp_path / "final.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph("Final paragraph")
    properties = paragraph._p.get_or_add_pPr()
    run_properties = OxmlElement("w:rPr")
    run_properties.append(OxmlElement("w:del"))
    properties.append(run_properties)
    document.save(path)

    with pytest.raises(revisions_module.AcceptRevisionsError, match="no following"):
        revisions_module.accept_all_revisions(path, tmp_path / "out.docx")

    assert not (tmp_path / "out.docx").exists()


def test_reject_all_revisions_fails_closed_on_unmergeable_paragraph_mark(
    tmp_path: Path,
) -> None:
    path = tmp_path / "final.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph("Final paragraph")
    properties = paragraph._p.get_or_add_pPr()
    run_properties = OxmlElement("w:rPr")
    run_properties.append(OxmlElement("w:ins"))
    properties.append(run_properties)
    document.save(path)

    with pytest.raises(RejectRevisionsError, match="no following"):
        reject_all_revisions(path, tmp_path / "out.docx")

    assert not (tmp_path / "out.docx").exists()


@pytest.mark.parametrize("operation_name", ["accept", "reject"])
def test_revision_operations_reject_duplicate_package_members(
    tmp_path: Path, operation_name: str
) -> None:
    source = tmp_path / "duplicate.docx"
    document_xml = (
        b'<w:document xmlns:w="http://schemas.openxmlformats.org/'
        b'wordprocessingml/2006/main"><w:body/></w:document>'
    )
    with pytest.warns(UserWarning, match="Duplicate name"):
        with ZipFile(source, "w") as archive:
            archive.writestr("word/document.xml", document_xml)
            archive.writestr("word/document.xml", document_xml)
    destination = tmp_path / "out.docx"
    operation = (
        revisions_module.accept_all_revisions
        if operation_name == "accept"
        else reject_all_revisions
    )
    error_type = (
        revisions_module.AcceptRevisionsError
        if operation_name == "accept"
        else RejectRevisionsError
    )

    with pytest.raises(error_type, match="duplicate"):
        operation(source, destination)

    assert not destination.exists()


@pytest.mark.parametrize("operation_name", ["accept", "reject"])
def test_revision_operations_leave_no_output_for_non_docx_package(
    tmp_path: Path, operation_name: str
) -> None:
    source = tmp_path / "not-docx.docx"
    with ZipFile(source, "w") as archive:
        archive.writestr("word/styles.xml", b"<styles/>")
    destination = tmp_path / "out.docx"
    operation = (
        revisions_module.accept_all_revisions
        if operation_name == "accept"
        else reject_all_revisions
    )

    with pytest.raises(BadZipFile):
        operation(source, destination)

    assert not destination.exists()


def test_reject_all_revisions_removes_inserted_content_control_paragraph(
    tmp_path: Path,
) -> None:
    path = tmp_path / "content-control.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    run_properties = OxmlElement("w:rPr")
    run_properties.append(OxmlElement("w:ins"))
    properties.append(run_properties)
    content_control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    insertion = OxmlElement("w:ins")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Inserted controlled clause"
    run.append(text)
    insertion.append(run)
    content.append(insertion)
    content_control.append(content)
    paragraph._p.append(content_control)
    document.add_paragraph("Following paragraph")
    document.save(path)

    restored = reject_all_revisions(path, tmp_path / "out.docx")

    assert inspect_markup(restored).is_clean
    assert [paragraph.text for paragraph in DocxDocument(restored).paragraphs] == [
        "Following paragraph"
    ]


def test_reject_all_revisions_merges_paragraphs_across_content_controls(
    tmp_path: Path,
) -> None:
    path = tmp_path / "controlled-blocks.docx"
    document = DocxDocument()
    body = document._body._element
    for value, tracked_mark in (("First ", True), ("second", False)):
        control = OxmlElement("w:sdt")
        content = OxmlElement("w:sdtContent")
        paragraph = OxmlElement("w:p")
        if tracked_mark:
            properties = OxmlElement("w:pPr")
            run_properties = OxmlElement("w:rPr")
            run_properties.append(OxmlElement("w:ins"))
            properties.append(run_properties)
            paragraph.append(properties)
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        paragraph.append(run)
        content.append(paragraph)
        control.append(content)
        body.insert(len(body) - 1, control)
    document.save(path)

    restored = reject_all_revisions(path, tmp_path / "out.docx")

    with ZipFile(restored) as archive:
        root = revision_package.parse_xml(archive.read("word/document.xml"))
    paragraphs = list(root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"))
    assert len(paragraphs) == 1
    assert "".join(paragraphs[0].itertext()) == "First second"


def test_reject_all_revisions_removes_inserted_blank_paragraph(tmp_path: Path) -> None:
    path = tmp_path / "blank.docx"
    document = DocxDocument()
    paragraph = document.add_paragraph(" ")
    properties = paragraph._p.get_or_add_pPr()
    run_properties = OxmlElement("w:rPr")
    run_properties.append(OxmlElement("w:ins"))
    properties.append(run_properties)
    deletion = OxmlElement("w:del")
    run = OxmlElement("w:r")
    text = OxmlElement("w:delText")
    text.text = " "
    run.append(text)
    deletion.append(run)
    paragraph._p.append(deletion)
    document.add_paragraph("Following paragraph")
    document.save(path)

    restored = reject_all_revisions(path, tmp_path / "out.docx")

    assert [paragraph.text for paragraph in DocxDocument(restored).paragraphs] == [
        "Following paragraph"
    ]


def test_reject_all_revisions_removes_numbered_shell_from_inserted_text(
    tmp_path: Path,
) -> None:
    path = tmp_path / "reviewed.docx"
    document = DocxDocument()
    document.add_paragraph("Lead-in")
    paragraph = document.add_paragraph()
    _add_numbering(paragraph)
    _add_inserted_text(paragraph, "New numbered point")
    document.add_paragraph("Tail")
    document.save(path)

    restored = reject_all_revisions(path, tmp_path / "out.docx")

    assert [paragraph.text for paragraph in DocxDocument(restored).paragraphs] == [
        "Lead-in",
        "Tail",
    ]


def test_reject_all_revisions_removes_literal_number_linked_to_insertion(
    tmp_path: Path,
) -> None:
    path = tmp_path / "reviewed.docx"
    document = DocxDocument()
    document.add_paragraph("Lead-in")
    paragraph = document.add_paragraph("7.")
    _add_inserted_text(paragraph, " New point")
    document.add_paragraph("Tail")
    document.save(path)

    restored = reject_all_revisions(path, tmp_path / "out.docx")

    assert [paragraph.text for paragraph in DocxDocument(restored).paragraphs] == [
        "Lead-in",
        "Tail",
    ]


def test_reject_all_revisions_ignores_duplicate_paragraph_ids(
    tmp_path: Path,
) -> None:
    path = tmp_path / "reviewed.docx"
    document = DocxDocument()
    unrelated = document.add_paragraph()
    unrelated._p.set(qn("w14:paraId"), "DUPL0001")
    _add_numbering(unrelated)
    inserted = document.add_paragraph()
    inserted._p.set(qn("w14:paraId"), "DUPL0001")
    _add_numbering(inserted)
    _add_inserted_text(inserted, "New numbered point")
    document.add_paragraph("Tail")
    document.save(path)

    restored = reject_all_revisions(path, tmp_path / "out.docx")

    assert [paragraph.text for paragraph in DocxDocument(restored).paragraphs] == ["", "Tail"]


def test_reject_all_revisions_merges_consecutive_inserted_paragraph_marks(
    tmp_path: Path,
) -> None:
    path = tmp_path / "reviewed.docx"
    document = DocxDocument()
    for text in ("One ", "two "):
        paragraph = document.add_paragraph(text)
        properties = paragraph._p.get_or_add_pPr()
        run_properties = OxmlElement("w:rPr")
        run_properties.append(OxmlElement("w:ins"))
        properties.append(run_properties)
    document.add_paragraph("three.")
    document.save(path)

    restored = reject_all_revisions(path, tmp_path / "out.docx")

    assert [paragraph.text for paragraph in DocxDocument(restored).paragraphs] == [
        "One two three."
    ]
