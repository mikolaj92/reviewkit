from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

import pytest
from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reviewkit.markup_purity import inspect_markup
from reviewkit.models import (
    ActionStatus,
    ReviewAction,
    ReviewActionType,
    ReviewLocator,
    ReviewScope,
)
from reviewkit.parser_docx import load_docx
from reviewkit.renderer_docx import render_reviewed_docx
from reviewkit.revisions import (
    AcceptRevisionsError,
    RejectRevisionsError,
    accept_all_revisions,
    apply_reviewed_markup,
    reject_all_revisions,
)

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# --- helpers --------------------------------------------------------------------------


def _saved_docx(tmp_path: Path, name: str, *paragraphs: str) -> Path:
    path = tmp_path / name
    docx = DocxDocument()
    for text in paragraphs:
        docx.add_paragraph(text)
    docx.save(path)
    return path


def _body_paragraph_texts(path: Path) -> list[str]:
    with ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/document.xml"))
    body = root.find(f"{_W}body")
    assert body is not None
    texts: list[str] = []
    for child in body:
        if child.tag == f"{_W}p":
            texts.append("".join(text.text or "" for text in child.iter(f"{_W}t")))
    return texts


# --- accepting inline tracked edits ---------------------------------------------------


def test_accept_all_revisions_applies_inline_replace_and_is_clean(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path, "input.docx", "The quick brown fox jumps.")
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE_TEXT,
        node_id=paragraph.id,
        original_text="fox",
        replacement_text="cat",
        locator=ReviewLocator(node_id=paragraph.id, char_start=16, char_end=19),
        status=ActionStatus.APPLIED,
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    # Sanity: the reviewed copy really does carry tracked markup before we flatten it.
    assert inspect_markup(reviewed).has_tracked_revisions

    corrected = accept_all_revisions(reviewed, tmp_path / "corrected.docx")

    assert inspect_markup(corrected).is_clean
    assert _body_paragraph_texts(corrected) == ["The quick brown cat jumps."]


def test_accept_all_revisions_drops_deleted_text(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path, "input.docx", "Alpha beta gamma.")
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.DELETE_TEXT,
        node_id=paragraph.id,
        original_text="beta",
        locator=ReviewLocator(node_id=paragraph.id, char_start=6, char_end=10),
        status=ActionStatus.APPLIED,
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    corrected = accept_all_revisions(reviewed, tmp_path / "corrected.docx")

    assert inspect_markup(corrected).is_clean
    assert "beta" not in _body_paragraph_texts(corrected)[0]


def test_accept_all_revisions_clean_copy_is_byte_reproducible(tmp_path: Path) -> None:
    # The clean copy must hash identically for identical input, so downstream attestation /
    # caching stays stable. accept_all_revisions used to copy the reviewed input's per-entry
    # zip timestamps through, which carry the wall-clock mtime and defeat that. Assert both the
    # pinned epoch and whole-file byte equality across two runs.
    source = _saved_docx(tmp_path, "input.docx", "The quick brown fox jumps.")
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE_TEXT,
        node_id=paragraph.id,
        original_text="fox",
        replacement_text="cat",
        locator=ReviewLocator(node_id=paragraph.id, char_start=16, char_end=19),
        status=ActionStatus.APPLIED,
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")

    first = accept_all_revisions(reviewed, tmp_path / "corrected-a.docx")
    second = accept_all_revisions(reviewed, tmp_path / "corrected-b.docx")

    with ZipFile(first) as archive:
        assert all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in archive.infolist())
    assert first.read_bytes() == second.read_bytes()


# --- the linchpin: a stand-alone clause survives as its own paragraph ------------------


def test_accept_all_revisions_keeps_new_paragraph_standalone(tmp_path: Path) -> None:
    # End-to-end proof that a stand-alone clause insert becomes a REAL separate
    # paragraph in the clean copy - not glued onto the anchor's text. This is exactly
    # what dike's czystopis relies on when it flattens an auto-applied catalogue clause.
    source = _saved_docx(tmp_path, "input.docx", "Anchor clause heading.", "Following paragraph.")
    document = load_docx(source)
    anchor = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.INSERT_AFTER,
        node_id=anchor.id,
        replacement_text="§20a. The inserted clause.",
        new_paragraph=True,
        status=ActionStatus.APPLIED,
        apply_to_corrected=True,
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    corrected = accept_all_revisions(reviewed, tmp_path / "corrected.docx")

    assert inspect_markup(corrected).is_clean
    assert _body_paragraph_texts(corrected) == [
        "Anchor clause heading.",
        "§20a. The inserted clause.",
        "Following paragraph.",
    ]


def test_accept_all_revisions_multiline_clause_becomes_multiple_paragraphs(
    tmp_path: Path,
) -> None:
    source = _saved_docx(tmp_path, "input.docx", "Anchor.", "Tail.")
    document = load_docx(source)
    anchor = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.INSERT_AFTER,
        node_id=anchor.id,
        replacement_text="Clause line one.\nClause line two.",
        new_paragraph=True,
        status=ActionStatus.APPLIED,
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    corrected = accept_all_revisions(reviewed, tmp_path / "corrected.docx")

    assert inspect_markup(corrected).is_clean
    assert _body_paragraph_texts(corrected) == [
        "Anchor.",
        "Clause line one.",
        "Clause line two.",
        "Tail.",
    ]


# --- comments -------------------------------------------------------------------------


def test_accept_all_revisions_drops_comments_by_default(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path, "input.docx", "The quick brown fox jumps.")
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.COMMENT,
        node_id=paragraph.id,
        comment="A reviewer note.",
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    assert inspect_markup(reviewed).has_comments

    corrected = accept_all_revisions(reviewed, tmp_path / "corrected.docx")

    report = inspect_markup(corrected)
    assert report.is_clean
    assert report.comment_count == 0
    # The clean package contains no comment payload, relationship, content-type entry,
    # or dangling body anchors. Keeping an empty comments.xml still exposes reviewer-only
    # metadata to clients and makes the primary artifact structurally unclean.
    with ZipFile(corrected) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml")
        relationships = archive.read("word/_rels/document.xml.rels")
        content_types = archive.read("[Content_Types].xml")
    assert not any(name.startswith(("word/comments", "word/people.xml")) for name in names)
    assert b"comments" not in relationships
    assert b"comments" not in content_types
    assert b"commentReference" not in document_xml
    assert b"commentRangeStart" not in document_xml


def test_accept_all_revisions_keeps_comments_when_asked(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path, "input.docx", "The quick brown fox jumps.")
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.COMMENT,
        node_id=paragraph.id,
        comment="A reviewer note.",
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")

    corrected = accept_all_revisions(reviewed, tmp_path / "corrected.docx", drop_comments=False)

    report = inspect_markup(corrected)
    # Revisions are always accepted; comments are preserved on request.
    assert not report.has_tracked_revisions
    assert report.has_comments


# --- fail-closed guards ---------------------------------------------------------------


def _paragraph_mark(kind: str, paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    run_properties = OxmlElement("w:rPr")
    mark = OxmlElement(f"w:{kind}")
    mark.set(qn("w:id"), "1")
    mark.set(qn("w:author"), "reviewer")
    mark.set(qn("w:date"), "1970-01-01T00:00:00Z")
    run_properties.append(mark)
    properties.insert(0, run_properties)


def test_accept_all_revisions_merges_paragraph_mark_deletion(tmp_path: Path) -> None:
    path = tmp_path / "merged.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph("Hello ")
    _paragraph_mark("del", paragraph)
    docx.add_paragraph("world.")
    docx.save(path)

    corrected = accept_all_revisions(path, tmp_path / "out.docx")
    assert inspect_markup(corrected).is_clean
    assert _body_paragraph_texts(corrected) == ["Hello world."]


def test_accept_all_revisions_rejects_cell_deletion(tmp_path: Path) -> None:
    path = tmp_path / "table.docx"
    docx = DocxDocument()
    table = docx.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = "cell text"
    tc_properties = cell._tc.get_or_add_tcPr()
    cell_del = OxmlElement("w:cellDel")
    cell_del.set(qn("w:id"), "1")
    cell_del.set(qn("w:author"), "reviewer")
    cell_del.set(qn("w:date"), "1970-01-01T00:00:00Z")
    tc_properties.append(cell_del)
    docx.save(path)

    with pytest.raises(AcceptRevisionsError):
        accept_all_revisions(path, tmp_path / "out.docx")


# --- ergonomics -----------------------------------------------------------------------


def test_accept_all_revisions_in_place(tmp_path: Path) -> None:
    source = _saved_docx(tmp_path, "input.docx", "The quick brown fox jumps.")
    document = load_docx(source)
    paragraph = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE_TEXT,
        node_id=paragraph.id,
        original_text="fox",
        replacement_text="cat",
        locator=ReviewLocator(node_id=paragraph.id, char_start=16, char_end=19),
        status=ActionStatus.APPLIED,
    )
    reviewed = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")

    # out_path == reviewed_path: rewrites the same file.
    result = accept_all_revisions(reviewed, reviewed)

    assert result == reviewed
    assert inspect_markup(reviewed).is_clean
    assert _body_paragraph_texts(reviewed) == ["The quick brown cat jumps."]


def test_apply_reviewed_markup_is_the_same_operation(tmp_path: Path) -> None:
    assert apply_reviewed_markup is accept_all_revisions


def test_accept_all_revisions_clean_document_untouched(tmp_path: Path) -> None:
    # A document with no markup flattens to an equivalent clean document.
    source = _saved_docx(tmp_path, "input.docx", "Nothing to accept here.", "Second line.")
    corrected = accept_all_revisions(source, tmp_path / "corrected.docx")

    assert inspect_markup(corrected).is_clean
    assert _body_paragraph_texts(corrected) == ["Nothing to accept here.", "Second line."]


def _add_tracked_run(paragraph, tag: str, text: str, *, deleted: bool = False) -> None:
    wrapper = OxmlElement(tag)
    wrapper.set(qn("w:id"), "9")
    wrapper.set(qn("w:author"), "reviewer")
    run = OxmlElement("w:r")
    node = OxmlElement("w:delText" if deleted else "w:t")
    node.text = text
    run.append(node)
    wrapper.append(run)
    paragraph._p.append(wrapper)


def test_reject_all_revisions_restores_deleted_text_and_drops_insertions(
    tmp_path: Path,
) -> None:
    path = tmp_path / "reviewed.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph("Original clause. ")
    _add_tracked_run(paragraph, "w:del", "old phrase", deleted=True)
    _add_tracked_run(paragraph, "w:ins", "lawyer phrase")
    docx.save(path)

    restored = reject_all_revisions(path, tmp_path / "input.docx")
    assert inspect_markup(restored).is_clean
    assert _body_paragraph_texts(restored) == ["Original clause. old phrase"]


def test_reject_all_revisions_merges_inserted_paragraph_mark(tmp_path: Path) -> None:
    path = tmp_path / "reviewed.docx"
    docx = DocxDocument()
    first = docx.add_paragraph("Hello ")
    _paragraph_mark("ins", first)
    docx.add_paragraph("world.")
    docx.save(path)

    restored = reject_all_revisions(path, tmp_path / "input.docx")
    assert inspect_markup(restored).is_clean
    assert _body_paragraph_texts(restored) == ["Hello world."]


def test_reject_all_revisions_drops_empty_numbered_leftover(tmp_path: Path) -> None:
    path = tmp_path / "reviewed.docx"
    docx = DocxDocument()
    docx.add_paragraph("Lead-in.")
    inserted = docx.add_paragraph()
    numbering = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    numbering.append(ilvl)
    numbering.append(num_id)
    inserted._p.get_or_add_pPr().append(numbering)
    _add_tracked_run(inserted, "w:ins", "new numbered item")
    docx.add_paragraph("Tail.")
    docx.save(path)

    restored = reject_all_revisions(path, tmp_path / "input.docx")
    assert inspect_markup(restored).is_clean
    assert _body_paragraph_texts(restored) == ["Lead-in.", "Tail."]


def test_reject_all_revisions_keeps_unrelated_empty_numbered_item(tmp_path: Path) -> None:
    path = tmp_path / "reviewed.docx"
    docx = DocxDocument()
    empty = docx.add_paragraph()
    numbering = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    numbering.append(ilvl)
    numbering.append(num_id)
    empty._p.get_or_add_pPr().append(numbering)
    docx.add_paragraph("Kept.")
    docx.save(path)

    restored = reject_all_revisions(path, tmp_path / "input.docx")
    assert inspect_markup(restored).is_clean
    assert _body_paragraph_texts(restored) == ["", "Kept."]


def test_reject_all_revisions_rejects_cell_deletion(tmp_path: Path) -> None:
    path = tmp_path / "table.docx"
    docx = DocxDocument()
    table = docx.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = "cell text"
    tc_properties = cell._tc.get_or_add_tcPr()
    cell_del = OxmlElement("w:cellDel")
    cell_del.set(qn("w:id"), "1")
    cell_del.set(qn("w:author"), "reviewer")
    cell_del.set(qn("w:date"), "1970-01-01T00:00:00Z")
    tc_properties.append(cell_del)
    docx.save(path)

    with pytest.raises(RejectRevisionsError):
        reject_all_revisions(path, tmp_path / "out.docx")
    assert not (tmp_path / "out.docx").exists()


def test_reject_all_revisions_in_place(tmp_path: Path) -> None:
    path = tmp_path / "reviewed.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph("Original clause. ")
    _add_tracked_run(paragraph, "w:ins", "lawyer phrase")
    docx.save(path)

    result = reject_all_revisions(path, path)
    assert result == path
    assert inspect_markup(path).is_clean
    assert _body_paragraph_texts(path) == ["Original clause. "]
