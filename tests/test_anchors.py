import pytest
from docx import Document
from docx.document import Document as DocxDocument

from reviewkit.anchors import (
    ANCHOR_LAST,
    SignatureBlockStart,
    find_body_paragraph,
    find_paragraph_by_locator,
    find_signature_block_start,
    is_supported_anchor,
    parse_body_anchor_index,
)

_SIGNATURE_KEYWORDS = (
    "signatur*",
    "date",
    "podpis*",
    "załączn*",
    "data",
    "miejsce",
    "attach*",
    "place",
)


def _make_document(texts: list[str]) -> DocxDocument:
    document = Document()
    for text in texts:
        document.add_paragraph(text)
    return document


@pytest.mark.parametrize(
    ("anchor", "expected"),
    [
        ("body:p:0", 0),
        ("body:p:7", 7),
        ("body:p:last", None),
        ("body:p:-1", None),
        ("body:p:1x", None),
        ("body:p:²", None),
        ("body:p:①", None),
        ("body:p:", None),
        ("body:p:0:extra", None),
        ("table:0:row:0:cell:0:p:0", None),
        ("", None),
    ],
)
def test_parse_body_anchor_index(anchor: str, expected: int | None) -> None:
    assert parse_body_anchor_index(anchor) == expected


def test_is_supported_anchor() -> None:
    assert is_supported_anchor(ANCHOR_LAST)
    assert is_supported_anchor("body:p:3")
    assert not is_supported_anchor("body:p:-3")
    assert not is_supported_anchor("body:p:²")
    assert not is_supported_anchor("header:0:p:0")
    assert not is_supported_anchor("")


def test_find_body_paragraph_resolves_in_range_and_fails_closed() -> None:
    document = _make_document(["first", "second"])
    paragraph = find_body_paragraph(document, 1)
    assert paragraph is not None
    assert paragraph.text == "second"
    assert find_body_paragraph(document, 2) is None
    assert find_body_paragraph(document, -1) is None


def test_find_paragraph_by_locator_body() -> None:
    document = _make_document(["first", "second"])
    paragraph = find_paragraph_by_locator(document, "body:p:1")
    assert paragraph is not None
    assert paragraph.text == "second"


def test_find_paragraph_by_locator_table_cell() -> None:
    document = _make_document(["intro"])
    table = document.add_table(rows=2, cols=2)
    table.rows[1].cells[0].paragraphs[0].add_run("cell text")
    paragraph = find_paragraph_by_locator(document, "table:0:row:1:cell:0:p:0")
    assert paragraph is not None
    assert paragraph.text == "cell text"


def test_find_paragraph_by_locator_header_and_footer() -> None:
    document = _make_document(["body text"])
    section = document.sections[0]
    section.header.paragraphs[0].add_run("header text")
    section.footer.paragraphs[0].add_run("footer text")
    header_paragraph = find_paragraph_by_locator(document, "header:0:p:0")
    footer_paragraph = find_paragraph_by_locator(document, "footer:0:p:0")
    assert header_paragraph is not None
    assert header_paragraph.text == "header text"
    assert footer_paragraph is not None
    assert footer_paragraph.text == "footer text"


@pytest.mark.parametrize(
    "locator",
    [
        None,
        "",
        "body:p:last",
        "body:p:9",
        "body:p:-1",
        "body:p:²",
        "body:x:0",
        "table:0:row:0:cell:0:q:0",
        "table:0:row:0:p:0",
        "table:x:row:0:cell:0:p:0",
        "table:²:row:0:cell:0:p:0",
        "header:x:p:0",
        "header:²:p:0",
        "header:5:p:0",
        "footnote:0:p:0",
    ],
)
def test_find_paragraph_by_locator_fails_closed(locator: str | None) -> None:
    document = _make_document(["only paragraph"])
    document.add_table(rows=1, cols=1)
    assert find_paragraph_by_locator(document, locator) is None


def test_signature_scan_finds_topmost_matching_tail_paragraph() -> None:
    document = _make_document(["Body clause.", "Signature: ____", "Date: ____"])
    start = find_signature_block_start(document, signature_keywords=_SIGNATURE_KEYWORDS)
    assert start is not None
    assert "Signature" in start.text


def test_signature_scan_without_keywords_finds_nothing() -> None:
    document = _make_document(["Body clause.", "Signature: ____"])
    assert find_signature_block_start(document) is None


def test_signature_scan_skips_whitespace_paragraphs() -> None:
    document = _make_document(["Body clause.", "Signature: ____", "   ", ""])
    start = find_signature_block_start(document, signature_keywords=_SIGNATURE_KEYWORDS)
    assert start is not None
    assert "Signature" in start.text


def test_signature_scan_treats_ignored_texts_like_whitespace() -> None:
    document = _make_document(
        ["Body clause.", "Signature: ____", "Inserted clause text.", "Date: ____"]
    )
    without_ignore = find_signature_block_start(document, signature_keywords=_SIGNATURE_KEYWORDS)
    assert without_ignore is not None
    assert "Date" in without_ignore.text

    with_ignore = find_signature_block_start(
        document,
        ignore_texts=["Inserted clause text."],
        signature_keywords=_SIGNATURE_KEYWORDS,
    )
    assert with_ignore is not None
    assert "Signature" in with_ignore.text


def test_signature_scan_respects_tail_window() -> None:
    document = _make_document(["Signature: ____"] + [f"Trailing clause {i}." for i in range(5)])
    assert find_signature_block_start(document, signature_keywords=_SIGNATURE_KEYWORDS) is None


def test_signature_scan_whitespace_does_not_consume_tail_window() -> None:
    document = _make_document(
        ["Signature: ____", " ", "Trailing 1.", "Trailing 2.", "Trailing 3.", "Trailing 4."]
    )
    start = find_signature_block_start(document, signature_keywords=_SIGNATURE_KEYWORDS)
    assert start is not None
    assert "Signature" in start.text


def test_signature_scan_stops_at_first_non_match_above_block() -> None:
    document = _make_document(["Date: ____", "Body clause.", "Signature: ____", "Date: ____"])
    start = find_signature_block_start(document, signature_keywords=_SIGNATURE_KEYWORDS)
    assert start is not None
    assert start.text == document.paragraphs[2].text
    assert "Signature" in start.text


def test_exact_keyword_matches_whole_word_only() -> None:
    # "date" without a stem star must not match "Dated" or "update".
    dated = _make_document(["Body clause.", "Dated 15 March ____"])
    assert find_signature_block_start(dated, signature_keywords=("date",)) is None

    exact = _make_document(["Body clause.", "Date: ____"])
    start = find_signature_block_start(exact, signature_keywords=("date",))
    assert start is not None
    assert "Date" in start.text

    update = _make_document(["Body clause.", "We update the register."])
    assert find_signature_block_start(update, signature_keywords=("date",)) is None


def test_stem_keyword_matches_inflections_and_case() -> None:
    document = _make_document(["Body clause.", "Załączniki:"])
    start = find_signature_block_start(document, signature_keywords=("załączn*",))
    assert start is not None
    assert "Załączniki" in start.text


def test_keyword_metacharacters_are_escaped() -> None:
    document = _make_document(["Body clause.", "cxv here"])
    assert find_signature_block_start(document, signature_keywords=("c.v.",)) is None


def test_empty_and_star_only_keywords_are_skipped_fail_closed() -> None:
    document = _make_document(["Body clause.", "Plain trailing text."])
    assert find_signature_block_start(document, signature_keywords=("", "*")) is None


def test_empty_keyword_tuple_finds_nothing() -> None:
    document = _make_document(["Body clause.", "Signature: ____"])
    assert find_signature_block_start(document, signature_keywords=()) is None


def test_signature_block_start_handle_is_opaque() -> None:
    document = _make_document(["Body clause.", "Signature: ____"])
    start = find_signature_block_start(document, signature_keywords=_SIGNATURE_KEYWORDS)
    assert isinstance(start, SignatureBlockStart)
    assert "Signature" in start.text
    # No element plumbing leaks through the handle.
    assert not hasattr(start, "addprevious")
