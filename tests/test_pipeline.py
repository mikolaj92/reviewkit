from pathlib import Path
from typing import Any
from xml.etree import ElementTree
from zipfile import ZipFile

import pytest
from docx import Document as DocxDocument

from reviewkit import ReviewResult, parser_docx, review_document
from reviewkit.context import ReviewContext, ReviewContextProvider
from reviewkit.document import ParagraphNode, ReviewDocument, SectionNode, SentenceNode
from reviewkit.llm import MockLLMClient
from reviewkit.models import (
    ActionStatus,
    ReviewAction,
    ReviewActionType,
    ReviewFinding,
    ReviewScope,
)
from reviewkit.pipeline import _unresolved_finding_id_warnings
from reviewkit.profile import ReviewProfile
from reviewkit.reviewer import HierarchicalReviewer
from reviewkit.state import ReviewState

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def test_hierarchical_review_passes_lower_level_results(tmp_path: Path) -> None:
    input_path = _make_docx(tmp_path, "Ala ma kota.")
    reviewed_path = tmp_path / "reviewed.docx"
    corrected_path = tmp_path / "corrected.docx"
    llm = MockLLMClient(
        responses=[
            {
                "actions": [
                    {
                        "id": "a-sentence",
                        "scope": "sentence",
                        "action_type": "replace",
                        "node_id": "p1.s1",
                        "original_text": "kota",
                        "replacement_text": "psa",
                        "reason": "Zmiana testowa.",
                        "category": "typo",
                        "confidence": 0.9,
                        "apply_hint": True,
                    }
                ],
                "summary": "Zdanie sprawdzone.",
            },
            {
                "actions": [
                    {
                        "id": "a-paragraph",
                        "scope": "paragraph",
                        "action_type": "comment",
                        "node_id": "p1",
                        "comment": "Akapit jest zrozumiały.",
                        "confidence": 0.8,
                    }
                ],
                "summary": "Akapit sprawdzony.",
            },
            {
                "actions": [
                    {
                        "id": "a-section",
                        "scope": "section",
                        "action_type": "summary",
                        "node_id": "s1",
                        "comment": "Sekcja jest krótka.",
                        "confidence": 0.8,
                    }
                ],
                "summary": "Sekcja sprawdzona.",
            },
            {
                "actions": [],
                "summary": "Dokument sprawdzony.",
            },
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=reviewed_path,
        out_corrected=corrected_path,
    )

    assert len(result.actions) == 3
    assert result.actions[0].status == ActionStatus.APPLIED
    assert result.document_summary == "Dokument sprawdzony."
    assert "sentence_review_results" in llm.calls[1].content
    assert "a-sentence" in llm.calls[1].content
    assert "paragraph_review_results" in llm.calls[2].content
    assert "a-paragraph" in llm.calls[2].content
    assert "section_review_results" in llm.calls[3].content
    assert "a-section" in llm.calls[3].content


def test_review_document_accepts_a_review_profile_object(tmp_path: Path) -> None:
    # A caller that already holds a ReviewProfile (built in memory or cached) should not have
    # to round-trip it through a folder on disk. Passing the object directly must work exactly
    # like passing its folder path.
    from reviewkit.profile import load_profile

    input_path = _make_docx(tmp_path, "The cat sat.")
    profile = load_profile("examples/profiles/story.teacher")
    llm = MockLLMClient(
        responses=[
            {"actions": [], "summary": "sentence"},
            {"actions": [], "summary": "paragraph"},
            {"actions": [], "summary": "section"},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path=profile,
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    assert isinstance(result, ReviewResult)
    assert result.document_summary == "Dokument sprawdzony."


def test_overlapping_edits_from_different_scopes_both_escalate(tmp_path: Path) -> None:
    # A sentence-scope edit and a paragraph-scope edit each auto-apply in isolation - they
    # run in separate LLM responses, so prepare_actions never compares them. But both land
    # on paragraph p1 ("The cat" at [0,7] vs "cat sat" at [4,11] overlap), and applying both
    # would clobber one silently. The post-hierarchy cross-scope pass must escalate both.
    input_path = _make_docx(tmp_path, "The cat sat.")
    llm = MockLLMClient(
        responses=[
            {
                "actions": [
                    {
                        "id": "a-sentence",
                        "scope": "sentence",
                        "action_type": "replace",
                        "node_id": "p1.s1",
                        "original_text": "The cat",
                        "replacement_text": "A feline",
                        "category": "typo",
                        "confidence": 1.0,
                        "apply_hint": True,
                    }
                ],
                "summary": "Sentence checked.",
            },
            {
                "actions": [
                    {
                        "id": "a-paragraph",
                        "scope": "paragraph",
                        "action_type": "replace",
                        "node_id": "p1",
                        "original_text": "cat sat",
                        "replacement_text": "dog ran",
                        "category": "typo",
                        "confidence": 1.0,
                        "apply_hint": True,
                    }
                ],
                "summary": "Paragraph checked.",
            },
            {"actions": [], "summary": "Section checked."},
            {"actions": [], "summary": "Document checked."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    statuses = {action.id: action.status for action in result.actions}
    assert statuses["a-sentence"] == ActionStatus.CONFLICT
    assert statuses["a-paragraph"] == ActionStatus.CONFLICT
    # Neither clobbering edit reached the clean copy.
    assert _docx_text(result.corrected_docx) == "The cat sat."


def test_subset_pipeline_rolls_lower_actions_up_to_the_next_enabled_scope() -> None:
    # A profile may enable only sentence + document (a subset the API permits). The
    # sentence-level action must still reach the document review even though the
    # intermediate paragraph and section scopes are skipped.
    document = ReviewDocument(
        sections=[
            SectionNode(
                id="s1",
                paragraphs=[
                    ParagraphNode(
                        id="p1",
                        text="The cat sat.",
                        section_id="s1",
                        sentences=[
                            SentenceNode(id="p1.s1", text="The cat sat.", paragraph_id="p1")
                        ],
                    )
                ],
            )
        ]
    )
    profile = ReviewProfile(
        name="generic",
        language="en",
        document_type="generic document",
        reviewer_role="generic reviewer",
        review_pipeline=[ReviewScope.SENTENCE, ReviewScope.DOCUMENT],
    )
    llm = MockLLMClient(
        responses=[
            {
                "actions": [
                    {
                        "id": "a-sentence",
                        "scope": "sentence",
                        "action_type": "comment",
                        "node_id": "p1.s1",
                        "comment": "Observed at sentence level.",
                        "confidence": 0.9,
                    }
                ],
                "summary": "Sentence checked.",
            },
            {"actions": [], "summary": "Document checked."},
        ]
    )

    reviewer = HierarchicalReviewer(profile=profile, llm=llm)
    reviewer.review(document)

    # Exactly two LLM calls: sentence then document (paragraph + section skipped).
    assert len(llm.calls) == 2
    document_prompt = llm.calls[1].content
    assert "section_review_results" in document_prompt
    assert "a-sentence" in document_prompt


def test_malformed_action_does_not_discard_valid_findings_summary_and_actions() -> None:
    # A real LLM emits one out-of-spec action among good items. Only the bad action must
    # be dropped; the valid action, the finding, and the summary must all survive.
    class _RawDictLLM:
        def complete_json(self, messages: list[dict[str, str]], schema: type[Any]) -> Any:
            return {
                "actions": [
                    {
                        "id": "good",
                        "scope": "paragraph",
                        "action_type": "comment",
                        "node_id": "p1",
                        "comment": "Valid action.",
                        "confidence": 0.9,
                    },
                    {
                        "id": "bad",
                        "scope": "paragraph",
                        "action_type": "not_a_real_action_type",
                        "node_id": "p1",
                        "confidence": 0.9,
                    },
                ],
                "findings": [
                    {
                        "finding_id": "f1",
                        "node_id": "p1",
                        "title": "Observation",
                        "description": "Salvaged despite the bad action.",
                        "dimension": "clarity",
                        "severity": "low",
                    }
                ],
                "summary": "Paragraph checked.",
            }

    document = ReviewDocument(
        sections=[
            SectionNode(
                id="s1",
                paragraphs=[ParagraphNode(id="p1", text="The cat sat.", section_id="s1")],
            )
        ]
    )
    profile = ReviewProfile(
        name="generic",
        language="en",
        document_type="generic document",
        reviewer_role="generic reviewer",
        review_pipeline=[ReviewScope.PARAGRAPH],
    )

    reviewer = HierarchicalReviewer(profile=profile, llm=_RawDictLLM())
    findings, actions, state = reviewer.review(document)

    assert [action.id for action in actions] == ["good"]
    assert [finding.finding_id for finding in findings] == ["f1"]
    assert state.paragraph_summaries.get("p1") == "Paragraph checked."
    assert any("action 1" in warning for warning in state.warnings)


def test_sentence_review_adds_action(tmp_path: Path) -> None:
    result = _run_with_single_sentence_action(
        tmp_path,
        action={
            "id": "a1",
            "scope": "sentence",
            "action_type": "replace",
            "node_id": "p1.s1",
            "original_text": "bład",
            "replacement_text": "błąd",
            "category": "typo",
            "confidence": 1.0,
        },
        text="To jest bład.",
    )

    assert len(result.actions) == 1
    assert result.actions[0].id == "a1"
    assert result.actions[0].action_type == ReviewActionType.REPLACE


def test_applied_actions_are_written_to_corrected_docx(tmp_path: Path) -> None:
    result = _run_with_single_sentence_action(
        tmp_path,
        action={
            "id": "a1",
            "scope": "sentence",
            "action_type": "replace",
            "node_id": "p1.s1",
            "original_text": "bład",
            "replacement_text": "błąd",
            "category": "typo",
            "confidence": 1.0,
            "apply_hint": True,
        },
        text="To jest bład.",
    )

    corrected_text = _docx_text(result.corrected_docx)
    assert "błąd" in corrected_text
    assert "bład" not in corrected_text


def test_suggestion_text_edits_are_tracked_in_reviewed_but_not_applied_in_corrected(
    tmp_path: Path,
) -> None:
    result = _run_with_single_sentence_action(
        tmp_path,
        action={
            "id": "a1",
            "scope": "sentence",
            "action_type": "replace",
            "node_id": "p1.s1",
            "original_text": "bardzo",
            "replacement_text": "wyjątkowo",
            "category": "style",
            "confidence": 1.0,
        },
        text="To jest bardzo dobre.",
    )

    corrected_text = _docx_text(result.corrected_docx)
    reviewed_comments = _docx_comments(result.reviewed_docx)
    reviewed_xml = _docx_document_xml(result.reviewed_docx)
    assert "[DELETE:" not in reviewed_xml
    assert _revision_texts(result.reviewed_docx, "del", "delText") == ["bardzo"]
    assert _revision_texts(result.reviewed_docx, "ins", "t") == ["wyjątkowo"]
    assert "SUGGESTION" in reviewed_comments
    assert result.actions[0].status == ActionStatus.NOT_APPLIED
    assert "bardzo" in corrected_text
    assert "wyjątkowo" not in corrected_text


def test_conflict_is_not_applied(tmp_path: Path) -> None:
    result = _run_with_single_sentence_action(
        tmp_path,
        action={
            "id": "a1",
            "scope": "sentence",
            "action_type": "replace",
            "node_id": "p1.s1",
            "original_text": "kot",
            "replacement_text": "pies",
            "category": "typo",
            "confidence": 1.0,
        },
        text="kot i kot.",
    )

    corrected_text = _docx_text(result.corrected_docx)
    assert result.actions[0].status == ActionStatus.CONFLICT
    assert "kot i kot." in corrected_text
    assert "pies" not in corrected_text


def test_document_type_action_policy_can_change_review_status(tmp_path: Path) -> None:
    result = _run_with_single_sentence_action(
        tmp_path,
        action={
            "id": "a1",
            "scope": "sentence",
            "action_type": "replace",
            "node_id": "p1.s1",
            "original_text": "bład",
            "replacement_text": "błąd",
            "category": "typo",
            "severity": "high",
            "confidence": 1.0,
            "apply_hint": True,
        },
        text="To jest bład.",
    )

    assert result.actions[0].status == ActionStatus.NEEDS_HUMAN_DECISION
    assert "exceeds policy threshold" in (result.actions[0].policy_reason or "")
    corrected_text = _docx_text(result.corrected_docx)
    assert "bład" in corrected_text
    assert "błąd" not in corrected_text


def test_policy_guard_blocks_corrected_when_protected_placeholder_changes(
    tmp_path: Path,
) -> None:
    input_path = _make_docx(tmp_path, "[OSOBA_1] podpisał umowę.")
    llm = MockLLMClient(
        responses=[
            {
                "actions": [
                    {
                        "id": "a1",
                        "scope": "sentence",
                        "action_type": "replace",
                        "node_id": "p1.s1",
                        "original_text": "[OSOBA_1]",
                        "replacement_text": "Jan Kowalski",
                        "category": "typo",
                        "confidence": 1.0,
                        "apply_to_corrected": True,
                    }
                ],
                "summary": "Zdanie sprawdzone.",
            },
            {"actions": [], "summary": "Akapit sprawdzony."},
            {"actions": [], "summary": "Sekcja sprawdzona."},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/employment-contract.lawyer",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    assert result.actions[0].status == ActionStatus.NEEDS_HUMAN_DECISION
    assert result.actions[0].metadata["blocked_from_corrected"] is True
    corrected_text = _docx_text(result.corrected_docx)
    assert "[OSOBA_1] podpisał umowę." in corrected_text
    assert "Jan Kowalski" not in corrected_text


def test_context_provider_is_included_in_prompts(tmp_path: Path) -> None:
    input_path = _make_docx(tmp_path, "Ala ma kota.")
    llm = MockLLMClient(
        responses=[
            {"actions": [], "summary": "Zdanie sprawdzone."},
            {"actions": [], "summary": "Akapit sprawdzony."},
            {"actions": [], "summary": "Sekcja sprawdzona."},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )

    review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
        context_provider=_StaticContextProvider(),
    )

    assert "external_review_context" in llm.calls[0].content
    assert "Dike-style grounding" in llm.calls[0].content


def test_tracked_revision_inputs_are_reported_as_warning(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(parser_docx, "_contains_tracked_revisions", lambda path: True)

    result = _run_with_single_sentence_action(
        tmp_path,
        action={
            "id": "a1",
            "scope": "sentence",
            "action_type": "comment",
            "node_id": "p1.s1",
            "comment": "Sprawdź historię zmian.",
            "confidence": 1.0,
        },
        text="To jest tekst.",
    )

    assert result.warnings == ["Input DOCX contains tracked revisions."]


def test_action_referencing_unknown_finding_id_is_reported_as_warning(tmp_path: Path) -> None:
    # The finding<->action linkage is the archetype's audit trail. An action pointing at a
    # finding_id no finding carries is a broken link and must surface as a warning; an action
    # whose finding_id resolves must not.
    input_path = _make_docx(tmp_path, "The cat sat.")
    llm = MockLLMClient(
        responses=[
            {
                "findings": [
                    {
                        "finding_id": "finding-real",
                        "node_id": "p1.s1",
                        "title": "Observation",
                        "description": "Worth reviewing.",
                    }
                ],
                "actions": [
                    {
                        "id": "a-linked",
                        "scope": "sentence",
                        "action_type": "comment",
                        "node_id": "p1.s1",
                        "finding_id": "finding-real",
                        "comment": "Responds to the finding.",
                        "confidence": 0.9,
                    },
                    {
                        "id": "a-dangling",
                        "scope": "sentence",
                        "action_type": "comment",
                        "node_id": "p1.s1",
                        "finding_id": "finding-ghost",
                        "comment": "Points at a finding that does not exist.",
                        "confidence": 0.9,
                    },
                ],
                "summary": "Sentence checked.",
            },
            {"actions": [], "summary": "Paragraph checked."},
            {"actions": [], "summary": "Section checked."},
            {"actions": [], "summary": "Document checked."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    assert result.warnings == [
        "Action a-dangling references unknown finding_id 'finding-ghost'."
    ]


def test_action_referencing_a_merged_away_finding_id_is_not_flagged() -> None:
    # When a duplicate finding is merged away, its finding_id is preserved on the survivor as
    # an alias, so an action that referenced the merged-away copy still resolves and must not
    # be reported as a dangling reference.
    survivor = ReviewFinding(
        finding_id="finding-a",
        node_id="p1",
        title="T",
        description="D",
        metadata={"merged_finding_ids": ["finding-b"]},
    )
    action = ReviewAction(
        id="a-1",
        scope=ReviewScope.SENTENCE,
        action_type=ReviewActionType.COMMENT,
        node_id="p1",
        finding_id="finding-b",
        comment="Responds to the merged-away finding.",
    )

    assert _unresolved_finding_id_warnings([survivor], [action]) == []


def test_core_system_prompt_requests_finding_id_linkage(tmp_path: Path) -> None:
    # The prompt must actually ask the model to emit finding_ids and link actions to them,
    # otherwise the linkage the archetype relies on is never populated in real runs.
    input_path = _make_docx(tmp_path, "The cat sat.")
    llm = MockLLMClient(
        responses=[
            {"actions": [], "summary": "Sentence checked."},
            {"actions": [], "summary": "Paragraph checked."},
            {"actions": [], "summary": "Section checked."},
            {"actions": [], "summary": "Document checked."},
        ]
    )

    review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    system_prompt = llm.calls[0].messages[0]["content"]
    assert llm.calls[0].messages[0]["role"] == "system"
    assert "set that action's finding_id to the same value" in system_prompt


def test_sentence_offset_edit_targets_the_correct_sentence(tmp_path: Path) -> None:
    input_path = _make_docx(tmp_path, "First sentence. Second sentence.")
    llm = MockLLMClient(
        responses=[
            {"actions": [], "summary": "s1."},
            {
                "actions": [
                    {
                        "id": "a1",
                        "scope": "sentence",
                        "action_type": "replace",
                        "node_id": "p1.s2",
                        "original_text": "Second",
                        "replacement_text": "Next",
                        "category": "typo",
                        "confidence": 1.0,
                        "apply_hint": True,
                        "locator": {"char_start": 0, "char_end": 6},
                    }
                ],
                "summary": "s2.",
            },
            {"actions": [], "summary": "Akapit sprawdzony."},
            {"actions": [], "summary": "Sekcja sprawdzona."},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    corrected_text = _docx_text(result.corrected_docx)
    assert result.actions[0].status == ActionStatus.APPLIED
    assert corrected_text == "First sentence. Next sentence."
    assert "Nextsentence" not in corrected_text


def test_sentence_string_edit_targets_the_matching_sentence(tmp_path: Path) -> None:
    input_path = _make_docx(tmp_path, "The cat sat. The cat ran.")
    llm = MockLLMClient(
        responses=[
            {"actions": [], "summary": "s1."},
            {
                "actions": [
                    {
                        "id": "a1",
                        "scope": "sentence",
                        "action_type": "replace",
                        "node_id": "p1.s2",
                        "original_text": "cat",
                        "replacement_text": "dog",
                        "category": "typo",
                        "confidence": 1.0,
                        "apply_hint": True,
                    }
                ],
                "summary": "s2.",
            },
            {"actions": [], "summary": "Akapit sprawdzony."},
            {"actions": [], "summary": "Sekcja sprawdzona."},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    corrected_text = _docx_text(result.corrected_docx)
    assert result.actions[0].status == ActionStatus.APPLIED
    assert corrected_text == "The cat sat. The dog ran."


def test_one_failing_node_does_not_abort_the_review(tmp_path: Path) -> None:
    class _PartiallyFailingLLM:
        def __init__(self) -> None:
            self.calls = 0

        def complete_json(self, messages: list[dict[str, str]], schema: type[Any]) -> Any:
            self.calls += 1
            if self.calls == 1:  # sentence node raises
                raise RuntimeError("simulated LLM failure")
            if self.calls == 2:  # paragraph node still produces a real, applicable edit
                return schema.model_validate(
                    {
                        "actions": [
                            {
                                "id": "p1a",
                                "scope": "paragraph",
                                "action_type": "replace",
                                "node_id": "p1",
                                "original_text": "bład",
                                "replacement_text": "błąd",
                                "category": "typo",
                                "confidence": 1.0,
                                "apply_hint": True,
                            }
                        ],
                        "summary": "ok",
                    }
                )
            return schema()

    input_path = _make_docx(tmp_path, "To jest bład.")
    llm = _PartiallyFailingLLM()

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    # Every node was still visited despite the sentence failure.
    assert llm.calls == 4
    # The failure is surfaced as a warning, naming the node.
    assert any("sentence p1.s1" in warning for warning in result.warnings)
    assert any("RuntimeError" in warning for warning in result.warnings)
    # The other nodes produced results: the paragraph edit was applied.
    assert result.stats.applied_count == 1
    assert _docx_text(result.corrected_docx) == "To jest błąd."


def test_propagate_llm_errors_reraises_on_first_client_failure() -> None:
    # Fail-closed opt-in (the peer of the resilient default above): a raising client
    # aborts the whole review on the FIRST failing node instead of degrading to a
    # partial result, and re-raises the original exception unchanged.
    class _RaisingLLM:
        def __init__(self) -> None:
            self.calls = 0

        def complete_json(self, messages: list[dict[str, str]], schema: type[Any]) -> Any:
            self.calls += 1
            raise RuntimeError("simulated LLM failure")

    document = ReviewDocument(
        sections=[
            SectionNode(
                id="s1",
                paragraphs=[
                    ParagraphNode(id="p1", text="The cat sat.", section_id="s1"),
                    ParagraphNode(id="p2", text="The dog ran.", section_id="s1"),
                ],
            )
        ]
    )
    profile = ReviewProfile(
        name="generic",
        language="en",
        document_type="generic document",
        reviewer_role="generic reviewer",
        review_pipeline=[ReviewScope.PARAGRAPH],
    )

    llm = _RaisingLLM()
    reviewer = HierarchicalReviewer(profile=profile, llm=llm, propagate_llm_errors=True)

    with pytest.raises(RuntimeError, match="simulated LLM failure"):
        reviewer.review(document)

    # Fail-fast: it stopped at the first failing node, never visiting the second.
    assert llm.calls == 1


def test_propagate_llm_errors_default_false_still_degrades() -> None:
    # The default (flag absent) keeps the resilient behavior: the same raising client
    # surfaces a warning and completes rather than aborting.
    class _RaisingLLM:
        def complete_json(self, messages: list[dict[str, str]], schema: type[Any]) -> Any:
            raise RuntimeError("simulated LLM failure")

    document = ReviewDocument(
        sections=[
            SectionNode(
                id="s1",
                paragraphs=[ParagraphNode(id="p1", text="The cat sat.", section_id="s1")],
            )
        ]
    )
    profile = ReviewProfile(
        name="generic",
        language="en",
        document_type="generic document",
        reviewer_role="generic reviewer",
        review_pipeline=[ReviewScope.PARAGRAPH],
    )

    findings, actions, state = HierarchicalReviewer(
        profile=profile, llm=_RaisingLLM()
    ).review(document)

    assert actions == []
    assert any("RuntimeError" in warning for warning in state.warnings)


def test_identical_finding_surfaced_at_two_levels_appears_once(tmp_path: Path) -> None:
    input_path = _make_docx(tmp_path, "Ala ma kota.")
    sentence_finding = {
        "finding_id": "dup-1",
        "node_id": "p1.s1",
        "title": "Repeated observation",
        "description": "The same issue seen at two levels.",
        "dimension": "clarity",
        "severity": "low",
    }
    # The paragraph level re-surfaces the same finding (same finding_id) at its own node.
    paragraph_finding = {**sentence_finding, "node_id": "p1"}
    llm = MockLLMClient(
        responses=[
            {"findings": [sentence_finding], "summary": "Zdanie sprawdzone."},
            {"findings": [paragraph_finding], "summary": "Akapit sprawdzony."},
            {"summary": "Sekcja sprawdzona."},
            {"summary": "Dokument sprawdzony."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )

    assert [finding.finding_id for finding in result.findings] == ["dup-1"]


def test_review_document_threads_an_injected_action_policy(tmp_path: Path) -> None:
    # An edit that auto-applies with the profile's config-only policy must escalate once a
    # caller injects an ActionPolicy carrying a fail-closed guard -- proving review_document
    # threads the policy hook through to prepare_actions.
    from reviewkit.policy import ActionPolicy
    from reviewkit.profile import load_profile

    input_path = _make_docx(tmp_path, "To jest bład.")
    profile = load_profile("examples/profiles/story.teacher")

    def _block_all_writes(action, node_text):  # type: ignore[no-untyped-def]
        return "guard: no automatic writes in this run"

    policy = ActionPolicy.from_profile(profile, guards=[_block_all_writes])
    llm = MockLLMClient(
        responses=[
            {
                "actions": [
                    {
                        "id": "a1",
                        "scope": "sentence",
                        "action_type": "replace",
                        "node_id": "p1.s1",
                        "original_text": "bład",
                        "replacement_text": "błąd",
                        "category": "typo",
                        "confidence": 1.0,
                        "apply_hint": True,
                    }
                ],
                "summary": "Zdanie sprawdzone.",
            },
            {"actions": [], "summary": "Akapit sprawdzony."},
            {"actions": [], "summary": "Sekcja sprawdzona."},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
        action_policy=policy,
    )

    assert result.actions[0].status == ActionStatus.NEEDS_HUMAN_DECISION
    assert "guard" in (result.actions[0].policy_reason or "")
    # Escalated edits never reach the clean copy.
    assert "bład" in _docx_text(result.corrected_docx)


def test_identical_runs_produce_byte_identical_json_reports(tmp_path: Path) -> None:
    # A uuid4 default made every report unique even for identical input, defeating
    # diffing/caching. With content-derived ids two identical runs must serialize to
    # byte-identical JSON. The action and finding omit ids so the derivation is exercised.
    input_path = _make_docx(tmp_path, "To jest zdanie.")
    reviewed_path = tmp_path / "reviewed.docx"
    corrected_path = tmp_path / "corrected.docx"

    def _report(dest: Path) -> bytes:
        llm = MockLLMClient(
            responses=[
                {
                    "actions": [
                        {
                            "scope": "sentence",
                            "action_type": "comment",
                            "node_id": "p1.s1",
                            "comment": "Dobre zdanie.",
                            "confidence": 0.9,
                        }
                    ],
                    "findings": [
                        {
                            "node_id": "p1.s1",
                            "title": "Observation",
                            "description": "A stable finding.",
                            "dimension": "clarity",
                            "severity": "low",
                        }
                    ],
                    "summary": "Zdanie sprawdzone.",
                },
                {"actions": [], "summary": "Akapit sprawdzony."},
                {"actions": [], "summary": "Sekcja sprawdzona."},
                {"actions": [], "summary": "Dokument sprawdzony."},
            ]
        )
        result = review_document(
            input_path=input_path,
            profile_path="examples/profiles/story.teacher",
            llm=llm,
            out_reviewed=reviewed_path,
            out_corrected=corrected_path,
        )
        return result.save_json(dest).read_bytes()

    assert _report(tmp_path / "a.json") == _report(tmp_path / "b.json")


def test_extra_action_is_tracked_in_reviewed_and_applied_in_corrected(tmp_path: Path) -> None:
    # A valid extra action (deterministic caller, no LLM involved) must flow through the same
    # machinery as reviewer output: validated, policy-checked, rendered as a tracked change in
    # reviewed.docx and applied in corrected.docx, with its source_system preserved.
    input_path = _make_docx(tmp_path, "To jest bład.")
    extra = ReviewAction(
        id="extra-1",
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id="p1",
        original_text="bład",
        replacement_text="błąd",
        category="typo",
        confidence=1.0,
        apply_hint=True,
        source_system="deterministic-checker",
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=_empty_llm(),
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
        extra_actions=[extra],
    )

    assert [action.id for action in result.actions] == ["extra-1"]
    assert result.actions[0].status == ActionStatus.APPLIED
    assert result.actions[0].source_system == "deterministic-checker"
    assert _revision_texts(result.reviewed_docx, "del", "delText") == ["bład"]
    assert _revision_texts(result.reviewed_docx, "ins", "t") == ["błąd"]
    assert _docx_text(result.corrected_docx) == "To jest błąd."


def test_extra_action_overlapping_an_llm_action_escalates_both(tmp_path: Path) -> None:
    # An extra action overlapping an LLM edit must demote per the existing overlap semantics:
    # both actions in the cluster become CONFLICT and neither reaches the clean copy.
    input_path = _make_docx(tmp_path, "The cat sat.")
    llm = MockLLMClient(
        responses=[
            {
                "actions": [
                    {
                        "id": "a-llm",
                        "scope": "sentence",
                        "action_type": "replace",
                        "node_id": "p1.s1",
                        "original_text": "The cat",
                        "replacement_text": "A feline",
                        "category": "typo",
                        "confidence": 1.0,
                        "apply_hint": True,
                    }
                ],
                "summary": "Sentence checked.",
            },
            {"actions": [], "summary": "Paragraph checked."},
            {"actions": [], "summary": "Section checked."},
            {"actions": [], "summary": "Document checked."},
        ]
    )
    extra = ReviewAction(
        id="a-extra",
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id="p1",
        original_text="cat sat",
        replacement_text="dog ran",
        category="typo",
        confidence=1.0,
        apply_hint=True,
        source_system="deterministic-checker",
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
        extra_actions=[extra],
    )

    statuses = {action.id: action.status for action in result.actions}
    assert statuses["a-llm"] == ActionStatus.CONFLICT
    assert statuses["a-extra"] == ActionStatus.CONFLICT
    extra_result = next(action for action in result.actions if action.id == "a-extra")
    assert extra_result.source_system == "deterministic-checker"
    # Neither clobbering edit reached the clean copy.
    assert _docx_text(result.corrected_docx) == "The cat sat."


def test_extra_action_with_unmatched_original_text_becomes_conflict(tmp_path: Path) -> None:
    # An extra action whose original_text does not exist in the document must surface as a
    # CONFLICT comment in reviewed.docx - never a silent apply - and leave corrected untouched.
    input_path = _make_docx(tmp_path, "To jest bład.")
    extra = ReviewAction(
        id="extra-ghost",
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id="p1",
        original_text="unicorn",
        replacement_text="horse",
        category="typo",
        confidence=1.0,
        apply_hint=True,
        source_system="deterministic-checker",
    )

    result = review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=_empty_llm(),
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
        extra_actions=[extra],
    )

    assert result.actions[0].status == ActionStatus.CONFLICT
    assert "found 0 matches" in (result.actions[0].policy_reason or "")
    comment_lines = _docx_comments(result.reviewed_docx).splitlines()
    assert any(line.startswith("CONFLICT:") for line in comment_lines), comment_lines
    assert _docx_text(result.corrected_docx) == "To jest bład."


def test_extra_actions_none_or_empty_matches_omitting_the_parameter(tmp_path: Path) -> None:
    # extra_actions=None (and []) must be byte-identical to today's behavior: same report
    # JSON and same reviewed/corrected artifacts as a call without the parameter.
    input_path = _make_docx(tmp_path, "To jest bład.")
    reviewed_path = tmp_path / "reviewed.docx"
    corrected_path = tmp_path / "corrected.docx"

    def _artifacts(suffix: str, **kwargs: Any) -> tuple[bytes, ...]:
        llm = MockLLMClient(
            responses=[
                {
                    "actions": [
                        {
                            "id": "a1",
                            "scope": "sentence",
                            "action_type": "replace",
                            "node_id": "p1.s1",
                            "original_text": "bład",
                            "replacement_text": "błąd",
                            "category": "typo",
                            "confidence": 1.0,
                            "apply_hint": True,
                        }
                    ],
                    "summary": "Zdanie sprawdzone.",
                },
                {"actions": [], "summary": "Akapit sprawdzony."},
                {"actions": [], "summary": "Sekcja sprawdzona."},
                {"actions": [], "summary": "Dokument sprawdzony."},
            ]
        )
        result = review_document(
            input_path=input_path,
            profile_path="examples/profiles/story.teacher",
            llm=llm,
            out_reviewed=reviewed_path,
            out_corrected=corrected_path,
            **kwargs,
        )
        report = result.save_json(tmp_path / f"report-{suffix}.json").read_bytes()
        with ZipFile(reviewed_path) as archive:
            reviewed = (archive.read("word/document.xml"), archive.read("word/comments.xml"))
        with ZipFile(corrected_path) as archive:
            corrected = archive.read("word/document.xml")
        return (report, *reviewed, corrected)

    baseline = _artifacts("baseline")
    assert _artifacts("none", extra_actions=None) == baseline
    assert _artifacts("empty", extra_actions=[]) == baseline


def _empty_llm() -> MockLLMClient:
    return MockLLMClient(
        responses=[
            {"actions": [], "summary": "Zdanie sprawdzone."},
            {"actions": [], "summary": "Akapit sprawdzony."},
            {"actions": [], "summary": "Sekcja sprawdzona."},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )


def _run_with_single_sentence_action(
    tmp_path: Path,
    action: dict[str, Any],
    text: str,
) -> ReviewResult:
    input_path = _make_docx(tmp_path, text)
    llm = MockLLMClient(
        responses=[
            {"actions": [action], "summary": "Zdanie sprawdzone."},
            {"actions": [], "summary": "Akapit sprawdzony."},
            {"actions": [], "summary": "Sekcja sprawdzona."},
            {"actions": [], "summary": "Dokument sprawdzony."},
        ]
    )
    return review_document(
        input_path=input_path,
        profile_path="examples/profiles/story.teacher",
        llm=llm,
        out_reviewed=tmp_path / "reviewed.docx",
        out_corrected=tmp_path / "corrected.docx",
    )


def _make_docx(tmp_path: Path, text: str) -> Path:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph(text)
    docx.save(input_path)
    return input_path


def _docx_text(path: Path | None) -> str:
    assert path is not None
    docx = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in docx.paragraphs)


def _docx_comments(path: Path | None) -> str:
    assert path is not None
    docx = DocxDocument(str(path))
    return "\n".join(comment.text for comment in docx.comments)


def _docx_document_xml(path: Path | None) -> str:
    assert path is not None
    with ZipFile(path) as archive:
        return archive.read("word/document.xml").decode()


def _revision_texts(path: Path | None, revision_tag: str, text_tag: str) -> list[str]:
    root = ElementTree.fromstring(_docx_document_xml(path))
    return [
        "".join(element.itertext())
        for element in root.findall(f".//{_W}{revision_tag}")
        if element.find(f".//{_W}{text_tag}") is not None
    ]


class _StaticContextProvider(ReviewContextProvider):
    def context_for(
        self,
        *,
        profile: ReviewProfile,
        document: ReviewDocument,
        state: ReviewState,
        scope,
        node,
    ) -> ReviewContext:
        return ReviewContext(
            scope=scope,
            node_id=node.id,
            data={"grounding": "Dike-style grounding", "source": "test"},
        )
