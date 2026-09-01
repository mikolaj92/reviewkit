from pathlib import Path
import pytest
from docx import Document
from docxtor import RevisionAuthor, RevisionPosition, insert_revision, publish_docx
from reviewkit import has_comments, has_suggestion_marker, has_tracked_revisions, inspect_markup


def _docx(path: Path, text: str = "Plain text.") -> Path:
    d = Document()
    d.add_paragraph(text)
    d.save(path)
    return path


def test_clean_document_has_no_markup(tmp_path) -> None:
    p = _docx(tmp_path / "x.docx")
    assert inspect_markup(p).is_clean


def test_typed_revision_is_detected(tmp_path) -> None:
    p = _docx(tmp_path / "x.docx")
    result = insert_revision(
        p.read_bytes(), RevisionPosition("body:p:0", 5), "new ", RevisionAuthor("R")
    )
    publish_docx(result.data, p)
    assert has_tracked_revisions(p)


def test_comment_is_detected(tmp_path) -> None:
    from docxtor import CommentAuthor, add_paragraph_comment

    p = _docx(tmp_path / "x.docx")
    result = add_paragraph_comment(p.read_bytes(), "body:p:0", "Note", CommentAuthor("R"))
    publish_docx(result.data, p)
    assert has_comments(p)


def test_suggestion_marker_is_detected(tmp_path) -> None:
    p = _docx(tmp_path / "x.docx", "[SUGGESTION: reason] text")
    assert has_suggestion_marker(p)


def test_unreadable_package_fails_closed(tmp_path) -> None:
    p = tmp_path / "x.docx"
    p.write_bytes(b"not docx")
    with pytest.raises(Exception):
        inspect_markup(p)
