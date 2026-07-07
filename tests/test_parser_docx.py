from pathlib import Path
from zipfile import ZipFile

from docx import Document as DocxDocument

from reviewkit.parser_docx import DocxFootnote, load_docx, read_footnotes, split_sentences

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _write_footnotes_part(path: Path, body: str) -> Path:
    docx = DocxDocument()
    docx.add_paragraph("Body prose.")
    docx.save(path)
    part = (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        f"<w:footnotes xmlns:w='{_W_NS}'>{body}</w:footnotes>"
    )
    with ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", part)
    return path


def test_document_is_split_into_sections_paragraphs_and_sentences(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_heading("Rozdział 1", level=1)
    docx.add_paragraph("Ala ma kota. Kot ma dom.")
    docx.save(input_path)

    document = load_docx(input_path)

    assert len(document.sections) == 1
    assert document.sections[0].id == "s1"
    assert document.sections[0].title == "Rozdział 1"
    assert len(document.sections[0].paragraphs) == 1
    paragraph = document.sections[0].paragraphs[0]
    assert paragraph.id == "p1"
    assert paragraph.locator == "body:p:1"
    assert document.metadata["tracked_revisions_detected"] == "false"
    assert [sentence.text for sentence in paragraph.sentences] == ["Ala ma kota.", "Kot ma dom."]


def test_decimal_numbers_do_not_split_a_sentence() -> None:
    assert split_sentences("Pi is 3.14 today.") == ["Pi is 3.14 today."]


def test_abbreviations_do_not_over_split() -> None:
    assert split_sentences("Sp. z o.o.") == ["Sp. z o.o."]
    assert split_sentences("See J. R. R. Tolkien.") == ["See J. R. R. Tolkien."]


def test_genuine_sentence_boundaries_still_split() -> None:
    assert split_sentences("First one. Second one!") == ["First one.", "Second one!"]


def test_non_latin_terminators_split_sentences() -> None:
    # The sentence tier must not silently disappear for non-Latin scripts (contract #2975).
    # CJK writes no inter-sentence space, so these strong terminators split without one.
    assert split_sentences("这是第一句。这是第二句。") == ["这是第一句。", "这是第二句。"]
    assert split_sentences("पहला वाक्य। दूसरा वाक्य।") == ["पहला वाक्य।", "दूसरा वाक्य।"]
    assert split_sentences("هل هذا صحيح؟ نعم.") == ["هل هذا صحيح؟", "نعم."]


def test_english_styled_heading_starts_a_new_section(tmp_path: Path) -> None:
    input_path = tmp_path / "headings.docx"
    docx = DocxDocument()
    docx.add_heading("Introduction", level=1)
    docx.add_paragraph("First body paragraph.")
    docx.add_heading("Conclusion", level=2)
    docx.add_paragraph("Second body paragraph.")
    docx.save(input_path)

    document = load_docx(input_path)

    assert [section.title for section in document.sections] == ["Introduction", "Conclusion"]
    assert [len(section.paragraphs) for section in document.sections] == [1, 1]


def test_core_source_contains_no_polish_or_domain_vocabulary() -> None:
    src_root = Path(__file__).resolve().parent.parent / "src" / "reviewkit"
    forbidden = [
        "naglowek",
        "nagłówek",
        "niski",
        "średni",
        "sredni",
        "wysoki",
        "opowiadanie",
        "nauczyciel",
        "uczeń",
        "ucznia",
        "literówk",
    ]
    offenders: list[str] = []
    for path in sorted(src_root.rglob("*.py")):
        text = path.read_text(encoding="utf-8").lower()
        for token in forbidden:
            if token in text:
                offenders.append(f"{path.relative_to(src_root)}: {token!r}")
    assert not offenders, "domain/language vocabulary must not appear in core src: " + ", ".join(
        offenders
    )


def test_docx_parser_reads_table_paragraphs_with_locators(tmp_path: Path) -> None:
    input_path = tmp_path / "table.docx"
    docx = DocxDocument()
    table = docx.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Treść w tabeli."
    docx.save(input_path)

    document = load_docx(input_path)

    paragraph = document.sections[0].paragraphs[0]
    assert paragraph.text == "Treść w tabeli."
    assert paragraph.locator == "table:0:row:0:cell:0:p:0"


def test_table_lands_under_its_authoring_section(tmp_path: Path) -> None:
    # Body content is walked in true document order, so a table interleaved between two
    # headings must land under the heading that authored it (the first), not be appended
    # to whatever section is open when body iteration ends (previously the last).
    input_path = tmp_path / "interleaved.docx"
    docx = DocxDocument()
    docx.add_heading("First section", level=1)
    docx.add_paragraph("Intro paragraph.")
    table = docx.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell content."
    docx.add_heading("Second section", level=1)
    docx.add_paragraph("Later paragraph.")
    docx.save(input_path)

    document = load_docx(input_path)

    by_title = {section.title: section for section in document.sections}
    first_locators = [p.locator for p in by_title["First section"].paragraphs]
    second_locators = [p.locator for p in by_title["Second section"].paragraphs]

    assert "table:0:row:0:cell:0:p:0" in first_locators
    assert "body:p:1" in first_locators  # "Intro paragraph." keeps its true body index
    assert first_locators == ["body:p:1", "table:0:row:0:cell:0:p:0"]
    assert second_locators == ["body:p:3"]  # heading at p:2 consumes an index


def test_header_and_footer_get_dedicated_sections(tmp_path: Path) -> None:
    # Header/footer lines must not be misread as body prose appended to the trailing body
    # section: they belong in synthetic sections keyed by source, with locators unchanged.
    input_path = tmp_path / "with_header_footer.docx"
    docx = DocxDocument()
    docx.add_paragraph("Body prose.")
    section = docx.sections[0]
    section.header.paragraphs[0].text = "Header line."
    section.footer.paragraphs[0].text = "Footer line."
    docx.save(input_path)

    document = load_docx(input_path)

    header = next(s for s in document.sections if s.metadata.get("source") == "header")
    footer = next(s for s in document.sections if s.metadata.get("source") == "footer")
    assert [p.text for p in header.paragraphs] == ["Header line."]
    assert [p.text for p in footer.paragraphs] == ["Footer line."]
    assert header.paragraphs[0].locator == "header:0:p:0"
    assert footer.paragraphs[0].locator == "footer:0:p:0"
    # The synthetic section title must not fabricate an English word ("Header"/"Footer") into
    # the reviewable tree: that would leak language into the language-blind core and reach the
    # LLM as if it were document prose. The header/footer distinction lives only in the
    # machine-readable metadata["source"] key, exactly like a titleless body section.
    assert header.title is None
    assert footer.title is None
    # The body section must hold only the body prose, not the header/footer lines.
    body = next(s for s in document.sections if s.id == "s1")
    assert [p.text for p in body.paragraphs] == ["Body prose."]


def test_tracked_revision_only_in_a_header_is_detected(tmp_path: Path) -> None:
    # A tracked change living only in a header part must still be surfaced so the pipeline
    # can warn the human; scanning a fixed allowlist of parts missed headers/footers.
    input_path = tmp_path / "header_revision.docx"
    docx = DocxDocument()
    docx.add_paragraph("Body paragraph with no revisions.")
    docx.save(input_path)

    header_xml = (
        b"<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        b"<w:hdr xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
        b"<w:p><w:ins w:id='1' w:author='reviewer'><w:r><w:t>added</w:t></w:r></w:ins></w:p>"
        b"</w:hdr>"
    )
    with ZipFile(input_path, "a") as archive:
        archive.writestr("word/header1.xml", header_xml)

    document = load_docx(input_path)

    assert document.metadata["tracked_revisions_detected"] == "true"


def test_merged_table_cells_are_walked_exactly_once(tmp_path: Path) -> None:
    # row.cells yields a merged cell once per grid position it spans, so without dedup a
    # merged cell's paragraphs are emitted multiple times: reviewed twice, edited twice.
    input_path = tmp_path / "merged.docx"
    docx = DocxDocument()
    table = docx.add_table(rows=2, cols=3)
    for row in range(2):
        for col in range(3):
            table.cell(row, col).paragraphs[0].add_run(f"r{row}c{col}")
    table.cell(0, 0).merge(table.cell(0, 1))  # horizontal merge
    table.cell(0, 2).merge(table.cell(1, 2))  # vertical merge
    docx.save(input_path)

    document = load_docx(input_path)
    table_paragraphs = [
        paragraph
        for paragraph in document.iter_paragraphs()
        if (paragraph.locator or "").startswith("table:")
    ]

    # 4 physical cells survive the two merges; each merged cell keeps both its paragraphs
    # but is walked once, so 6 nodes total with no text (and no locator) emitted twice.
    assert len(table_paragraphs) == 6
    texts = [paragraph.text for paragraph in table_paragraphs]
    assert len(texts) == len(set(texts))
    locators = [paragraph.locator for paragraph in table_paragraphs]
    assert len(locators) == len(set(locators))


def test_read_footnotes_returns_only_content_notes(tmp_path: Path) -> None:
    # Word seeds every document with two structural footnotes (the separator lines drawn
    # above footnote text). They carry no authored content and must never surface.
    path = _write_footnotes_part(
        tmp_path / "notes.docx",
        "<w:footnote w:type='separator' w:id='-1'><w:p><w:r><w:separator/></w:r></w:p></w:footnote>"
        "<w:footnote w:type='continuationSeparator' w:id='0'>"
        "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>"
        "<w:footnote w:id='1'><w:p><w:r><w:t>First note.</w:t></w:r></w:p></w:footnote>"
        "<w:footnote w:id='2'><w:p><w:r><w:t>Second note.</w:t></w:r></w:p></w:footnote>",
    )

    assert read_footnotes(path) == [
        DocxFootnote(id="1", text="First note."),
        DocxFootnote(id="2", text="Second note."),
    ]


def test_read_footnotes_joins_runs_and_renders_breaks_and_tabs(tmp_path: Path) -> None:
    # Visible text spans multiple runs; tabs and line breaks are structural whitespace that
    # must render as \t / \n so callers see the note as the reader sees it.
    path = _write_footnotes_part(
        tmp_path / "rich.docx",
        "<w:footnote w:id='1'><w:p>"
        "<w:r><w:t>See</w:t></w:r><w:r><w:tab/><w:t>clause</w:t></w:r>"
        "<w:r><w:br/><w:t>below.</w:t></w:r>"
        "</w:p></w:footnote>",
    )

    assert read_footnotes(path) == [DocxFootnote(id="1", text="See\tclause\nbelow.")]


def test_read_footnotes_absent_part_returns_empty(tmp_path: Path) -> None:
    path = tmp_path / "plain.docx"
    docx = DocxDocument()
    docx.add_paragraph("No footnotes here.")
    docx.save(path)

    assert read_footnotes(path) == []


def test_read_footnotes_missing_or_corrupt_file_returns_empty(tmp_path: Path) -> None:
    assert read_footnotes(tmp_path / "does_not_exist.docx") == []
    corrupt = tmp_path / "corrupt.docx"
    corrupt.write_bytes(b"not a zip archive")
    assert read_footnotes(corrupt) == []
