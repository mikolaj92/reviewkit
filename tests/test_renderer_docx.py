from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

import pytest
from docx import Document as DocxDocument
from docx.oxml import OxmlElement

from reviewkit.actions import prepare_actions
from reviewkit.document import ParagraphNode, ReviewDocument, SectionNode
from reviewkit.models import (
    ActionStatus,
    ReviewAction,
    ReviewActionType,
    ReviewFinding,
    ReviewLocator,
    ReviewReference,
    ReviewResult,
    ReviewScope,
)
from reviewkit.parser_docx import load_docx
from reviewkit.profile import ActionPolicyConfig, ReviewProfile
from reviewkit.renderer_docx import (
    RenderIntegrityError,
    render_corrected_docx,
    render_reviewed_docx,
)

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# The core contract: reviewed.docx must mark EVERY review action. Trackable edits land as
# w:ins/w:del tracked changes; advisory actions land as labelled Word comments. This table
# enumerates the expected marker for every ReviewActionType so a regression that silently
# stopped emitting one is caught. (tracked tags expected in document.xml, comment label)
_EVERY_ACTION_MARKER = [
    (ReviewActionType.REPLACE_TEXT, {"ins", "del"}, "SUGGESTION"),
    (ReviewActionType.REPLACE, {"ins", "del"}, "SUGGESTION"),
    (ReviewActionType.DELETE_TEXT, {"del"}, "SUGGESTION"),
    (ReviewActionType.DELETE, {"del"}, "SUGGESTION"),
    (ReviewActionType.INSERT_TEXT, {"ins"}, "SUGGESTION"),
    (ReviewActionType.INSERT_BEFORE, {"ins"}, "SUGGESTION"),
    (ReviewActionType.INSERT_AFTER, {"ins"}, "SUGGESTION"),
    (ReviewActionType.COMMENT, set(), "COMMENT"),
    (ReviewActionType.QUESTION, set(), "QUESTION"),
    (ReviewActionType.RISK, set(), "RISK"),
    (ReviewActionType.SUGGESTION, set(), "SUGGESTION"),
    (ReviewActionType.PRAISE, set(), "PRAISE"),
    (ReviewActionType.SUMMARY, set(), "SUMMARY"),
    (ReviewActionType.FLAG, set(), "COMMENT"),
]


@pytest.mark.parametrize(
    ("action_type", "tracked_tags", "label"),
    _EVERY_ACTION_MARKER,
    ids=[case[0].value for case in _EVERY_ACTION_MARKER],
)
def test_reviewed_docx_marks_every_action_type(
    tmp_path: Path,
    action_type: ReviewActionType,
    tracked_tags: set[str],
    label: str,
) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("The quick brown fox jumps.")
    docx.save(input_path)

    document = load_docx(input_path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=action_type,
        node_id="p1",
        original_text="fox",
        replacement_text="cat",
        comment="Reviewer note.",
    )

    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    document_xml = _part_xml(reviewed_path, "word/document.xml")
    root = ElementTree.fromstring(document_xml)

    # Trackable edits produce the expected tracked-change markup; advisory actions must not.
    for tag in ("ins", "del"):
        found = root.find(f".//{_W}{tag}") is not None
        assert found is (tag in tracked_tags), f"{action_type.value}: unexpected w:{tag}={found}"

    # Every action - trackable or advisory - carries a labelled comment marker.
    comments = _comment_texts(reviewed_path)
    assert any(text.startswith(f"{label}:") for text in comments), (
        f"{action_type.value}: no comment labelled {label!r}; got {comments}"
    )


@pytest.mark.parametrize(
    ("status", "metadata", "label"),
    [
        (ActionStatus.CONFLICT, {}, "CONFLICT"),
        (ActionStatus.NEEDS_HUMAN_DECISION, {"blocked_from_corrected": True}, "HUMAN_DECISION"),
    ],
    ids=["conflict", "blocked_human_decision"],
)
def test_reviewed_docx_surfaces_unapplied_edits_as_labelled_comments(
    tmp_path: Path,
    status: ActionStatus,
    metadata: dict[str, object],
    label: str,
) -> None:
    # A writing edit that cannot be auto-applied (CONFLICT, or blocked from corrected) must
    # NOT become a silent tracked change - it surfaces as a labelled comment for a human.
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("The quick brown fox jumps.")
    docx.save(input_path)

    document = load_docx(input_path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id="p1",
        original_text="fox",
        replacement_text="cat",
        reason="Contested wording.",
        status=status,
        metadata=metadata,
    )

    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    root = ElementTree.fromstring(_part_xml(reviewed_path, "word/document.xml"))

    assert root.find(f".//{_W}ins") is None
    assert root.find(f".//{_W}del") is None
    comments = _comment_texts(reviewed_path)
    assert any(text.startswith(f"{label}:") for text in comments), (
        f"expected a comment labelled {label!r}; got {comments}"
    )


def test_reviewed_docx_surfaces_scope_level_conflict_as_a_comment(tmp_path: Path) -> None:
    # A section/document-scoped edit that ends up CONFLICT must still be surfaced. It used to
    # vanish (anchor returned None, and the scope-comment loops skip original_text actions),
    # losing the escalation for exactly the ambiguous, higher-scope case that needs a human.
    from reviewkit.document import ParagraphNode, ReviewDocument, SectionNode

    document = ReviewDocument(
        sections=[
            SectionNode(
                id="s1",
                paragraphs=[
                    ParagraphNode(id="p1", text="The quick brown fox jumps.", section_id="s1"),
                    ParagraphNode(id="p2", text="Another line entirely.", section_id="s1"),
                ],
            )
        ]
    )
    action = ReviewAction(
        scope=ReviewScope.SECTION,
        action_type=ReviewActionType.REPLACE,
        node_id="s1",
        original_text="fox",
        replacement_text="cat",
        reason="Ambiguous at section scope.",
        status=ActionStatus.CONFLICT,
    )

    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    root = ElementTree.fromstring(_part_xml(reviewed_path, "word/document.xml"))

    # Surfaced as a labelled comment, never a silent tracked change.
    assert root.find(f".//{_W}ins") is None
    assert root.find(f".//{_W}del") is None
    comments = _comment_texts(reviewed_path)
    assert any(text.startswith("CONFLICT:") for text in comments), comments


def test_reviewed_docx_renders_delete_and_insert_text_revisions(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Alpha beta gamma.")
    docx.save(input_path)

    document = load_docx(input_path)
    paragraph = document.sections[0].paragraphs[0]
    actions = [
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.DELETE_TEXT,
            node_id=paragraph.id,
            original_text="beta",
            locator=ReviewLocator(node_id=paragraph.id, char_start=6, char_end=10),
            reason="Remove redundant word.",
            status=ActionStatus.NOT_APPLIED,
        ),
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.INSERT_AFTER,
            node_id=paragraph.id,
            original_text="gamma",
            replacement_text="!",
            locator=ReviewLocator(node_id=paragraph.id, char_start=11, char_end=16),
            reason="Add emphasis.",
            status=ActionStatus.NOT_APPLIED,
        ),
    ]

    reviewed_path = render_reviewed_docx(document, actions, tmp_path / "reviewed.docx")
    document_xml = _part_xml(reviewed_path, "word/document.xml")

    assert _revision_texts(document_xml, "del", "delText") == ["beta"]
    assert _revision_texts(document_xml, "ins", "t") == ["!"]


def test_insert_honors_locator_when_anchor_text_repeats(tmp_path: Path) -> None:
    # "beta" occurs twice; the locator disambiguates to the *second* one. The renderer
    # must honor the locator (like apply_action_to_text) instead of blindly inserting
    # before the first find() match, which would land the revision on the wrong "beta".
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("beta alpha beta gamma.")
    docx.save(input_path)

    document = load_docx(input_path)
    paragraph = document.sections[0].paragraphs[0]
    actions = [
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.INSERT_BEFORE,
            node_id=paragraph.id,
            original_text="beta",
            replacement_text="X",
            locator=ReviewLocator(node_id=paragraph.id, char_start=11, char_end=15),
            reason="Mark the second occurrence.",
            status=ActionStatus.NOT_APPLIED,
        ),
    ]

    reviewed_path = render_reviewed_docx(document, actions, tmp_path / "reviewed.docx")
    document_xml = _part_xml(reviewed_path, "word/document.xml")

    assert _accepted_paragraph_text(document_xml) == "beta alpha Xbeta gamma."


def test_reviewed_docx_patches_original_and_preserves_run_formatting(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph()
    paragraph.add_run("Plain ")
    target = paragraph.add_run("target")
    target.bold = True
    paragraph.add_run(" text")
    docx.save(input_path)

    document = load_docx(input_path)
    reviewed_path = render_reviewed_docx(
        document,
        [
            ReviewAction(
                scope=ReviewScope.PARAGRAPH,
                action_type=ReviewActionType.REPLACE,
                node_id="p1",
                original_text="target",
                replacement_text="replacement",
                reason="Use clearer wording.",
                status=ActionStatus.NOT_APPLIED,
            )
        ],
        tmp_path / "reviewed.docx",
    )

    document_xml = _part_xml(reviewed_path, "word/document.xml")
    root = ElementTree.fromstring(document_xml)
    paragraph_xml = root.find(f".//{_W}p")
    assert paragraph_xml is not None
    assert [child.tag for child in paragraph_xml] == [
        f"{_W}r",
        f"{_W}commentRangeStart",
        f"{_W}del",
        f"{_W}ins",
        f"{_W}commentRangeEnd",
        f"{_W}r",
        f"{_W}r",
    ]
    assert _revision_texts(document_xml, "del", "delText") == ["target"]
    assert _revision_texts(document_xml, "ins", "t") == ["replacement"]
    assert root.find(f".//{_W}del/{_W}r/{_W}rPr/{_W}b") is not None
    assert "[DELETE:" not in document_xml


def test_reviewed_docx_comment_carries_a_references_line(tmp_path: Path) -> None:
    # An action's references are the citation carrier for downstream users: the comment body
    # in word/comments.xml must include a "References:" line listing each reference by label
    # (falling back to source when no label is set).
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("The quick brown fox jumps.")
    docx.save(input_path)

    document = load_docx(input_path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.COMMENT,
        node_id="p1",
        comment="Grounded observation.",
        references=[
            ReviewReference(source="KC", label="art. 385(1)"),
            ReviewReference(source="unlabelled-source"),
        ],
    )

    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")

    comments = _comment_texts(reviewed_path)
    assert any(
        "References: art. 385(1), unlabelled-source" in text for text in comments
    ), comments


def test_reviewed_docx_interleaves_several_tracked_edits_in_one_paragraph(tmp_path: Path) -> None:
    # Several trackable edits landing on ONE paragraph must interleave correctly: each edit's
    # w:del/w:ins pair sits inside its own commentRangeStart/commentRangeEnd, untouched text
    # survives between them, and every revision carries a fresh id (strictly increasing in
    # document order for edits applied left to right).
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Alpha beta gamma delta.")
    docx.save(input_path)

    document = load_docx(input_path)
    actions = [
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.REPLACE,
            node_id="p1",
            original_text="beta",
            replacement_text="B",
            reason="First edit.",
            status=ActionStatus.NOT_APPLIED,
        ),
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.DELETE_TEXT,
            node_id="p1",
            original_text=" gamma",
            reason="Second edit.",
            status=ActionStatus.NOT_APPLIED,
        ),
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.INSERT_AFTER,
            node_id="p1",
            original_text="delta",
            replacement_text="!",
            reason="Third edit.",
            status=ActionStatus.NOT_APPLIED,
        ),
    ]

    reviewed_path = render_reviewed_docx(document, actions, tmp_path / "reviewed.docx")
    document_xml = _part_xml(reviewed_path, "word/document.xml")
    root = ElementTree.fromstring(document_xml)
    paragraph_xml = root.find(f".//{_W}p")
    assert paragraph_xml is not None

    # Exact interleaving: kept runs between the edits, each edit wrapped in its own
    # comment range (replace = del+ins, delete = del only, insert = ins only).
    assert [child.tag for child in paragraph_xml] == [
        f"{_W}r",  # "Alpha "
        f"{_W}commentRangeStart",
        f"{_W}del",  # "beta"
        f"{_W}ins",  # "B"
        f"{_W}commentRangeEnd",
        f"{_W}r",  # comment reference
        f"{_W}commentRangeStart",
        f"{_W}del",  # " gamma"
        f"{_W}commentRangeEnd",
        f"{_W}r",  # comment reference
        f"{_W}r",  # " delta"
        f"{_W}commentRangeStart",
        f"{_W}ins",  # "!"
        f"{_W}commentRangeEnd",
        f"{_W}r",  # comment reference
        f"{_W}r",  # "."
    ]
    assert _revision_texts(document_xml, "del", "delText") == ["beta", " gamma"]
    assert _revision_texts(document_xml, "ins", "t") == ["B", "!"]

    # Every revision gets a fresh id: strictly increasing in document order.
    revision_ids = [
        int(child.get(f"{_W}id") or -1)
        for child in paragraph_xml
        if child.tag in {f"{_W}del", f"{_W}ins"}
    ]
    assert len(revision_ids) == 4
    assert all(earlier < later for earlier, later in zip(revision_ids, revision_ids[1:]))

    # Comment ranges pair up: starts and ends carry the same ids in the same order,
    # one distinct comment per edit.
    start_ids = [
        child.get(f"{_W}id")
        for child in paragraph_xml
        if child.tag == f"{_W}commentRangeStart"
    ]
    end_ids = [
        child.get(f"{_W}id") for child in paragraph_xml if child.tag == f"{_W}commentRangeEnd"
    ]
    assert start_ids == end_ids
    assert len(set(start_ids)) == 3


def test_reviewed_docx_patches_table_header_and_footer_paragraphs(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Body text.")
    table = docx.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table bład."
    section = docx.sections[0]
    section.header.paragraphs[0].text = "Header bład."
    section.footer.paragraphs[0].text = "Footer bład."
    docx.save(input_path)

    document = load_docx(input_path)
    actions = [
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.REPLACE,
            node_id=paragraph.id,
            original_text="bład",
            replacement_text="błąd",
            reason="Fix typo.",
            status=ActionStatus.APPLIED,
        )
        for paragraph in document.iter_paragraphs()
        if "bład" in paragraph.text
    ]

    reviewed_path = render_reviewed_docx(document, actions, tmp_path / "reviewed.docx")

    document_xml = _part_xml(reviewed_path, "word/document.xml")
    header_xml = _part_xml(reviewed_path, "word/header1.xml")
    footer_xml = _part_xml(reviewed_path, "word/footer1.xml")
    assert _revision_texts(document_xml, "del", "delText") == ["bład"]
    assert _revision_texts(document_xml, "ins", "t") == ["błąd"]
    assert _revision_texts(header_xml, "del", "delText") == ["bład"]
    assert _revision_texts(header_xml, "ins", "t") == ["błąd"]
    assert _revision_texts(footer_xml, "del", "delText") == ["bład"]
    assert _revision_texts(footer_xml, "ins", "t") == ["błąd"]


def test_reviewed_docx_stamps_caller_supplied_reviewer_identity(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Plain target text")
    docx.save(input_path)

    document = load_docx(input_path)
    reviewed_path = render_reviewed_docx(
        document,
        [
            ReviewAction(
                scope=ReviewScope.PARAGRAPH,
                action_type=ReviewActionType.REPLACE,
                node_id="p1",
                original_text="target",
                replacement_text="replacement",
                reason="Use clearer wording.",
                status=ActionStatus.NOT_APPLIED,
            )
        ],
        tmp_path / "reviewed.docx",
        comment_author="External Reviewer",
        comment_initials="ER",
    )

    document_xml = _part_xml(reviewed_path, "word/document.xml")
    revision = ElementTree.fromstring(document_xml).find(f".//{_W}ins")
    assert revision is not None
    assert revision.get(f"{_W}author") == "External Reviewer"

    comments_xml = _part_xml(reviewed_path, "word/comments.xml")
    comment = ElementTree.fromstring(comments_xml).find(f".//{_W}comment")
    assert comment is not None
    assert comment.get(f"{_W}author") == "External Reviewer"
    assert comment.get(f"{_W}initials") == "ER"


def test_reviewed_docx_uses_neutral_default_reviewer_identity(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Plain target text")
    docx.save(input_path)

    document = load_docx(input_path)
    reviewed_path = render_reviewed_docx(
        document,
        [
            ReviewAction(
                scope=ReviewScope.PARAGRAPH,
                action_type=ReviewActionType.REPLACE,
                node_id="p1",
                original_text="target",
                replacement_text="replacement",
                reason="Use clearer wording.",
                status=ActionStatus.NOT_APPLIED,
            )
        ],
        tmp_path / "reviewed.docx",
    )

    document_xml = _part_xml(reviewed_path, "word/document.xml")
    revision = ElementTree.fromstring(document_xml).find(f".//{_W}ins")
    assert revision is not None
    assert revision.get(f"{_W}author") == "Reviewer"

    comments_xml = _part_xml(reviewed_path, "word/comments.xml")
    comment = ElementTree.fromstring(comments_xml).find(f".//{_W}comment")
    assert comment is not None
    assert comment.get(f"{_W}author") == "Reviewer"
    assert comment.get(f"{_W}initials") == "RV"


def test_reviewed_docx_revision_dates_are_deterministic_by_default(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Plain target text")
    docx.save(input_path)

    document = load_docx(input_path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id="p1",
        original_text="target",
        replacement_text="replacement",
        reason="Use clearer wording.",
        status=ActionStatus.NOT_APPLIED,
    )

    first = render_reviewed_docx(document, [action], tmp_path / "first.docx")
    second = render_reviewed_docx(document, [action], tmp_path / "second.docx")

    first_date = ElementTree.fromstring(_part_xml(first, "word/document.xml")).find(f".//{_W}ins")
    second_date = ElementTree.fromstring(_part_xml(second, "word/document.xml")).find(f".//{_W}ins")
    assert first_date is not None and second_date is not None
    assert first_date.get(f"{_W}date") == second_date.get(f"{_W}date")
    assert first_date.get(f"{_W}date") == "1970-01-01T00:00:00+00:00"


def test_reviewed_docx_accepts_an_explicit_revision_timestamp(tmp_path: Path) -> None:
    from datetime import UTC, datetime

    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Plain target text")
    docx.save(input_path)

    document = load_docx(input_path)
    stamp = datetime(2026, 7, 3, 9, 30, tzinfo=UTC)
    reviewed_path = render_reviewed_docx(
        document,
        [
            ReviewAction(
                scope=ReviewScope.PARAGRAPH,
                action_type=ReviewActionType.REPLACE,
                node_id="p1",
                original_text="target",
                replacement_text="replacement",
                reason="Use clearer wording.",
                status=ActionStatus.NOT_APPLIED,
            )
        ],
        tmp_path / "reviewed.docx",
        revision_timestamp=stamp,
    )

    revision = ElementTree.fromstring(_part_xml(reviewed_path, "word/document.xml")).find(
        f".//{_W}ins"
    )
    assert revision is not None
    assert revision.get(f"{_W}date") == "2026-07-03T09:30:00+00:00"


def test_reviewed_docx_comment_dates_are_deterministic_by_default(tmp_path: Path) -> None:
    # python-docx stamps comment w:date with wall-clock time; without the override the
    # comments.xml part differs on every run. Two renders of identical inputs must match.
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Plain target text")
    docx.save(input_path)

    document = load_docx(input_path)
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id="p1",
        original_text="target",
        replacement_text="replacement",
        reason="Use clearer wording.",
        status=ActionStatus.NOT_APPLIED,
    )

    first = render_reviewed_docx(document, [action], tmp_path / "first.docx")
    second = render_reviewed_docx(document, [action], tmp_path / "second.docx")

    first_comments = _part_xml(first, "word/comments.xml")
    second_comments = _part_xml(second, "word/comments.xml")
    assert first_comments == second_comments

    comment = ElementTree.fromstring(first_comments).find(f".//{_W}comment")
    assert comment is not None
    assert comment.get(f"{_W}date") == "1970-01-01T00:00:00+00:00"


def test_reviewed_docx_comment_uses_explicit_revision_timestamp(tmp_path: Path) -> None:
    from datetime import UTC, datetime

    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Plain target text")
    docx.save(input_path)

    document = load_docx(input_path)
    stamp = datetime(2026, 7, 3, 9, 30, tzinfo=UTC)
    reviewed_path = render_reviewed_docx(
        document,
        [
            ReviewAction(
                scope=ReviewScope.PARAGRAPH,
                action_type=ReviewActionType.REPLACE,
                node_id="p1",
                original_text="target",
                replacement_text="replacement",
                reason="Use clearer wording.",
                status=ActionStatus.NOT_APPLIED,
            )
        ],
        tmp_path / "reviewed.docx",
        revision_timestamp=stamp,
    )

    comment = ElementTree.fromstring(_part_xml(reviewed_path, "word/comments.xml")).find(
        f".//{_W}comment"
    )
    assert comment is not None
    assert comment.get(f"{_W}date") == "2026-07-03T09:30:00+00:00"


def test_corrected_docx_preserves_original_structure_and_formatting(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_heading("Rozdział pierwszy", level=1)
    paragraph = docx.add_paragraph()
    paragraph.add_run("Ala ma ")
    kept = paragraph.add_run("kota")
    kept.bold = True
    paragraph.add_run(" oraz bład.")
    table = docx.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Tabela bład."
    docx.save(input_path)

    document = load_docx(input_path)
    actions = [
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.REPLACE,
            node_id=paragraph_node.id,
            original_text="bład",
            replacement_text="błąd",
            reason="Fix typo.",
            status=ActionStatus.APPLIED,
        )
        for paragraph_node in document.iter_paragraphs()
        if "bład" in paragraph_node.text
    ]

    corrected_path = render_corrected_docx(document, actions, tmp_path / "corrected.docx")

    corrected = DocxDocument(corrected_path)
    # Structure preserved: the heading and the table survive the round-trip.
    assert any(p.style.name.startswith("Heading") for p in corrected.paragraphs)
    assert any("Rozdział pierwszy" == p.text for p in corrected.paragraphs)
    assert len(corrected.tables) == 1
    assert corrected.tables[0].cell(0, 0).text == "Tabela błąd."
    # Edit applied cleanly with no tracked-change markup.
    body_text = "\n".join(p.text for p in corrected.paragraphs)
    assert "Ala ma kota oraz błąd." in body_text
    document_xml = _part_xml(corrected_path, "word/document.xml")
    assert "<w:ins" not in document_xml
    assert "<w:del" not in document_xml
    # Run-level formatting on the untouched bold run is retained.
    root = ElementTree.fromstring(document_xml)
    assert root.find(f".//{_W}r/{_W}rPr/{_W}b") is not None


def test_locator_edit_aligns_offsets_past_leading_whitespace(tmp_path: Path) -> None:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph()
    paragraph.add_run("   ")  # leading whitespace preserved in the source run
    paragraph.add_run("Hello world.")
    docx.save(input_path)

    document = load_docx(input_path)
    paragraph_node = document.sections[0].paragraphs[0]
    assert paragraph_node.text == "Hello world."
    # Offsets are relative to the stripped node text: "world" is [6, 11).
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id=paragraph_node.id,
        original_text="world",
        replacement_text="planet",
        locator=ReviewLocator(node_id=paragraph_node.id, char_start=6, char_end=11),
        reason="Clarify.",
        status=ActionStatus.APPLIED,
    )

    corrected_path = render_corrected_docx(document, [action], tmp_path / "corrected.docx")
    corrected = DocxDocument(corrected_path)
    corrected_text = corrected.paragraphs[0].text
    # The edit lands on "world" despite the 3-char leading whitespace offset skew.
    assert corrected_text.strip() == "Hello planet."


def test_edited_paragraph_preserves_images_hyperlinks_tabs_and_breaks(tmp_path: Path) -> None:
    from docx.oxml import OxmlElement

    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph()
    paragraph.add_run("Ala ma ")
    image_run = paragraph.add_run()  # inline image lives in its own run
    image_run._r.append(OxmlElement("w:drawing"))
    paragraph.add_run("kota")
    tab_run = paragraph.add_run()
    tab_run._r.append(OxmlElement("w:tab"))
    paragraph.add_run("oraz")
    break_run = paragraph.add_run()
    break_run._r.append(OxmlElement("w:br"))
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink_run = OxmlElement("w:r")
    hyperlink_text = OxmlElement("w:t")
    hyperlink_text.text = "link"
    hyperlink_run.append(hyperlink_text)
    hyperlink.append(hyperlink_run)
    paragraph._p.append(hyperlink)
    paragraph.add_run(" koniec")
    docx.save(input_path)

    document = load_docx(input_path)
    node = document.sections[0].paragraphs[0]
    assert node.text == "Ala ma kota\toraz\nlink koniec"
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE,
        node_id=node.id,
        original_text="kota",
        replacement_text="pies",
        reason="Change the animal.",
        status=ActionStatus.APPLIED,
    )

    # corrected.docx: the edit applies cleanly and every inline object survives.
    corrected_path = render_corrected_docx(document, [action], tmp_path / "corrected.docx")
    corrected_xml = _part_xml(corrected_path, "word/document.xml")
    assert "<w:drawing" in corrected_xml
    assert "<w:tab" in corrected_xml
    assert "<w:br" in corrected_xml
    assert "<w:hyperlink" in corrected_xml
    corrected_text = DocxDocument(corrected_path).paragraphs[0].text
    assert "pies" in corrected_text
    assert "kota" not in corrected_text
    assert "link" in corrected_text
    assert "<w:ins" not in corrected_xml and "<w:del" not in corrected_xml

    # reviewed.docx: tracked change is applied and every inline object survives.
    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    reviewed_xml = _part_xml(reviewed_path, "word/document.xml")
    assert "<w:drawing" in reviewed_xml
    assert "<w:tab" in reviewed_xml
    assert "<w:br" in reviewed_xml
    assert "<w:hyperlink" in reviewed_xml
    assert _revision_texts(reviewed_xml, "del", "delText") == ["kota"]
    assert _revision_texts(reviewed_xml, "ins", "t") == ["pies"]


def test_reviewed_and_corrected_artifacts_are_byte_reproducible(tmp_path: Path) -> None:
    # Determinism is a core contract: identical document + actions must yield byte-identical
    # artifacts so downstream diffing/caching stays stable. A nondeterministic comment id,
    # revision id, dict ordering or timestamp would break this with no other failing test.
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("The quick brown fox jumps.")
    docx.save(input_path)

    document = load_docx(input_path)
    actions = [
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.REPLACE_TEXT,
            node_id="p1",
            original_text="fox",
            replacement_text="cat",
            comment="Prefer cat.",
            apply_hint=True,
        ),
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.COMMENT,
            node_id="p1",
            comment="Advisory note.",
        ),
    ]

    def reviewed_parts(suffix: str) -> tuple[bytes, bytes]:
        path = render_reviewed_docx(document, actions, tmp_path / f"reviewed-{suffix}.docx")
        with ZipFile(path) as archive:
            return archive.read("word/document.xml"), archive.read("word/comments.xml")

    first_document, first_comments = reviewed_parts("a")
    second_document, second_comments = reviewed_parts("b")
    assert first_document == second_document
    assert first_comments == second_comments

    def corrected_document(suffix: str) -> bytes:
        path = render_corrected_docx(document, actions, tmp_path / f"corrected-{suffix}.docx")
        with ZipFile(path) as archive:
            return archive.read("word/document.xml")

    assert corrected_document("a") == corrected_document("b")

    result = ReviewResult(
        findings=[ReviewFinding(node_id="p1", title="Word choice", description="fox -> cat")],
        actions=actions,
    )
    first_json = result.save_json(tmp_path / "report-a.json").read_bytes()
    second_json = result.save_json(tmp_path / "report-b.json").read_bytes()
    assert first_json == second_json


def test_rendered_docx_packages_are_byte_reproducible_whole_file(tmp_path: Path) -> None:
    # The test above compares inner part XML, which stays identical regardless of the zip
    # container's per-entry mtime. python-docx stamps that mtime with the wall clock, so the
    # WHOLE-FILE bytes only match run-to-run once those timestamps are pinned. Assert both the
    # pinned epoch (robust, timing-independent) and whole-file byte equality (the contract).
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("The quick brown fox jumps.")
    docx.save(input_path)

    document = load_docx(input_path)
    actions = [
        ReviewAction(
            scope=ReviewScope.PARAGRAPH,
            action_type=ReviewActionType.REPLACE_TEXT,
            node_id="p1",
            original_text="fox",
            replacement_text="cat",
            comment="Prefer cat.",
            apply_hint=True,
        ),
    ]

    for render in (render_reviewed_docx, render_corrected_docx):
        first = render(document, actions, tmp_path / f"{render.__name__}-a.docx")
        second = render(document, actions, tmp_path / f"{render.__name__}-b.docx")
        with ZipFile(first) as archive:
            assert all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in archive.infolist())
        assert first.read_bytes() == second.read_bytes()


def test_precise_comment_anchor_survives_opaque_segment_before_the_quote(tmp_path: Path) -> None:
    # A comment carrying original_text anchors to just the quoted span. When an opaque inline
    # object (here a line break) precedes the quote, _mark_text_comment must count its visible
    # width the same way _visible_text does; otherwise the two coordinate systems desync, the
    # anchor indexes come back empty, and the comment silently degrades to whole-paragraph.
    from docx.oxml import OxmlElement

    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph()
    paragraph.add_run("Ala")
    break_run = paragraph.add_run()
    break_run._r.append(OxmlElement("w:br"))
    paragraph.add_run("beta target here")
    docx.save(input_path)

    document = load_docx(input_path)
    node = document.sections[0].paragraphs[0]
    assert node.text == "Ala\nbeta target here"

    reviewed_path = render_reviewed_docx(
        document,
        [
            ReviewAction(
                scope=ReviewScope.PARAGRAPH,
                action_type=ReviewActionType.RISK,
                node_id=node.id,
                original_text="target",
                comment="Ambiguous term.",
            )
        ],
        tmp_path / "reviewed.docx",
    )

    paragraph_xml = ElementTree.fromstring(
        _part_xml(reviewed_path, "word/document.xml")
    ).find(f".//{_W}p")
    assert paragraph_xml is not None
    # Only "target" is inside the comment range - not the whole paragraph.
    assert _commented_run_text(paragraph_xml) == "target"


# --- Render integrity (issue #142): a writing action that must leave a trace can never
# --- be skipped silently - every miss raises a typed RenderIntegrityError naming the
# --- action, so consumers can fail closed without re-parsing the artifact.


def _saved_docx(tmp_path: Path, *paragraphs: str) -> Path:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    for text in paragraphs:
        docx.add_paragraph(text)
    docx.save(input_path)
    return input_path


def _writing_action(**overrides: object) -> ReviewAction:
    fields: dict = {
        "scope": ReviewScope.PARAGRAPH,
        "action_type": ReviewActionType.REPLACE,
        "node_id": "p1",
        "original_text": "fox",
        "replacement_text": "cat",
        "status": ActionStatus.APPLIED,
    }
    fields.update(overrides)
    return ReviewAction(**fields)


@pytest.mark.parametrize("renderer", [render_reviewed_docx, render_corrected_docx])
def test_render_raises_when_writing_action_routes_to_no_paragraph(tmp_path, renderer) -> None:
    # An action whose node_id names no paragraph/sentence/scope is never picked up by
    # actions_for_paragraph, so it used to vanish from both artifacts without a trace.
    document = load_docx(_saved_docx(tmp_path, "The quick brown fox jumps."))
    action = _writing_action(id="a-ghost", node_id="ghost-node")

    with pytest.raises(RenderIntegrityError) as excinfo:
        renderer(document, [action], tmp_path / "out.docx")
    assert "a-ghost" in str(excinfo.value)
    assert "ghost-node" in str(excinfo.value)


@pytest.mark.parametrize("renderer", [render_reviewed_docx, render_corrected_docx])
def test_conflict_action_on_unknown_node_still_renders_without_error(tmp_path, renderer) -> None:
    # prepare_actions demotes unknown-node actions to CONFLICT; those are advisory-only
    # (never trackable, never applied to corrected), so rendering them must NOT raise.
    document = load_docx(_saved_docx(tmp_path, "The quick brown fox jumps."))
    action = _writing_action(id="a-conflict", node_id="ghost-node", status=ActionStatus.CONFLICT)

    assert renderer(document, [action], tmp_path / "out.docx").exists()


@pytest.mark.parametrize(
    "action_type",
    [ReviewActionType.REPLACE, ReviewActionType.DELETE_TEXT, ReviewActionType.INSERT_AFTER],
)
@pytest.mark.parametrize("status", [ActionStatus.APPLIED, ActionStatus.NOT_APPLIED])
def test_reviewed_docx_raises_when_tracked_edit_fails_to_anchor(
    tmp_path: Path, action_type: ReviewActionType, status: ActionStatus
) -> None:
    # The original issue-#142 hole: _track_action could not find original_text in the
    # paragraph and returned unchanged segments, silently dropping the tracked change.
    document = load_docx(_saved_docx(tmp_path, "The quick brown fox jumps."))
    action = _writing_action(
        id="a-miss", action_type=action_type, original_text="unicorn", status=status
    )

    with pytest.raises(RenderIntegrityError) as excinfo:
        render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    assert "a-miss" in str(excinfo.value)
    assert "'p1'" in str(excinfo.value)


@pytest.mark.parametrize(
    "action_type",
    [ReviewActionType.REPLACE, ReviewActionType.DELETE_TEXT, ReviewActionType.INSERT_AFTER],
)
def test_corrected_docx_raises_when_applied_edit_fails_to_anchor(
    tmp_path: Path, action_type: ReviewActionType
) -> None:
    # Sibling hole in the clean rewrite: _apply_clean_corrections used the same silent
    # _track_action skip, so corrected.docx shipped without an edit the report claimed.
    document = load_docx(_saved_docx(tmp_path, "The quick brown fox jumps."))
    action = _writing_action(id="a-miss", action_type=action_type, original_text="unicorn")

    with pytest.raises(RenderIntegrityError) as excinfo:
        render_corrected_docx(document, [action], tmp_path / "corrected.docx")
    assert "a-miss" in str(excinfo.value)


def test_corrected_docx_without_source_raises_when_applied_edit_fails_to_anchor(
    tmp_path: Path,
) -> None:
    # No-source documents take the apply_corrections_to_text path, where an unmatched
    # original_text used to be a silent str.replace no-op.
    from reviewkit.document import ParagraphNode, ReviewDocument, SectionNode

    document = ReviewDocument(
        sections=[
            SectionNode(
                id="s1",
                paragraphs=[
                    ParagraphNode(id="p1", text="The quick brown fox jumps.", section_id="s1")
                ],
            )
        ]
    )
    action = _writing_action(id="a-miss", original_text="unicorn")

    with pytest.raises(RenderIntegrityError) as excinfo:
        render_corrected_docx(document, [action], tmp_path / "corrected.docx")
    assert "a-miss" in str(excinfo.value)


@pytest.mark.parametrize("renderer", [render_reviewed_docx, render_corrected_docx])
def test_render_raises_when_paragraph_locator_does_not_resolve(tmp_path, renderer) -> None:
    # With a source document every parsed locator must resolve; the old fallback landed
    # the edit in a detached appended paragraph while the real one stayed untouched.
    document = load_docx(_saved_docx(tmp_path, "The quick brown fox jumps."))
    document.sections[0].paragraphs[0].locator = "body:p:99"
    action = _writing_action(id="a-locator")

    with pytest.raises(RenderIntegrityError) as excinfo:
        renderer(document, [action], tmp_path / "out.docx")
    assert "a-locator" in str(excinfo.value)
    assert "body:p:99" in str(excinfo.value)


def test_overlap_consumed_suggestion_degrades_to_comment_not_error(tmp_path: Path) -> None:
    # Two NOT_APPLIED suggestions on overlapping spans are legitimate reviewer output
    # (prepare_actions only demotes overlapping APPLIED edits). The first tracked change
    # consumes the second one's anchor; that documented degrade to a labelled comment
    # must survive the fail-closed guards - it anchored fine in the pristine text.
    document = load_docx(_saved_docx(tmp_path, "The quick brown fox jumps."))
    actions = [
        _writing_action(
            id="a-first",
            original_text="brown fox",
            replacement_text="red cat",
            status=ActionStatus.NOT_APPLIED,
        ),
        _writing_action(id="a-second", status=ActionStatus.NOT_APPLIED),
    ]

    reviewed_path = render_reviewed_docx(document, actions, tmp_path / "reviewed.docx")

    document_xml = _part_xml(reviewed_path, "word/document.xml")
    assert _revision_texts(document_xml, "del", "delText") == ["brown fox"]
    # The consumed suggestion still surfaces as a labelled comment carrying its content.
    comments = _comment_texts(reviewed_path)
    assert any("Original: 'fox'" in text for text in comments), comments


def _commented_run_text(paragraph_xml: ElementTree.Element) -> str:
    inside = False
    parts: list[str] = []
    for child in paragraph_xml:
        if child.tag == f"{_W}commentRangeStart":
            inside = True
        elif child.tag == f"{_W}commentRangeEnd":
            inside = False
        elif inside and child.tag == f"{_W}r":
            parts.extend(text.text or "" for text in child.iter(f"{_W}t"))
    return "".join(parts)


def _part_xml(path: Path, member: str) -> str:
    with ZipFile(path) as archive:
        return archive.read(member).decode()


def _comment_texts(path: Path) -> list[str]:
    root = ElementTree.fromstring(_part_xml(path, "word/comments.xml"))
    return ["".join(comment.itertext()) for comment in root.findall(f".//{_W}comment")]


def _accepted_paragraph_text(xml: str) -> str:
    # Visible text of the first paragraph with all insertions accepted: every w:t in
    # document order (w:delText under w:del is a different tag, so deletions drop out).
    root = ElementTree.fromstring(xml)
    paragraph = root.find(f".//{_W}p")
    assert paragraph is not None
    return "".join(text.text or "" for text in paragraph.iter(f"{_W}t"))


def _revision_texts(xml: str, revision_tag: str, text_tag: str) -> list[str]:
    root = ElementTree.fromstring(xml)
    return [
        "".join(element.itertext())
        for element in root.findall(f".//{_W}{revision_tag}")
        if element.find(f".//{_W}{text_tag}") is not None
    ]


# --- Fail-closed pins: anchor consumption, opaque inline content, application order ---


def _auto_profile() -> ReviewProfile:
    return ReviewProfile(
        name="generic",
        language="en",
        document_type="generic document",
        reviewer_role="generic reviewer",
        action_policy=ActionPolicyConfig(
            apply_policy={"safe_edit": "apply"},
            require_llm_apply_hint=True,
            min_confidence_for_auto_apply=0.85,
            max_severity_for_auto_apply="medium",
        ),
    )


def _safe_edit(**overrides: object) -> ReviewAction:
    fields: dict = {
        "scope": ReviewScope.PARAGRAPH,
        "action_type": ReviewActionType.REPLACE_TEXT,
        "node_id": "p1",
        "category": "safe_edit",
        "confidence": 1.0,
        "apply_hint": True,
    }
    fields.update(overrides)
    return ReviewAction(**fields)


def _tabbed_docx(tmp_path: Path) -> Path:
    # One paragraph reading "A\tB" where the tab is a real w:tab element.
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph()
    paragraph.add_run("A")
    paragraph.add_run().add_tab()
    paragraph.add_run("B")
    docx.save(input_path)
    return input_path


def _hyperlink_docx(tmp_path: Path) -> Path:
    # One paragraph reading "See the appendix for details." where "the appendix"
    # is the text of a real w:hyperlink child (python-docx includes it in .text).
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    paragraph = docx.add_paragraph("See ")
    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "the appendix"
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    paragraph.add_run(" for details.")
    docx.save(input_path)
    return input_path


def test_corrected_no_source_raises_when_earlier_applied_edit_consumes_anchor(
    tmp_path: Path,
) -> None:
    # No-source path: each anchor must be checked against the EVOLVING text in
    # application order. Deleting "abc" from "abcabc" leaves "abc", which consumes
    # the second APPLIED delete's "bca" anchor - str.replace would silently no-op
    # it while the report claims it APPLIED, so render must raise instead.
    document = ReviewDocument(
        sections=[
            SectionNode(
                id="s1",
                paragraphs=[ParagraphNode(id="p1", text="abcabc", section_id="s1")],
            )
        ]
    )
    actions = [
        _writing_action(
            id="a-first",
            action_type=ReviewActionType.DELETE_TEXT,
            original_text="abc",
            replacement_text=None,
        ),
        _writing_action(
            id="a-second",
            action_type=ReviewActionType.DELETE_TEXT,
            original_text="bca",
            replacement_text=None,
        ),
    ]

    with pytest.raises(RenderIntegrityError, match="a-second"):
        render_corrected_docx(document, actions, tmp_path / "corrected.docx")


@pytest.mark.parametrize("renderer", [render_reviewed_docx, render_corrected_docx])
@pytest.mark.parametrize(
    ("original", "replacement"),
    [("\t", " - "), ("A\tB", "AB")],
    ids=["tab-only-range", "range-spanning-tab"],
)
def test_applied_replace_covering_opaque_tab_raises_instead_of_misrendering(
    tmp_path: Path, renderer, original: str, replacement: str
) -> None:
    # An APPLIED replace whose range covers opaque inline content (a real w:tab)
    # cannot be honored: the tab survives verbatim, so marking/applying only the
    # editable remainder would ship a wrong paragraph (historically the insertion
    # even landed at the paragraph END). A hand-built APPLIED action that bypasses
    # prepare must fail closed in BOTH renderers, never mis-render.
    document = load_docx(_tabbed_docx(tmp_path))
    action = _writing_action(
        id="a-tab", original_text=original, replacement_text=replacement
    )

    with pytest.raises(RenderIntegrityError, match="a-tab"):
        renderer(document, [action], tmp_path / "out.docx")


def test_prepare_demotes_edit_over_tab_and_both_artifacts_still_render(
    tmp_path: Path,
) -> None:
    # The supported flow: prepare detects that the edit's span covers the paragraph's
    # opaque tab (parser-recorded opaque_ranges), demotes it to CONFLICT, and both
    # artifacts render - the tab intact, the edit surfaced as a labelled comment.
    document = load_docx(_tabbed_docx(tmp_path))
    paragraph = next(document.iter_paragraphs())
    assert paragraph.text == "A\tB"
    assert paragraph.opaque_ranges == [(1, 2)]

    action = _safe_edit(original_text="\t", replacement_text=" - ")
    prepared = prepare_actions(document, _auto_profile(), [action])
    assert prepared[0].status == ActionStatus.CONFLICT

    reviewed_path = render_reviewed_docx(document, prepared, tmp_path / "reviewed.docx")
    corrected_path = render_corrected_docx(document, prepared, tmp_path / "corrected.docx")

    assert DocxDocument(str(corrected_path)).paragraphs[0].text == "A\tB"
    assert any(text.startswith("CONFLICT:") for text in _comment_texts(reviewed_path))


def test_hyperlink_edit_demotes_to_conflict_and_both_artifacts_render(
    tmp_path: Path,
) -> None:
    # paragraph.text includes hyperlink text, so validation anchors an APPLIED delete
    # inside it - but the hyperlink is opaque to both renderers. This used to raise
    # RenderIntegrityError and abort the whole run with no artifacts; the fail-closed
    # outcome is a CONFLICT at prepare: both artifacts render, the hyperlink survives,
    # and the edit surfaces as a CONFLICT-labelled comment.
    document = load_docx(_hyperlink_docx(tmp_path))
    paragraph = next(document.iter_paragraphs())
    assert paragraph.text == "See the appendix for details."
    assert paragraph.opaque_ranges == [(4, 16)]

    action = _safe_edit(
        action_type=ReviewActionType.DELETE_TEXT, original_text="the appendix"
    )
    prepared = prepare_actions(document, _auto_profile(), [action])
    assert prepared[0].status == ActionStatus.CONFLICT

    reviewed_path = render_reviewed_docx(document, prepared, tmp_path / "reviewed.docx")
    corrected_path = render_corrected_docx(document, prepared, tmp_path / "corrected.docx")

    corrected = DocxDocument(str(corrected_path))
    assert corrected.paragraphs[0].text == "See the appendix for details."
    assert "<w:hyperlink" in _part_xml(corrected_path, "word/document.xml")
    assert any(text.startswith("CONFLICT:") for text in _comment_texts(reviewed_path))


def test_applied_insert_after_abutting_applied_replace_renders_in_both_artifacts(
    tmp_path: Path,
) -> None:
    # Abutting zero-width spans are deliberately compatible in prepare, so both actions
    # arrive APPLIED. If the tracked replace ran first it would consume the insert's
    # find-based anchor and abort the run; zero-width inserts must apply first so both
    # edits land, identically in reviewed (accepted) and corrected output.
    document = load_docx(_saved_docx(tmp_path, "The foo bar."))
    actions = [
        _safe_edit(id="a-replace", original_text="foo", replacement_text="qux"),
        _safe_edit(
            id="a-insert",
            action_type=ReviewActionType.INSERT_AFTER,
            original_text="foo",
            replacement_text=" indeed",
        ),
    ]
    prepared = prepare_actions(document, _auto_profile(), actions)
    assert [action.status for action in prepared] == [
        ActionStatus.APPLIED,
        ActionStatus.APPLIED,
    ]

    reviewed_path = render_reviewed_docx(document, prepared, tmp_path / "reviewed.docx")
    corrected_path = render_corrected_docx(document, prepared, tmp_path / "corrected.docx")

    reviewed_xml = _part_xml(reviewed_path, "word/document.xml")
    assert _accepted_paragraph_text(reviewed_xml) == "The qux indeed bar."
    assert DocxDocument(str(corrected_path)).paragraphs[0].text == "The qux indeed bar."


def test_applied_delete_tracks_before_listed_overlapping_suggestion(
    tmp_path: Path,
) -> None:
    # A non-APPLIED suggestion listed BEFORE an APPLIED delete it overlaps: the
    # APPLIED edit must always land as a tracked change (it is what corrected.docx
    # applies), so it applies first; the suggestion whose anchor it consumed keeps
    # the documented degrade to a labelled comment. Historically the suggestion
    # tracked first and the APPLIED delete aborted the run.
    document = load_docx(_saved_docx(tmp_path, "The quick brown fox jumps."))
    suggestion = ReviewAction(
        id="a-suggest",
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.REPLACE_TEXT,
        node_id="p1",
        original_text="brown fox",
        replacement_text="red cat",
        category="style",  # not in apply_policy -> stays non-APPLIED
        confidence=0.9,
    )
    applied = _safe_edit(
        id="a-applied", action_type=ReviewActionType.DELETE_TEXT, original_text="fox"
    )
    prepared = prepare_actions(document, _auto_profile(), [suggestion, applied])
    assert prepared[0].status != ActionStatus.APPLIED
    assert prepared[1].status == ActionStatus.APPLIED

    reviewed_path = render_reviewed_docx(document, prepared, tmp_path / "reviewed.docx")
    corrected_path = render_corrected_docx(document, prepared, tmp_path / "corrected.docx")

    reviewed_xml = _part_xml(reviewed_path, "word/document.xml")
    assert _revision_texts(reviewed_xml, "del", "delText") == ["fox"]
    assert any("Original: 'brown fox'" in text for text in _comment_texts(reviewed_path))
    assert DocxDocument(str(corrected_path)).paragraphs[0].text == "The quick brown  jumps."


# --- new_paragraph: stand-alone clause inserts render as a NEW tracked paragraph ---


def _body_paragraphs(xml: str) -> list[ElementTree.Element]:
    root = ElementTree.fromstring(xml)
    body = root.find(f"{_W}body")
    assert body is not None
    return [child for child in body if child.tag == f"{_W}p"]


def _paragraph_visible_text(paragraph: ElementTree.Element) -> str:
    # Text with every insertion accepted: w:t survives, w:delText does not.
    return "".join(text.text or "" for text in paragraph.iter(f"{_W}t"))


def _two_paragraph_document(tmp_path: Path) -> ReviewDocument:
    input_path = tmp_path / "input.docx"
    docx = DocxDocument()
    docx.add_paragraph("Anchor clause heading.")
    docx.add_paragraph("Following paragraph.")
    docx.save(input_path)
    return load_docx(input_path)


def test_new_paragraph_insert_after_renders_a_tracked_sibling_paragraph(tmp_path: Path) -> None:
    # The linchpin: a stand-alone clause insert must become a NEW paragraph between the
    # anchor and the next one - never glued inline onto the anchor's text - with both a
    # paragraph-mark insertion (pPr/rPr/w:ins) and run-level w:ins so accepting the
    # markup yields a real stand-alone paragraph.
    document = _two_paragraph_document(tmp_path)
    anchor = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.INSERT_AFTER,
        node_id=anchor.id,
        replacement_text="§20a. The inserted clause.",
        new_paragraph=True,
        status=ActionStatus.APPLIED,
        apply_to_corrected=True,
    )

    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    paragraphs = _body_paragraphs(_part_xml(reviewed_path, "word/document.xml"))

    # A third body paragraph appeared, spliced between the anchor and the following one.
    assert len(paragraphs) == 3
    assert _paragraph_visible_text(paragraphs[0]) == "Anchor clause heading."
    assert _paragraph_visible_text(paragraphs[2]) == "Following paragraph."

    inserted = paragraphs[1]
    # The anchor paragraph is untouched - the clause is NOT glued onto its text.
    assert _paragraph_visible_text(paragraphs[0]) == "Anchor clause heading."
    # The new paragraph's mark is an insertion: pPr/rPr/w:ins.
    mark = inserted.find(f"{_W}pPr/{_W}rPr/{_W}ins")
    assert mark is not None
    # Its content is a run-level insertion carrying exactly the clause text.
    run_ins = [ins for ins in inserted.findall(f"{_W}ins") if ins.find(f".//{_W}t") is not None]
    assert [ins.find(f".//{_W}t").text for ins in run_ins] == ["§20a. The inserted clause."]


def test_new_paragraph_insert_before_places_clause_ahead_of_anchor(tmp_path: Path) -> None:
    document = _two_paragraph_document(tmp_path)
    anchor = document.sections[0].paragraphs[1]  # the SECOND paragraph
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.INSERT_BEFORE,
        node_id=anchor.id,
        replacement_text="Preamble clause.",
        new_paragraph=True,
        status=ActionStatus.APPLIED,
    )

    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    paragraphs = _body_paragraphs(_part_xml(reviewed_path, "word/document.xml"))

    assert [_paragraph_visible_text(p) for p in paragraphs] == [
        "Anchor clause heading.",
        "Preamble clause.",
        "Following paragraph.",
    ]


def test_new_paragraph_insert_splits_multiline_clause_into_paragraphs(tmp_path: Path) -> None:
    document = _two_paragraph_document(tmp_path)
    anchor = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.INSERT_AFTER,
        node_id=anchor.id,
        replacement_text="First clause line.\nSecond clause line.\n",  # trailing newline dropped
        new_paragraph=True,
        status=ActionStatus.APPLIED,
    )

    reviewed_path = render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
    paragraphs = _body_paragraphs(_part_xml(reviewed_path, "word/document.xml"))

    assert [_paragraph_visible_text(p) for p in paragraphs] == [
        "Anchor clause heading.",
        "First clause line.",
        "Second clause line.",
        "Following paragraph.",
    ]


def test_new_paragraph_flag_ignored_without_replacement_text_raises(tmp_path: Path) -> None:
    # A stand-alone insert with no text to insert is an internal inconsistency, not a
    # silent no-op.
    document = _two_paragraph_document(tmp_path)
    anchor = document.sections[0].paragraphs[0]
    action = ReviewAction(
        scope=ReviewScope.PARAGRAPH,
        action_type=ReviewActionType.INSERT_AFTER,
        node_id=anchor.id,
        replacement_text=None,
        new_paragraph=True,
        status=ActionStatus.APPLIED,
    )

    with pytest.raises(RenderIntegrityError):
        render_reviewed_docx(document, [action], tmp_path / "reviewed.docx")
