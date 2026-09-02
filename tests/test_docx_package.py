"""Unit coverage for the deterministic-packaging primitive (reviewkit.docx_package)."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from lxml import etree

from reviewkit.docx_package import (
    normalize_docx_timestamps,
    restore_semantically_unchanged_xml_parts,
)
from reviewkit.parser_docx import load_docx
from reviewkit.renderer_docx import render_corrected_docx, render_reviewed_docx

_CUSTOM_XML_PART = "customXml/item1.xml"
_RELATIVE_NAMESPACE_XML = (
    b'<properties xmlns="http://schemas.microsoft.com/office/2006/metadata/properties" '
    b'xmlns:field="9a9b9e37-d08a-4d70-9c0a-e2c426ef4e4c">'
    b"<field:value>stable</field:value></properties>"
)
_EQUIVALENT_RELATIVE_NAMESPACE_XML = (
    b"<properties xmlns='http://schemas.microsoft.com/office/2006/metadata/properties' "
    b"xmlns:field='9a9b9e37-d08a-4d70-9c0a-e2c426ef4e4c'>"
    b"<field:value>stable</field:value></properties>"
)
_CHANGED_RELATIVE_NAMESPACE_XML = _RELATIVE_NAMESPACE_XML.replace(b"stable", b"changed")


def _write_docx(path: Path) -> Path:
    docx = DocxDocument()
    docx.add_paragraph("Alpha beta gamma.")
    docx.save(str(path))
    return path


def _write_xml_package(path: Path, payload: bytes) -> Path:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(_CUSTOM_XML_PART, payload)
    return path


def _part_bytes(path: Path, member: str = _CUSTOM_XML_PART) -> bytes:
    with zipfile.ZipFile(path) as archive:
        return archive.read(member)


def _write_relative_namespace_docx(tmp_path: Path) -> Path:
    generated = tmp_path / "generated.docx"
    source = tmp_path / "source.docx"
    docx = DocxDocument()
    docx.add_paragraph("First item", style="List Number")
    docx.add_paragraph("Second item", style="List Number")
    docx.save(str(generated))
    with (
        zipfile.ZipFile(generated) as incoming,
        zipfile.ZipFile(source, "w", zipfile.ZIP_DEFLATED) as outgoing,
    ):
        for info in incoming.infolist():
            payload = incoming.read(info.filename)
            if info.filename == "word/numbering.xml":
                payload = payload.replace(
                    b"<w:numbering ",
                    b'<w:numbering xmlns:field="9a9b9e37-d08a-4d70-9c0a-e2c426ef4e4c" ',
                    1,
                )
            outgoing.writestr(info, payload)
    return source


def test_restore_accepts_byte_identical_xml_with_relative_namespace(tmp_path: Path) -> None:
    # Given a valid Office custom XML part whose namespace is a bare UUID.
    source = _write_xml_package(tmp_path / "source.docx", _RELATIVE_NAMESPACE_XML)
    rendered = _write_xml_package(tmp_path / "rendered.docx", _RELATIVE_NAMESPACE_XML)

    # When the unchanged part is inspected for source-byte restoration.
    restore_semantically_unchanged_xml_parts(source, rendered)

    # Then byte equality is enough and unsupported C14N 1.0 is never required.
    assert _part_bytes(rendered) == _RELATIVE_NAMESPACE_XML


def test_restore_preserves_source_bytes_for_equivalent_relative_namespace_xml(
    tmp_path: Path,
) -> None:
    # Given semantically equivalent custom XML with different serialization.
    source = _write_xml_package(tmp_path / "source.docx", _RELATIVE_NAMESPACE_XML)
    rendered = _write_xml_package(tmp_path / "rendered.docx", _EQUIVALENT_RELATIVE_NAMESPACE_XML)

    # When source-byte restoration compares the two valid relative-namespace parts.
    restore_semantically_unchanged_xml_parts(source, rendered)

    # Then the rendered package carries the exact independently defined source bytes.
    assert _part_bytes(rendered) == _RELATIVE_NAMESPACE_XML


def test_restore_keeps_changed_relative_namespace_xml(tmp_path: Path) -> None:
    # Given custom XML whose content changed under the same relative namespace.
    source = _write_xml_package(tmp_path / "source.docx", _RELATIVE_NAMESPACE_XML)
    rendered = _write_xml_package(tmp_path / "rendered.docx", _CHANGED_RELATIVE_NAMESPACE_XML)

    # When source-byte restoration compares the changed part.
    restore_semantically_unchanged_xml_parts(source, rendered)

    # Then the rendered content is retained instead of being overwritten from source.
    assert _part_bytes(rendered) == _CHANGED_RELATIVE_NAMESPACE_XML


def test_restore_keeps_rendered_xml_when_source_is_malformed(tmp_path: Path) -> None:
    # Given source XML that cannot be parsed and a valid rendered replacement.
    source = _write_xml_package(tmp_path / "source.docx", b"<properties>")
    rendered = _write_xml_package(tmp_path / "rendered.docx", b"<properties />")

    # When source-byte restoration reaches the malformed source part.
    restore_semantically_unchanged_xml_parts(source, rendered)

    # Then failure to compare safely leaves the rendered bytes untouched.
    assert _part_bytes(rendered) == b"<properties />"


def test_restore_keeps_rendered_xml_when_canonicalization_is_unsupported(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # Given different XML bytes and a canonicalizer that rejects the valid tree.
    source = _write_xml_package(tmp_path / "source.docx", b"<properties />")
    rendered = _write_xml_package(tmp_path / "rendered.docx", b"<properties></properties>")

    def reject_canonicalization(
        _root: etree._Element,
        *,
        method: str,
        with_comments: bool,
    ) -> bytes:
        raise etree.C14NError(f"unsupported {method=}, {with_comments=}")

    monkeypatch.setattr(etree, "tostring", reject_canonicalization)

    # When source-byte restoration cannot canonicalize either representation.
    restore_semantically_unchanged_xml_parts(source, rendered)

    # Then safe non-restoration retains the rendered bytes.
    assert _part_bytes(rendered) == b"<properties></properties>"


@pytest.mark.parametrize("renderer", [render_reviewed_docx, render_corrected_docx])
def test_render_paths_preserve_relative_namespace_part(tmp_path: Path, renderer) -> None:
    # Given a source-backed DOCX part using a bare UUID namespace.
    source = _write_relative_namespace_docx(tmp_path)
    source_numbering = _part_bytes(source, "word/numbering.xml")

    # When either public renderer saves the document through python-docx.
    rendered = renderer(load_docx(source), [], tmp_path / "rendered.docx")

    # Then ReviewKit restores the exact source bytes after semantic comparison.
    assert _part_bytes(rendered, "word/numbering.xml") == source_numbering


def test_normalize_pins_every_entry_timestamp_to_the_zip_epoch(tmp_path: Path) -> None:
    path = _write_docx(tmp_path / "doc.docx")
    # python-docx stamps the wall clock, so pre-normalization entries are NOT the epoch.
    with zipfile.ZipFile(path) as archive:
        assert any(info.date_time != (1980, 1, 1, 0, 0, 0) for info in archive.infolist())

    normalize_docx_timestamps(path)

    with zipfile.ZipFile(path) as archive:
        assert all(info.date_time == (1980, 1, 1, 0, 0, 0) for info in archive.infolist())


def test_normalize_preserves_names_order_and_content(tmp_path: Path) -> None:
    path = _write_docx(tmp_path / "doc.docx")
    with zipfile.ZipFile(path) as archive:
        before = {info.filename: archive.read(info.filename) for info in archive.infolist()}
        order_before = [info.filename for info in archive.infolist()]

    normalize_docx_timestamps(path)

    with zipfile.ZipFile(path) as archive:
        after = {info.filename: archive.read(info.filename) for info in archive.infolist()}
        order_after = [info.filename for info in archive.infolist()]
    assert order_after == order_before
    assert after == before
    # The package still opens as a valid DOCX after the rewrite.
    assert DocxDocument(str(path)).paragraphs[0].text == "Alpha beta gamma."


def test_normalize_is_idempotent(tmp_path: Path) -> None:
    path = _write_docx(tmp_path / "doc.docx")
    normalize_docx_timestamps(path)
    once = path.read_bytes()
    normalize_docx_timestamps(path)
    assert path.read_bytes() == once
