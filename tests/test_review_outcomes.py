from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from lxml import etree

from reviewkit import (
    incorporated_comment_outcomes,
    measure_review_changes,
    read_metadata_marker,
    set_metadata_marker,
    strip_metadata_marker,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx() -> bytes:
    document = Document()
    paragraph = document.add_paragraph("text")
    document.comments.add_comment(text="accepted remark", author="A")
    paragraph._p.insert(0, etree.Element(f"{{{W}}}commentRangeStart", {f"{{{W}}}id": "0"}))
    paragraph._p.append(etree.Element(f"{{{W}}}commentRangeEnd", {f"{{{W}}}id": "0"}))
    ref = etree.SubElement(paragraph._p, f"{{{W}}}r")
    etree.SubElement(ref, f"{{{W}}}commentReference", {f"{{{W}}}id": "0"})
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _with_revision(data: bytes) -> bytes:
    output = BytesIO()
    with ZipFile(BytesIO(data)) as source, ZipFile(output, "w", ZIP_DEFLATED) as target:
        for info in source.infolist():
            payload = source.read(info)
            if info.filename == "word/document.xml":
                root = etree.fromstring(payload)
                paragraph = root.find(f".//{{{W}}}p")
                assert paragraph is not None
                start = paragraph.find(f"{{{W}}}commentRangeStart")
                assert start is not None
                run = next(item for item in paragraph if etree.QName(item).localname == "r")
                paragraph.remove(run)
                insertion = etree.Element(f"{{{W}}}ins")
                insertion.set(f"{{{W}}}id", "1")
                insertion.append(run)
                paragraph.insert(paragraph.index(start) + 1, insertion)
                payload = etree.tostring(
                    root, xml_declaration=True, encoding="UTF-8", standalone=True
                )
            target.writestr(info, payload)
    return output.getvalue()


def test_incorporated_comment_outcomes_use_typed_docxtor_association() -> None:
    outcomes = incorporated_comment_outcomes(_with_revision(_docx()))
    assert [(item.comment_id, item.text, item.revision_kinds) for item in outcomes] == [
        ("0", "accepted remark", ("ins",))
    ]


def test_metadata_marker_and_change_metrics(tmp_path: Path) -> None:
    original = _docx()
    marked = set_metadata_marker(original, prefix="review=", value="abc")
    assert read_metadata_marker(marked, prefix="review=") == "abc"
    stripped = strip_metadata_marker(marked, prefix="review=")
    assert read_metadata_marker(stripped, prefix="review=") is None
    before = tmp_path / "before.docx"
    after = tmp_path / "after.docx"
    before.write_bytes(original)
    after.write_bytes(stripped)
    metrics = measure_review_changes(before, after)
    assert metrics.before_paragraphs == metrics.after_paragraphs
    assert metrics.before_text_sha256 == metrics.after_text_sha256


def test_change_metrics_separate_text_from_block_order(tmp_path: Path) -> None:
    before = tmp_path / "before.docx"
    after = tmp_path / "after.docx"
    left = Document()
    left.add_paragraph("body")
    left.add_table(rows=1, cols=1).cell(0, 0).text = "cell"
    left.save(before)
    right = Document()
    right.add_table(rows=1, cols=1).cell(0, 0).text = "cell"
    right.add_paragraph("body")
    right.save(after)
    metrics = measure_review_changes(before, after)
    assert metrics.before_text_sha256 == metrics.after_text_sha256
    assert metrics.before_structure_sha256 != metrics.after_structure_sha256
